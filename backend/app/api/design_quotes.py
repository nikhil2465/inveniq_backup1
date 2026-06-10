"""
Design Quote Studio API — Interior & Architect Quotations.
Accessible only to the 'architect' role (module: designquote).
DB-first / demo-fallback pattern. Tables auto-created on first DB call.
"""
import io
import json
import logging
import base64
import math
import os
import smtplib
import asyncio
from datetime import datetime, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from typing import List, Optional

from fastapi import APIRouter, HTTPException, UploadFile, File, Form, Request

logger = logging.getLogger(__name__)
router = APIRouter(tags=["Design Quotes"])

# ── Table DDL (created on first DB call, idempotent) ─────────────────────────

_DESIGN_QUOTES_DDL = """
CREATE TABLE IF NOT EXISTS design_quotes (
    id              BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    quote_number    VARCHAR(50)     NOT NULL,
    client_name     VARCHAR(255)    NOT NULL,
    client_phone    VARCHAR(50)     DEFAULT '',
    client_email    VARCHAR(150)    DEFAULT '',
    project_name    VARCHAR(255)    DEFAULT '',
    project_address TEXT            DEFAULT '',
    project_type    VARCHAR(100)    DEFAULT 'Residential',
    designer_name   VARCHAR(255)    DEFAULT '',
    designer_company VARCHAR(255)   DEFAULT '',
    payment_terms   TEXT            DEFAULT '',
    validity_days   INT             DEFAULT 30,
    gst_rate        DECIMAL(5,2)    DEFAULT 18,
    include_gst     TINYINT(1)      DEFAULT 1,
    notes           TEXT            DEFAULT '',
    terms           TEXT            DEFAULT '',
    status          ENUM('DRAFT','SENT','APPROVED','REVISION','IN_PROGRESS','COMPLETED','CANCELLED')
                    DEFAULT 'DRAFT',
    subtotal        DECIMAL(14,2)   DEFAULT 0,
    gst_amount      DECIMAL(14,2)   DEFAULT 0,
    grand_total     DECIMAL(14,2)   DEFAULT 0,
    total_area_sqft DECIMAL(12,2)   DEFAULT 0,
    sections_json   LONGTEXT        DEFAULT '[]',
    margin_mode     VARCHAR(20)     DEFAULT 'per_line',
    overall_margin_pct DECIMAL(8,4) DEFAULT 0,
    created_at      DATETIME        DEFAULT CURRENT_TIMESTAMP,
    updated_at      DATETIME        DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
    INDEX idx_status       (status),
    INDEX idx_client       (client_name(50)),
    INDEX idx_quote_number (quote_number)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;
"""

_ARCHITECT_PROPOSALS_DDL = """
CREATE TABLE IF NOT EXISTS architect_proposals (
    id                  BIGINT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    proposal_number     VARCHAR(50)     NOT NULL,
    client_name         VARCHAR(255)    NOT NULL,
    client_phone        VARCHAR(50)     DEFAULT '',
    client_email        VARCHAR(150)    DEFAULT '',
    project_name        VARCHAR(255)    DEFAULT '',
    project_type        VARCHAR(50)     DEFAULT 'residential',
    typology            VARCHAR(100)    DEFAULT 'villa',
    plot_length         DECIMAL(10,2)   DEFAULT 0,
    plot_width          DECIMAL(10,2)   DEFAULT 0,
    plot_unit           VARCHAR(10)     DEFAULT 'feet',
    site_area_sqft      DECIMAL(12,2)   DEFAULT 0,
    floors              INT             DEFAULT 1,
    builtup_area_sqft   DECIMAL(12,2)   DEFAULT 0,
    carpet_area_sqft    DECIMAL(12,2)   DEFAULT 0,
    fee_model           VARCHAR(20)     DEFAULT 'percentage',
    fee_rate            DECIMAL(8,4)    DEFAULT 5.0,
    construction_cost   DECIMAL(14,2)   DEFAULT 0,
    total_fee           DECIMAL(12,2)   DEFAULT 0,
    gst_pct             DECIMAL(5,2)    DEFAULT 18,
    validity_days       INT             DEFAULT 30,
    notes               TEXT            DEFAULT '',
    status              ENUM('DRAFT','SENT','APPROVED','REVISION','COMPLETED','CANCELLED')
                        DEFAULT 'DRAFT',
    phases_json         LONGTEXT        DEFAULT '[]',
    boq_json            LONGTEXT        DEFAULT '[]',
    created_at          DATETIME        DEFAULT CURRENT_TIMESTAMP,
    updated_at          DATETIME        DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
    INDEX idx_status   (status),
    INDEX idx_client   (client_name(50))
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;
"""

_tables_initialized = False


async def _ensure_tables(pool) -> None:
    global _tables_initialized
    if _tables_initialized:
        return
    try:
        async with pool.acquire() as conn:
            async with conn.cursor() as cur:
                await cur.execute(_DESIGN_QUOTES_DDL)
                await cur.execute(_ARCHITECT_PROPOSALS_DDL)
                # Additive migrations — safe to run on existing installs; silently pass if column exists
                for stmt in (
                    "ALTER TABLE design_quotes ADD COLUMN margin_mode VARCHAR(20) DEFAULT 'per_line'",
                    "ALTER TABLE design_quotes ADD COLUMN overall_margin_pct DECIMAL(8,4) DEFAULT 0",
                ):
                    try:
                        await cur.execute(stmt)
                    except Exception:
                        pass  # column already exists (error 1060)
        _tables_initialized = True
        logger.info("design_quotes: tables verified OK")
    except Exception as exc:
        logger.debug("design_quotes: table init skipped — %s", exc)


async def _get_db_pool():
    """Return an initialised DB pool or None on any error.

    Combines pool acquisition + table migration in one call so every endpoint
    only needs `pool = await _get_db_pool()` — no repeated boilerplate.
    _ensure_tables is idempotent (guarded by _tables_initialized) so the
    overhead after the first call is a single boolean check.
    """
    try:
        from app.db.connection import get_pool as _get_pool
        pool = await _get_pool()
        if pool:
            await _ensure_tables(pool)
        return pool
    except Exception:
        return None


# ── Room templates & product options (static master data) ────────────────────

ROOM_TEMPLATES = {
    "Kitchen": {
        "icon": "kitchen", "color": "#f59e0b",
        "items": [
            {"item_name": "Base Cabinets",  "item_type": "cabinet",     "dim_type": "L",   "unit": "rft"},
            {"item_name": "Wall Cabinets",  "item_type": "cabinet",     "dim_type": "L",   "unit": "rft"},
            {"item_name": "Countertop",     "item_type": "countertop",  "dim_type": "LxW", "unit": "sqft"},
            {"item_name": "Dado Tiles",     "item_type": "dado_tiles",  "dim_type": "LxW", "unit": "sqft"},
        ],
    },
    "Master Bedroom": {
        "icon": "bed", "color": "#8b5cf6",
        "items": [
            {"item_name": "Wardrobe",       "item_type": "wardrobe",      "dim_type": "L",   "unit": "rft"},
            {"item_name": "False Ceiling",  "item_type": "false_ceiling", "dim_type": "LxW", "unit": "sqft"},
            {"item_name": "Wall Paneling",  "item_type": "wall_panel",    "dim_type": "LxH", "unit": "sqft"},
        ],
    },
    "Living Room": {
        "icon": "sofa", "color": "#3b82f6",
        "items": [
            {"item_name": "TV Unit",        "item_type": "tv_unit",       "dim_type": "L",   "unit": "rft"},
            {"item_name": "False Ceiling",  "item_type": "false_ceiling", "dim_type": "LxW", "unit": "sqft"},
            {"item_name": "Wall Paneling",  "item_type": "wall_panel",    "dim_type": "LxH", "unit": "sqft"},
            {"item_name": "Flooring",       "item_type": "flooring",      "dim_type": "LxW", "unit": "sqft"},
        ],
    },
    "Pooja Room": {
        "icon": "pooja", "color": "#ec4899",
        "items": [
            {"item_name": "Mandir Unit",    "item_type": "mandir",        "dim_type": "LxH", "unit": "sqft"},
            {"item_name": "False Ceiling",  "item_type": "false_ceiling", "dim_type": "LxW", "unit": "sqft"},
        ],
    },
    "Dining Room": {
        "icon": "dining", "color": "#f97316",
        "items": [
            {"item_name": "Crockery Unit",  "item_type": "cabinet",    "dim_type": "L",   "unit": "rft"},
            {"item_name": "Flooring",       "item_type": "flooring",   "dim_type": "LxW", "unit": "sqft"},
        ],
    },
    "Bedroom": {
        "icon": "bed", "color": "#16a34a",
        "items": [
            {"item_name": "Wardrobe",       "item_type": "wardrobe",      "dim_type": "L",   "unit": "rft"},
            {"item_name": "Bed",            "item_type": "bed",           "dim_type": "fixed","unit": "nos"},
            {"item_name": "False Ceiling",  "item_type": "false_ceiling", "dim_type": "LxW", "unit": "sqft"},
            {"item_name": "Wall Paneling",  "item_type": "wall_panel",    "dim_type": "LxH", "unit": "sqft"},
            {"item_name": "Flooring",       "item_type": "flooring",      "dim_type": "LxW", "unit": "sqft"},
        ],
    },
    "Kids Bedroom": {
        "icon": "kids", "color": "#06b6d4",
        "items": [
            {"item_name": "Study Unit",     "item_type": "cabinet",    "dim_type": "L",   "unit": "rft"},
            {"item_name": "Wardrobe",       "item_type": "wardrobe",   "dim_type": "L",   "unit": "rft"},
            {"item_name": "Bed",            "item_type": "bed",        "dim_type": "fixed","unit": "nos"},
        ],
    },
    "Bathroom": {
        "icon": "bath", "color": "#14b8a6",
        "items": [
            {"item_name": "Dado Tiles",     "item_type": "dado_tiles", "dim_type": "LxH", "unit": "sqft"},
            {"item_name": "Vanity Unit",    "item_type": "cabinet",    "dim_type": "L",   "unit": "rft"},
            {"item_name": "CP Accessories", "item_type": "hardware",   "dim_type": "fixed","unit": "set"},
        ],
    },
    "Study Room": {
        "icon": "study", "color": "#6366f1",
        "items": [
            {"item_name": "Study Unit",     "item_type": "cabinet",    "dim_type": "L",   "unit": "rft"},
            {"item_name": "Bookshelf",      "item_type": "cabinet",    "dim_type": "LxH", "unit": "sqft"},
            {"item_name": "Flooring",       "item_type": "flooring",   "dim_type": "LxW", "unit": "sqft"},
        ],
    },
    "Balcony": {
        "icon": "balcony", "color": "#84cc16",
        "items": [
            {"item_name": "Flooring",       "item_type": "flooring",   "dim_type": "LxW", "unit": "sqft"},
            {"item_name": "Planters",       "item_type": "cabinet",    "dim_type": "L",   "unit": "rft"},
        ],
    },
}

PRODUCT_OPTIONS = {
    "cabinet": [
        {"id": "bwp_acrylic",     "name": "BWP Board + Acrylic",         "base_rate": 8500,  "unit": "rft"},
        {"id": "hdhmr_acrylic",   "name": "HDHMR + Acrylic",             "base_rate": 7800,  "unit": "rft"},
        {"id": "bwp_laminate",    "name": "BWP Board + Laminate",        "base_rate": 5200,  "unit": "rft"},
        {"id": "mdf_laminate",    "name": "MDF + Laminate",              "base_rate": 4500,  "unit": "rft"},
        {"id": "hdhmr_pu",        "name": "HDHMR + PU Finish",           "base_rate": 9200,  "unit": "rft"},
        {"id": "bwp_veneer",      "name": "BWP Board + Veneer",          "base_rate": 10500, "unit": "rft"},
    ],
    "wardrobe": [
        {"id": "hdhmr_acrylic_w", "name": "HDHMR + Acrylic",            "base_rate": 6800,  "unit": "rft"},
        {"id": "hdhmr_lami_w",    "name": "HDHMR + Laminate",           "base_rate": 4800,  "unit": "rft"},
        {"id": "hdhmr_pu_w",      "name": "HDHMR + PU Finish",          "base_rate": 9000,  "unit": "rft"},
        {"id": "hdhmr_veneer_w",  "name": "HDHMR + Veneer",             "base_rate": 11000, "unit": "rft"},
        {"id": "mdf_lami_w",      "name": "MDF + Laminate (Budget)",    "base_rate": 3800,  "unit": "rft"},
    ],
    "tv_unit": [
        {"id": "mdf_acrylic_tv",  "name": "MDF + Acrylic",              "base_rate": 5500,  "unit": "rft"},
        {"id": "hdhmr_lami_tv",   "name": "HDHMR + Laminate",           "base_rate": 4200,  "unit": "rft"},
        {"id": "hdhmr_veneer_tv", "name": "HDHMR + Veneer",             "base_rate": 8200,  "unit": "rft"},
    ],
    "false_ceiling": [
        {"id": "gypsum_basic",    "name": "Gypsum Board (Basic)",       "base_rate": 65,    "unit": "sqft"},
        {"id": "gypsum_designer", "name": "Gypsum Board (Designer)",    "base_rate": 95,    "unit": "sqft"},
        {"id": "pop_plain",       "name": "POP Plain",                  "base_rate": 55,    "unit": "sqft"},
        {"id": "pop_designer",    "name": "POP Designer",               "base_rate": 90,    "unit": "sqft"},
        {"id": "acp_grid",        "name": "ACP / Grid Ceiling",         "base_rate": 120,   "unit": "sqft"},
    ],
    "flooring": [
        {"id": "vitrified_600",   "name": "Vitrified Tile 600×600",     "base_rate": 95,    "unit": "sqft"},
        {"id": "marble_italian",  "name": "Italian Marble",             "base_rate": 280,   "unit": "sqft"},
        {"id": "marble_indian",   "name": "Indian Marble",              "base_rate": 150,   "unit": "sqft"},
        {"id": "wood_laminate",   "name": "Wood Laminate",              "base_rate": 110,   "unit": "sqft"},
        {"id": "hardwood",        "name": "Hardwood Flooring",          "base_rate": 350,   "unit": "sqft"},
    ],
    "countertop": [
        {"id": "quartz_import",   "name": "Quartz Imported (per slab)", "base_rate": 18000, "unit": "slab"},
        {"id": "quartz_domestic", "name": "Quartz Domestic (per slab)", "base_rate": 9500,  "unit": "slab"},
        {"id": "granite_polished","name": "Granite Polished",           "base_rate": 120,   "unit": "sqft"},
        {"id": "marble_counter",  "name": "Marble Countertop",          "base_rate": 200,   "unit": "sqft"},
    ],
    "dado_tiles": [
        {"id": "ceramic_std",     "name": "Ceramic (Standard)",         "base_rate": 65,    "unit": "sqft"},
        {"id": "vitrified_dado",  "name": "Vitrified Tile",             "base_rate": 95,    "unit": "sqft"},
        {"id": "designer_dado",   "name": "Designer Ceramic",           "base_rate": 140,   "unit": "sqft"},
    ],
    "wall_panel": [
        {"id": "fabric_panel",    "name": "Fabric Wall Panel",          "base_rate": 180,   "unit": "sqft"},
        {"id": "wallpaper_basic", "name": "Wallpaper (Basic)",          "base_rate": 65,    "unit": "sqft"},
        {"id": "wallpaper_prem",  "name": "Wallpaper (Premium)",        "base_rate": 150,   "unit": "sqft"},
        {"id": "wpc_panel",       "name": "WPC / PVC Panel",            "base_rate": 120,   "unit": "sqft"},
        {"id": "stone_cladding",  "name": "Stone Cladding",             "base_rate": 250,   "unit": "sqft"},
    ],
    "mandir": [
        {"id": "wood_mandir",     "name": "Solid Wood Mandir",          "base_rate": 1800,  "unit": "sqft"},
        {"id": "marble_mandir",   "name": "Marble Mandir",              "base_rate": 2500,  "unit": "sqft"},
        {"id": "mdf_mandir",      "name": "MDF + Carving",              "base_rate": 1200,  "unit": "sqft"},
    ],
    "bed": [
        {"id": "hydraulic_bed",   "name": "Hydraulic Storage Bed",      "base_rate": 35000, "unit": "nos"},
        {"id": "box_bed_king",    "name": "Box Bed (King Size)",         "base_rate": 28000, "unit": "nos"},
        {"id": "box_bed_queen",   "name": "Box Bed (Queen Size)",        "base_rate": 22000, "unit": "nos"},
    ],
    "hardware": [
        {"id": "cp_set_std",      "name": "CP Fitting Set (Standard)",  "base_rate": 8000,  "unit": "set"},
        {"id": "cp_set_premium",  "name": "CP Fitting Set (Premium)",   "base_rate": 18000, "unit": "set"},
    ],
}

# ── Demo in-memory store ──────────────────────────────────────────────────────

_demo_quotes: dict = {}
_demo_proposals: dict = {}
_demo_q_counter = [3]   # next ID
_demo_p_counter = [2]


def _init_demo():
    if _demo_quotes:
        return
    today = datetime.now()
    _demo_quotes[1] = {
        "id": 1, "quote_number": "DQ-2026-001",
        "client_name": "Mr. Rajesh Sharma", "client_phone": "+91 98765 43210",
        "client_email": "rajesh@example.com",
        "project_name": "3BHK Flat — Koramangala", "project_type": "Residential",
        "project_address": "Block 5, Koramangala, Bangalore",
        "designer_name": "Architect Studio", "designer_company": "Arc Design Co.",
        "payment_terms": "50% Advance · 25% at Mid Stage · 25% on Completion",
        "validity_days": 30, "gst_rate": 18, "include_gst": True,
        "notes": "Premium acrylic finish preferred for kitchen.",
        "terms": "1. Design includes supply & installation.\n2. Structural changes billed separately.\n3. Civil works not included unless stated.\n4. 1-year workmanship warranty.",
        "status": "SENT",
        "subtotal": 485000, "gst_amount": 87300, "grand_total": 572300,
        "total_area_sqft": 1850,
        "margin_mode": "per_line", "overall_margin_pct": 0,
        "sections": [
            {"section_name": "Kitchen", "section_order": 0, "section_total": 165000, "items": []},
            {"section_name": "Master Bedroom", "section_order": 1, "section_total": 185000, "items": []},
            {"section_name": "Living Room", "section_order": 2, "section_total": 135000, "items": []},
        ],
        "created_at": (today - timedelta(days=11)).strftime("%Y-%m-%d"),
        "valid_till": (today - timedelta(days=11) + timedelta(days=30)).strftime("%Y-%m-%d"),
    }
    _demo_quotes[2] = {
        "id": 2, "quote_number": "DQ-2026-002",
        "client_name": "Ms. Priya Mehta", "client_phone": "+91 87654 32109",
        "client_email": "priya@example.com",
        "project_name": "Villa Interior — Whitefield", "project_type": "Residential",
        "project_address": "Palm Meadows, Whitefield, Bangalore",
        "designer_name": "Architect Studio", "designer_company": "Arc Design Co.",
        "payment_terms": "40% Advance · 30% at Mid Stage · 30% on Completion",
        "validity_days": 45, "gst_rate": 18, "include_gst": True,
        "notes": "Veneer finish for master bedroom wardrobe.",
        "terms": "1. Design includes supply & installation.\n2. Structural changes billed separately.\n3. 1-year workmanship warranty.",
        "status": "APPROVED",
        "subtotal": 1240000, "gst_amount": 223200, "grand_total": 1463200,
        "total_area_sqft": 4200,
        "margin_mode": "overall", "overall_margin_pct": 18.5,
        "sections": [
            {"section_name": "Kitchen", "section_order": 0, "section_total": 280000, "items": []},
            {"section_name": "Master Bedroom", "section_order": 1, "section_total": 320000, "items": []},
            {"section_name": "Living Room", "section_order": 2, "section_total": 420000, "items": []},
            {"section_name": "Kids Bedroom", "section_order": 3, "section_total": 220000, "items": []},
        ],
        "margin_mode": "overall", "overall_margin_pct": 22.0,
        "created_at": (today - timedelta(days=36)).strftime("%Y-%m-%d"),
        "valid_till": (today - timedelta(days=36) + timedelta(days=45)).strftime("%Y-%m-%d"),
    }
    _demo_proposals[1] = {
        "id": 1, "proposal_number": "FP-2026-001",
        "client_name": "Prestige Developers", "client_phone": "+91 76543 21098",
        "client_email": "prestige@example.com",
        "project_name": "G+3 Residential Complex — Sarjapur",
        "project_type": "residential", "typology": "apartment",
        "plot_length": 100, "plot_width": 120, "plot_unit": "feet",
        "site_area_sqft": 12000, "floors": 4,
        "builtup_area_sqft": 28800, "carpet_area_sqft": 20160,
        "fee_model": "percentage", "fee_rate": 5.5,
        "construction_cost": 14400000, "total_fee": 792000,
        "gst_pct": 18, "validity_days": 30, "notes": "RERA registration required.",
        "status": "SENT",
        "phases": [
            {"phase_name": "P1 Concept Design",    "pct_of_total": 10, "fee_amount": 79200,  "is_paid": True,  "due_date": "2026-06-15"},
            {"phase_name": "P2 Schematic Design",  "pct_of_total": 15, "fee_amount": 118800, "is_paid": False, "due_date": "2026-07-15"},
            {"phase_name": "P3 Design Development","pct_of_total": 20, "fee_amount": 158400, "is_paid": False, "due_date": "2026-08-30"},
            {"phase_name": "P4 Construction Docs", "pct_of_total": 25, "fee_amount": 198000, "is_paid": False, "due_date": "2026-10-15"},
            {"phase_name": "P5 Approvals",         "pct_of_total": 5,  "fee_amount": 39600,  "is_paid": False, "due_date": "2026-11-01"},
            {"phase_name": "P6 Site Supervision",  "pct_of_total": 25, "fee_amount": 198000, "is_paid": False, "due_date": "2027-03-31"},
        ],
        "boq": [],
        "created_at": (today - timedelta(days=25)).strftime("%Y-%m-%d"),
        "valid_till": (today + timedelta(days=5)).strftime("%Y-%m-%d"),
    }


def _next_quote_number() -> str:
    yr = datetime.now().year
    n = _demo_q_counter[0]
    _demo_q_counter[0] += 1
    return f"DQ-{yr}-{n:03d}"


def _next_proposal_number() -> str:
    yr = datetime.now().year
    n = _demo_p_counter[0]
    _demo_p_counter[0] += 1
    return f"FP-{yr}-{n:03d}"


def _valid_till(days: int) -> str:
    return (datetime.now() + timedelta(days=days)).strftime("%Y-%m-%d")


# ── AI prompts ────────────────────────────────────────────────────────────────

_INTERIOR_SCAN_PROMPT = """You are an expert estimator for the Indian hardware, sanitary fittings, CP fittings, and interior fit-out industry. Parse the client's description or extract from images into structured JSON covering room-by-room requirements.
Always respond with valid JSON only — no preamble, no explanation.

INDUSTRY KNOWLEDGE:
- CP FITTINGS: basin mixer, shower mixer, overhead shower, hand shower, health faucet, stop cock, bath spout, kitchen mixer. Brands: Jaquar, Grohe, Kohler, Hindware, Parryware, Cera, Essco. Unit = Nos or Set.
- SANITARY WARE: WC/EWC (one-piece, two-piece, wall-hung), wash basin (pedestal, counter-top, wall-hung), bathtub (free-standing, built-in), urinal. Brands: Kohler, Hindware, Jaquar, Parryware, Cera. Unit = Nos or Set.
- BATHROOM ACCESSORIES: towel bar/ring/hook, soap dispenser, toilet paper holder, mirror (plain/LED), floor drain, shower enclosure/glass partition. Unit = Nos or Set.
- HARDWARE: hinges (soft-close/clip-top), channels/drawer slides (soft-close/tandem), handles/knobs (size in mm: 96/128/160/320mm), locks/latches, door closers, floor springs. Brands: Hettich, Blum, Hafele, Ebco, Godrej. Unit = Nos or Pair or Set.
- PLUMBING: CPVC pipe 15/20/25mm, PVC drainage 2"/4", ball valves, GI fittings. Brands: Astral, Supreme, Finolex, Wavin. Unit = Mtr or Lot.
- TILES: ceramic/vitrified/PGVT/marble/granite in SqFt or SqMtr. Brands: Kajaria, Somany, Johnson, RAK. dim_type = LxW.
- WATERPROOFING: bathroom/terrace waterproofing. Unit = SqFt or Lot.
- INSTALLATION: plumbing per bathroom, complete fitting charges. Unit = Nos (per bathroom) or Lot.

Return this exact schema:
{
  "client_name": "string or null",
  "client_phone": "string or null",
  "client_email": "string or null",
  "project_name": "string or null",
  "project_address": "string or null",
  "project_type": "Residential|Commercial|Hospitality|Office|Industrial|Other",
  "designer_name": "string or null",
  "no_of_units": number_or_null,
  "no_of_bathrooms_per_unit": number_or_null,
  "budget_indication": "string or null",
  "notes": "string or null",
  "rooms": [
    {
      "room_name": "Master Bathroom|Guest Bathroom|Common Bathroom|Kids Bathroom|Kitchen|Utility / Wash Area|Balcony|Common Area|<custom>",
      "items": [
        {
          "item_name": "full brand + model + type (e.g. Jaquar Florentine Basin Mixer Chrome)",
          "description": "complete description with specs",
          "item_type": "cp_fittings|sanitary_ware|bathroom_accessories|hardware_hinges|hardware_channels|hardware_handles|hardware_locks|plumbing|tiles|waterproofing|installation|countertop|flooring|other",
          "length_ft": number_or_null,
          "width_ft": number_or_null,
          "height_ft": number_or_null,
          "dim_type": "LxW|LxH|L|W|fixed",
          "unit": "Nos|Set|Pair|Mtr|SqFt|SqMtr|Lot|RFT",
          "qty": number_or_null,
          "unit_price": number,
          "specifications": "finish, color, size, model number, HSN code if known",
          "material_preference": "brand preference or tier: Economy|Standard|Premium|Luxury",
          "inferred_hsn": "8481|6910|3922|8302|8301|3917|6907|3214|9954|null"
        }
      ]
    }
  ]
}

Rules:
- Always infer room_name for bathroom items (e.g. "master bath" → Master Bathroom)
- If units/flats are mentioned (e.g. "24 flats, 3 bathrooms each") populate no_of_units and no_of_bathrooms_per_unit
- CP fittings: dim_type = fixed, unit = Nos or Set
- Sanitary ware: dim_type = fixed, unit = Nos or Set
- Tiles/flooring/waterproofing: dim_type = LxW, unit = SqFt
- Plumbing complete lot: dim_type = fixed, unit = Lot
- If no dimensions given: all null and dim_type = fixed
- If qty is mentioned inline (e.g. "3 sets"), capture it in qty
- Keep all brand names exactly as mentioned

DOCUMENT / BOQ PARSING RULES (for PDF, Word, Excel, or structured text input):
- If input is a BOQ table or schedule, read each row as a separate item
- Extract Length×Width, Length×Height, or single dimension from any column labelled: L, W, H, Length, Width, Height, Size, Dim, Dimension, Area, Room Size
- Convert mm → ft by dividing by 304.8; cm → ft by dividing by 30.48; m → ft by multiplying by 3.281
- Column labelled "Qty", "Quantity", "Nos", "Units" → populate qty field
- Column labelled "Rate", "Unit Price", "MRP", "Price" → populate unit_price field
- Column labelled "Description", "Particulars", "Item" → populate item_name and description
- Column labelled "HSN", "HSN Code", "SAC" → populate inferred_hsn
- Group rows by room/area/section headers found in the document
- If client name, project name, or address is visible in the document header → populate those fields

UNIVERSAL EXTRACTION RULES — CRITICAL — APPLY TO ALL INPUT TYPES:
- NEVER return an empty rooms array. If no specific product names or SKUs are found, extract descriptive requirements as items with item_type="other".
- For ANY space/area/room mentioned (kitchen, bedroom, bathroom, living room, dining, hall, office, terrace, balcony, study, gym, store room, etc.) → create a room entry with that name
- For ANY work, task, material, finish, or scope described → create an item under the relevant room:
  * "Modular kitchen with grey shutters" → item_name="Modular Kitchen — Grey Shutter Finish", item_type="other", unit="Set"
  * "Granite / marble countertop" → item_name="Granite Countertop", item_type="countertop", unit="SqFt"
  * "Oak / teak / laminate flooring" → item_name="<Material> Flooring", item_type="flooring", unit="SqFt"
  * "False ceiling / gypsum / POP" → item_name="False Ceiling Works", item_type="other", unit="SqFt"
  * "Paint / texture / wallpaper" → item_name="Interior Painting Works", item_type="other", unit="SqFt"
  * "Electrical / lighting / wiring" → item_name="Electrical & Lighting Works", item_type="other", unit="Lot"
  * "Renovation / fitout / interior design" (vague) → item_name="General Interior Works", item_type="other", unit="Lot"
  * "Wardrobe / storage / shelving" → item_name="Custom Wardrobe / Storage Unit", item_type="other", unit="Nos"
  * "Door / window / glazing / shutter" → item_name="Door / Window Works", item_type="other", unit="Nos"
  * Any brand or material mentioned without specific product → still create an item using material + application as item_name
- DIMENSION EXTRACTION — mandatory; find and convert dimensions from ANY format anywhere in the text:
  * "10×12" or "10 by 12 feet" or "10ft x 12ft" or "10' x 12'" → length_ft=10, width_ft=12, dim_type="LxW"
  * "100 sqft" or "100 sq.ft" or "100 sq ft" → length_ft=10.0, width_ft=10.0, dim_type="LxW" (use sqrt approximation)
  * "8m × 6m" or "8 mtr x 6 mtr" → length_ft=26.25, width_ft=19.69, dim_type="LxW"
  * "height 9 feet" or "9ft ceiling" or "ceiling ht 9'" → height_ft=9
  * "20 running feet" or "20 RFT" or "20 rft" → length_ft=20, dim_type="L"
  * mm → ft: divide by 304.8 · cm → ft: divide by 30.48 · m → ft: multiply by 3.281
  * Attach dimensions to the specific item they describe; for room-level area dimensions, apply to the first area-based item in that room
- If input has no room names → infer from context (bathroom items → "Bathroom", kitchen items → "Kitchen", ungrouped general text → "General Requirements")
- If ALL input is general text, meeting notes, or a project description without specific products → still extract EVERY mentioned task, material, or requirement as a separate item in the appropriate room
- item_name for general/descriptive items: use clear, professional language — material + finish + work type (e.g. "Vitrified Tile Flooring — Matt Finish", "Soft-Close Modular Kitchen Cabinets — Grey Finish")
- qty defaults to 1 unless explicitly stated or calculable from no_of_units × no_of_bathrooms_per_unit
- notes field: always populate with any context not captured as line items — budget range, timeline, brand tier preference, special requirements, client remarks

RATE ESTIMATION — MANDATORY:
Set unit_price to your best INR estimate for every item. Never leave 0 unless it is a genuinely custom or one-off item with no market precedent. Adjust ±25% for Premium or Economy tier.
Per-unit rates only — do not multiply by no_of_units.

CP FITTINGS (Nos, Standard): Basin Mixer 7500 · Shower Mixer 9500 · Overhead Shower 3500
  Hand Shower 2500 · Health Faucet 1800 · Kitchen Mixer 6000 · Bath Spout 2200
  Stop Cock 400 · Angle Cock 300 · Full Set per bathroom 18000
SANITARY WARE (Nos, Standard): EWC / One-piece WC 18000 · Two-piece WC 12000
  Counter-top Basin 8000 · Pedestal Basin 5500 · Wall-hung Basin 9000 · Bathtub 45000
BATHROOM ACCESSORIES (Nos): Towel Bar 1500 · Towel Ring 900 · Soap Dispenser 1200
  TP Holder 800 · Mirror plain 3500 · LED Mirror 8000 · Shower Enclosure 22000
  Floor Drain 600 · Exhaust Fan 1500
HARDWARE (per piece / pair): Soft-close Hinge 180 · Channel 18" soft-close (pair) 1200
  Tandem Channel 21" (pair) 2200 · Handle 128mm 350 · Handle 320mm 600
  Door Lock / Latch 1800 · Floor Spring 4500 · Door Closer 2800
TILES (SqFt, supply + fix): Ceramic 2×2 85 · Vitrified PGVT 120 · Large format 4×8 160
  Marble / Granite 250 · Anti-skid 95
WATERPROOFING (SqFt): Bathroom 65 · Terrace 80
PLUMBING: Complete bathroom (Nos) 22000 · CPVC per point 350"""

_ARCHITECT_PARSE_PROMPT = """You are an expert architectural quantity surveyor. Parse the project brief into structured JSON.
Always respond with valid JSON only — no preamble, no explanation.

Return this exact schema:
{
  "project_type": "residential|commercial|institutional|landscape|renovation|interior_only",
  "typology": "villa|row_house|apartment|duplex|office|retail|hotel|school|hospital|other",
  "plot_length": number_or_null,
  "plot_width": number_or_null,
  "plot_unit": "feet|meter",
  "floors": number_or_null,
  "units": number_or_null,
  "bedrooms": number_or_null,
  "construction_budget": number_or_null,
  "scope_of_services": ["architecture","interior","structural","mep","landscape"],
  "fee_model_suggestion": "percentage|per_sqft|lump_sum",
  "suggested_fee_pct": number,
  "complexity": "simple|medium|complex",
  "notes": "string"
}

Fee % defaults: residential villa=5-6%, apartment=4-5%, commercial=6-8%, institutional=7-9%, interior only=3-4%
Fee model: % for most projects; per_sqft for large apartment complexes; lump_sum only if explicitly requested."""


# ── Area calculation helper ───────────────────────────────────────────────────

def _calc_areas(plot_length: float, plot_width: float, plot_unit: str, floors: int, typology: str) -> dict:
    if plot_unit == "meter":
        plot_length *= 3.281
        plot_width  *= 3.281
    site_area      = round(plot_length * plot_width, 2)
    builtup_area   = round(site_area * 0.60 * floors, 2)
    carpet_area    = round(builtup_area * 0.70, 2)
    super_builtup  = round(builtup_area * 1.25, 2) if typology == "apartment" else None
    floor_plate    = round(builtup_area / floors, 2) if floors else builtup_area
    return {
        "site_area_sqft":      site_area,
        "builtup_area_sqft":   builtup_area,
        "carpet_area_sqft":    carpet_area,
        "super_builtup_sqft":  super_builtup,
        "floor_plate_sqft":    floor_plate,
    }


def _calc_fee(fee_model: str, fee_rate: float, construction_cost: float, builtup_area: float) -> float:
    if fee_model == "percentage":
        return round(construction_cost * (fee_rate / 100), 2)
    elif fee_model == "per_sqft":
        return round(builtup_area * fee_rate, 2)
    else:
        return round(fee_rate, 2)


def _default_phases(total_fee: float) -> list:
    splits = [
        ("P1 Concept Design",     10),
        ("P2 Schematic Design",   15),
        ("P3 Design Development", 20),
        ("P4 Construction Docs",  25),
        ("P5 Approvals",           5),
        ("P6 Site Supervision",   25),
    ]
    today = datetime.now()
    phases = []
    cumulative_months = [0, 1, 2, 4, 7, 8, 14]
    for i, (name, pct) in enumerate(splits):
        due = (today + timedelta(days=cumulative_months[i + 1] * 30)).strftime("%Y-%m-%d")
        phases.append({
            "phase_name":  name,
            "pct_of_total": pct,
            "fee_amount":   round(total_fee * pct / 100, 2),
            "is_paid":      False,
            "due_date":     due,
        })
    return phases


def _generate_boq(areas: dict, project_type: str, complexity: str) -> list:
    """Auto-generate standard BOQ line items from area data."""
    ba  = areas["builtup_area_sqft"]
    ca  = areas["carpet_area_sqft"]
    fp  = areas["floor_plate_sqft"]
    sa  = areas["site_area_sqft"]
    floors = max(1, round(ba / fp) if fp else 1)

    # Rate modifiers by complexity
    rate_factor = {"simple": 0.85, "medium": 1.0, "complex": 1.25}.get(complexity, 1.0)
    # Steel kg/CUM
    steel_factor = 120 if project_type == "residential" else 150

    concrete_cum = round(fp * floors * 0.0116, 2)   # 125mm slab
    steel_kg     = round(concrete_cum * steel_factor, 2)
    col_beam_cum = round(fp * floors * 0.07, 2)
    masonry_cum  = round(fp * floors * 0.023, 2)
    plaster_sqmt = round(ba * 2.5 * 0.0929, 2)       # sqft→sqmt
    floor_sqft   = round(ca, 2)
    doors_nos    = max(1, round(ca / 150))
    windows_nos  = max(1, round(ca / 120))
    supply_rmt   = round(ca * 0.12, 2)
    drainage_rmt = round(ca * 0.10, 2)
    elec_rmt     = round(ca * 0.35, 2)
    elec_pts     = max(1, round(ca / 25))
    paint_sqft   = round(ba * 2.5, 2)

    def r(rate): return round(rate * rate_factor, 2)

    items = [
        # Package A
        {"work_package": "A", "item_code": "A01", "description": "Excavation in ordinary soil",
         "unit_type": "cum", "quantity": round(fp * 2.5 * 0.0283, 2), "unit_rate": r(180)},
        {"work_package": "A", "item_code": "A02", "description": "PCC M10 (1:3:6) bed",
         "unit_type": "cum", "quantity": round(fp * 0.1 * 0.0283, 2), "unit_rate": r(4500)},
        # Package B
        {"work_package": "B", "item_code": "B01", "description": "RCC M20 footing concrete",
         "unit_type": "cum", "quantity": round(col_beam_cum * 0.15, 2), "unit_rate": r(8200)},
        {"work_package": "B", "item_code": "B02", "description": "TMT Fe 500 reinforcement steel",
         "unit_type": "kg",  "quantity": round(steel_kg * 0.15, 2), "unit_rate": r(68)},
        # Package C
        {"work_package": "C", "item_code": "C01", "description": "RCC M20 columns",
         "unit_type": "cum", "quantity": round(col_beam_cum * 0.04 * floors, 2), "unit_rate": r(8200)},
        {"work_package": "C", "item_code": "C02", "description": "RCC M20 beams",
         "unit_type": "cum", "quantity": round(col_beam_cum * 0.03 * floors, 2), "unit_rate": r(8200)},
        {"work_package": "C", "item_code": "C03", "description": "RCC M20 slabs",
         "unit_type": "cum", "quantity": concrete_cum, "unit_rate": r(8200)},
        {"work_package": "C", "item_code": "C04", "description": "TMT Fe 500 superstructure steel",
         "unit_type": "kg",  "quantity": steel_kg, "unit_rate": r(68)},
        # Package D
        {"work_package": "D", "item_code": "D01", "description": "Fly ash brick masonry (external walls)",
         "unit_type": "cum", "quantity": masonry_cum, "unit_rate": r(5800)},
        {"work_package": "D", "item_code": "D02", "description": "AAC block masonry (internal walls)",
         "unit_type": "sqmt","quantity": plaster_sqmt * 0.4, "unit_rate": r(380)},
        # Package E
        {"work_package": "E", "item_code": "E01", "description": "Waterproofing membrane (terrace/roof)",
         "unit_type": "sqft","quantity": round(fp, 2), "unit_rate": r(85)},
        {"work_package": "E", "item_code": "E02", "description": "Terrace screed + slope",
         "unit_type": "sqft","quantity": round(fp, 2), "unit_rate": r(45)},
        # Package F
        {"work_package": "F", "item_code": "F01", "description": "Internal cement plaster 12mm",
         "unit_type": "sqmt","quantity": plaster_sqmt, "unit_rate": r(220)},
        {"work_package": "F", "item_code": "F02", "description": "Vitrified tile flooring 600×600",
         "unit_type": "sqft","quantity": floor_sqft, "unit_rate": r(95)},
        # Package G
        {"work_package": "G", "item_code": "G01", "description": "Flush door with frame (solid core)",
         "unit_type": "nos", "quantity": doors_nos, "unit_rate": r(12000)},
        {"work_package": "G", "item_code": "G02", "description": "uPVC sliding window (standard)",
         "unit_type": "sqft","quantity": windows_nos * 18, "unit_rate": r(480)},
        # Package H
        {"work_package": "H", "item_code": "H01", "description": "CPVC supply piping (25mm)",
         "unit_type": "rmt", "quantity": supply_rmt, "unit_rate": r(320)},
        {"work_package": "H", "item_code": "H02", "description": "SWR drainage piping (110mm)",
         "unit_type": "rmt", "quantity": drainage_rmt, "unit_rate": r(280)},
        {"work_package": "H", "item_code": "H03", "description": "Sanitary fixture set (EWC + washbasin)",
         "unit_type": "nos", "quantity": max(1, round(ca / 400)), "unit_rate": r(18000)},
        # Package I
        {"work_package": "I", "item_code": "I01", "description": "Concealed electrical wiring (per point)",
         "unit_type": "nos", "quantity": elec_pts, "unit_rate": r(1800)},
        {"work_package": "I", "item_code": "I02", "description": "Distribution board with MCBs (8-way)",
         "unit_type": "nos", "quantity": floors, "unit_rate": r(4500)},
        # Package J
        {"work_package": "J", "item_code": "J01", "description": "Internal emulsion paint (2 coats)",
         "unit_type": "sqft","quantity": paint_sqft, "unit_rate": r(28)},
        {"work_package": "J", "item_code": "J02", "description": "External texture paint",
         "unit_type": "sqft","quantity": round(sa * 0.6, 2), "unit_rate": r(45)},
        {"work_package": "J", "item_code": "J03", "description": "Compound wall (1.5m height)",
         "unit_type": "sqmt","quantity": round(math.sqrt(sa) * 4 * 1.5 * 0.0929, 2), "unit_rate": r(2800)},
    ]

    for it in items:
        it["amount"] = round(it["quantity"] * it["unit_rate"], 2)

    return items


# ── Document text extraction ──────────────────────────────────────────────────

_DOC_EXTS_TEXT = (".txt", ".csv", ".md", ".log", ".json", ".xml")
_DOC_EXTS_PDF  = (".pdf",)
_DOC_EXTS_DOCX = (".docx", ".doc")
_DOC_EXTS_XLSX = (".xlsx", ".xls", ".ods")


def _extract_document_text(filename: str, raw: bytes, content_type: str) -> Optional[str]:
    """
    Extract readable text from uploaded documents.
    Returns extracted text string (up to 6000 chars), or None on failure.
    Supports: PDF, DOCX/DOC, XLSX/XLS, CSV, plain text.
    Libraries (pypdf, python-docx, openpyxl) are already in requirements.txt.
    """
    fn  = (filename or "").lower()
    ct  = (content_type or "").lower()

    # ── PDF ──────────────────────────────────────────────────────────────────
    if fn.endswith(_DOC_EXTS_PDF) or "pdf" in ct:
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw))
            pages  = []
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    pages.append(t.strip())
            text = "\n\n--- Page Break ---\n\n".join(pages)
            if text.strip():
                logger.info("design_quotes: extracted %d chars from PDF '%s' (%d pages)", len(text), filename, len(pages))
                return text[:6000]
        except ImportError:
            logger.warning("design_quotes: pypdf not available — falling back to raw decode for %s", filename)
        except Exception as exc:
            logger.warning("design_quotes: PDF extraction failed for '%s' — %s", filename, exc)
        # Fallback: raw bytes occasionally contain readable ASCII text
        return raw.decode("utf-8", errors="ignore")[:4000] or None

    # ── DOCX / DOC ───────────────────────────────────────────────────────────
    if fn.endswith(_DOC_EXTS_DOCX) or "wordprocessingml" in ct or "msword" in ct:
        try:
            import docx as _docx
            document = _docx.Document(io.BytesIO(raw))
            parts = []
            for para in document.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
            if text.strip():
                logger.info("design_quotes: extracted %d chars from DOCX '%s'", len(text), filename)
                return text[:6000]
        except ImportError:
            logger.warning("design_quotes: python-docx not available for %s", filename)
        except Exception as exc:
            logger.warning("design_quotes: DOCX extraction failed for '%s' — %s", filename, exc)
        return None

    # ── XLSX / XLS ───────────────────────────────────────────────────────────
    if fn.endswith(_DOC_EXTS_XLSX) or "spreadsheet" in ct or "excel" in ct:
        try:
            import openpyxl
            wb    = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            lines = []
            for ws in list(wb.worksheets)[:5]:          # max 5 sheets
                lines.append(f"\n=== Sheet: {ws.title} ===")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        lines.append(" | ".join(cells))
            text = "\n".join(lines)
            if text.strip():
                logger.info("design_quotes: extracted %d chars from XLSX '%s'", len(text), filename)
                return text[:6000]
        except ImportError:
            logger.warning("design_quotes: openpyxl not available for %s", filename)
        except Exception as exc:
            logger.warning("design_quotes: XLSX extraction failed for '%s' — %s", filename, exc)
        return None

    # ── CSV / plain text ─────────────────────────────────────────────────────
    try:
        text = raw.decode("utf-8", errors="ignore")
        return text[:4000] if text.strip() else None
    except Exception:
        return None


def _pdf_to_vision_images(raw: bytes, max_pages: int = 6, dpi_scale: float = 2.5) -> list:
    """
    Render pages of a scanned/image-based PDF as PNG images for GPT-4o vision.
    Called when pypdf yields < 150 chars — i.e., the PDF is a scanned document.
    Returns list of {"b64": str, "ct": "image/png"} dicts.
    Returns empty list if PyMuPDF is not installed or rendering fails.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning(
            "design_quotes: PyMuPDF not installed — cannot render scanned PDF for vision; "
            "run: pip install PyMuPDF>=1.23.0"
        )
        return []

    try:
        doc    = fitz.open(stream=raw, filetype="pdf")
        result = []
        for page_num in range(min(max_pages, len(doc))):
            page      = doc[page_num]
            mat       = fitz.Matrix(dpi_scale, dpi_scale)
            pix       = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("png")
            # Downscale if rendered page is too large for the vision API (>4 MB encoded)
            if len(img_bytes) > 4 * 1024 * 1024:
                mat2      = fitz.Matrix(1.5, 1.5)
                pix       = page.get_pixmap(matrix=mat2, alpha=False)
                img_bytes = pix.tobytes("png")
            b64 = base64.b64encode(img_bytes).decode()
            result.append({"b64": b64, "ct": "image/png"})
        doc.close()
        logger.info("design_quotes: rendered %d PDF page(s) as PNG for vision processing", len(result))
        return result
    except Exception as exc:
        logger.warning("design_quotes: PDF→vision render failed — %s", exc)
        return []


# ── AI helpers ────────────────────────────────────────────────────────────────

def _sanitize_doc_text(text: str) -> str:
    """Remove control characters and normalise whitespace from extracted document text."""
    if not text:
        return text
    # Keep printable ASCII + common Unicode printables; remove control chars except \n \t
    cleaned = "".join(ch if (ch in ("\n", "\t") or (ord(ch) >= 32 and ord(ch) != 127)) else " " for ch in text)
    import re as _re
    cleaned = _re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _text_has_item_content(text: str) -> bool:
    """
    Heuristic: does extracted PDF text contain actual BOQ item content
    (not just a typed header / client name / project title)?
    Used to decide whether a PDF needs the vision path even when pypdf extracted some text.
    Returns True only when the text is substantive enough to be used directly for extraction.
    """
    if not text or len(text.strip()) < 300:
        return False
    import re as _re
    # Must have item-like numeric/unit patterns
    has_units   = bool(_re.search(r'\b(Nos|SqFt|SqMtr|RFT|Mtr|Lot|LS|Bag|Cft|Kg|Set|Pair)\b', text, _re.I))
    has_dims    = bool(_re.search(r'\d+\s*[xX×]\s*\d+|\d+\.?\d*\s*(ft|\'|"|mm|cm|m)\b', text))
    has_rates   = bool(_re.search(r'(rate|price|unit\s*cost|mrp|amount)\s*:?\s*[\d,]+', text, _re.I))
    # Must have enough content lines (≥6 non-empty lines after the header)
    content_lines = [l for l in text.split('\n') if len(l.strip()) > 10]
    has_enough_lines = len(content_lines) >= 6
    return (has_units or has_dims or has_rates) and has_enough_lines


async def _ai_scan(text: str, images: Optional[List[dict]] = None) -> dict:
    """images is a list of dicts with keys b64 (str) and ct (content-type str)."""
    from app.core.config import get_settings
    cfg = get_settings()
    if not cfg.openai_api_key:
        return _demo_scan_result()
    # Sanitise document text before sending to the model
    text = _sanitize_doc_text(text) if text else text
    try:
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=cfg.openai_api_key, timeout=90.0)
        if images:
            user_content: list = []
            for img in images:
                # Warn if individual image base64 is very large (>4 MB decoded ≈ 5.4 MB b64)
                if len(img["b64"]) > 5_400_000:
                    logger.warning("design_quotes: image too large (%d chars b64), skipping", len(img["b64"]))
                    continue
                user_content.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:{img['ct']};base64,{img['b64']}", "detail": "high"},
                })
            if not user_content:
                return {
                    "scan_error": "Image(s) are too large for AI processing. Please use photos under 4 MB (reduce resolution or compress before uploading).",
                    "extracted": _demo_scan_result()["extracted"],
                    "data_source": "error",
                }
            user_content.append({
                "type": "text",
                "text": (text.strip() + "\n\n" if text and text.strip() else "") +
                        "You are an expert estimator for the Indian interior design, architecture, hardware, sanitary, and CP fittings industry. "
                        "Examine every detail in these images. Extract EVERY product, material, fitting, and component visible — "
                        "including CP fittings, sanitary ware, tiles, bathroom accessories, hardware (hinges, channels, handles, locks), "
                        "plumbing, doors, windows, flooring, false ceiling, wall paneling, furniture, and any other interior elements. "
                        "For each item give full brand+model name, type, unit, quantity (per unit and total if project scale is given), "
                        "HSN code, and specifications. Do not omit anything visible.",
            })
            messages = [
                {"role": "system", "content": _INTERIOR_SCAN_PROMPT},
                {"role": "user", "content": user_content},
            ]
        else:
            messages = [
                {"role": "system", "content": _INTERIOR_SCAN_PROMPT},
                {"role": "user", "content":
                    (text.strip() if text and text.strip() else
                     "Extract all hardware, sanitary fittings, CP fittings, plumbing, tiles, and interior fitout requirements. "
                     "Be thorough — list every product mentioned including brand, model, finish, quantity, and HSN code.")
                },
            ]
        resp = await client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            max_tokens=4096,
            response_format={"type": "json_object"},
        )
        raw = resp.choices[0].message.content

        # GPT-4o can return content=None on content_filter or tool_call finish reasons.
        # Detect this early and retry without response_format constraint.
        if not raw or not raw.strip():
            finish = resp.choices[0].finish_reason
            logger.warning("design_quotes: GPT-4o returned null content (finish_reason=%s), retrying without json_object mode", finish)
            resp2 = await client.chat.completions.create(
                model="gpt-4o",
                messages=messages,
                max_tokens=4096,
                # No response_format — ask the model to produce JSON in the system prompt
            )
            raw = resp2.choices[0].message.content or ""
            # Strip markdown fences if any (```json ... ```)
            raw = raw.strip()
            if raw.startswith("```"):
                raw = raw.split("```", 2)[1]
                if raw.startswith("json"):
                    raw = raw[4:]
                raw = raw.split("```")[0].strip()

        if not raw:
            raise ValueError(
                "OpenAI returned an empty response on both attempts. "
                "This is usually a transient API issue — please retry."
            )

        extracted = json.loads(raw)
        return {"extracted": extracted, "data_source": "ai"}
    except Exception as exc:
        logger.error("design_quotes: AI scan failed — %s", exc)
        # Key exists but API call failed — return demo data WITHOUT the misleading
        # "configure OPENAI_API_KEY" message; surface the real error instead.
        demo = _demo_scan_result()
        demo["data_source"] = "error"
        demo["demo_note"] = None
        demo["scan_error"] = f"AI scan failed: {str(exc)[:200]}"
        return demo


async def _ai_parse_architect(text: str) -> dict:
    from app.core.config import get_settings
    cfg = get_settings()
    if not cfg.openai_api_key:
        return {"parsed": {}, "data_source": "demo"}
    try:
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=cfg.openai_api_key, timeout=60.0)
        resp = await client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": _ARCHITECT_PARSE_PROMPT},
                {"role": "user", "content": text},
            ],
            max_tokens=1000,
            response_format={"type": "json_object"},
        )
        return {"parsed": json.loads(resp.choices[0].message.content), "data_source": "ai"}
    except Exception as exc:
        logger.warning("design_quotes: architect parse failed — %s", exc)
        return {"parsed": {}, "data_source": "demo"}


def _demo_scan_result() -> dict:
    return {
        "demo_note": "Demo mode — configure OPENAI_API_KEY in backend/.env to use live AI scanning.",
        "extracted": {
            "client_name": "Ramesh Constructions Pvt. Ltd.",
            "client_phone": "+91 98450 12345",
            "project_name": "Brigade Lakefront — 3BHK (24 Units)",
            "project_type": "Residential",
            "project_address": "Yelahanka, Bangalore",
            "no_of_units": 24,
            "no_of_bathrooms_per_unit": 3,
            "budget_indication": "Standard-Premium tier, target ₹32,000 per unit",
            "notes": "Builder rate required. Delivery in 3 batches of 8 units.",
            "rooms": [
                {
                    "room_name": "Master Bathroom",
                    "items": [
                        {
                            "item_name": "Jaquar Florentine Single Lever Basin Mixer — Chrome",
                            "item_type": "cp_fittings", "dim_type": "fixed", "unit": "Nos", "qty": 24,
                            "specifications": "Hot & Cold, wall-mounted, 35mm cartridge",
                            "material_preference": "Standard", "inferred_hsn": "8481",
                        },
                        {
                            "item_name": "Jaquar Overhead Shower 6\" Round with Arm — Chrome",
                            "item_type": "cp_fittings", "dim_type": "fixed", "unit": "Set", "qty": 24,
                            "specifications": "SS arm 450mm, ABS shower head",
                            "material_preference": "Standard", "inferred_hsn": "8481",
                        },
                        {
                            "item_name": "Hindware Queo Wall-Hung EWC with Soft-Close Seat",
                            "item_type": "sanitary_ware", "dim_type": "fixed", "unit": "Set", "qty": 24,
                            "specifications": "Dual flush 3/6L, concealed cistern",
                            "material_preference": "Standard", "inferred_hsn": "6910",
                        },
                        {
                            "item_name": "Bathroom Accessories Set 5-Piece — Chrome",
                            "item_type": "bathroom_accessories", "dim_type": "fixed", "unit": "Set", "qty": 24,
                            "specifications": "Towel bar 24\" + Towel ring + Robe hook + Soap dish + TP holder",
                            "material_preference": "Standard", "inferred_hsn": "3922",
                        },
                    ],
                },
                {
                    "room_name": "Common Bathroom",
                    "items": [
                        {
                            "item_name": "Jaquar Essco Bib Cock — Chrome",
                            "item_type": "cp_fittings", "dim_type": "fixed", "unit": "Nos", "qty": 48,
                            "specifications": "Wall-mounted, 15mm",
                            "material_preference": "Economy", "inferred_hsn": "8481",
                        },
                        {
                            "item_name": "Parryware Calla Two-Piece WC",
                            "item_type": "sanitary_ware", "dim_type": "fixed", "unit": "Set", "qty": 48,
                            "specifications": "S-trap, 300mm, 6L flush, white",
                            "material_preference": "Economy", "inferred_hsn": "6910",
                        },
                    ],
                },
                {
                    "room_name": "Kitchen",
                    "items": [
                        {
                            "item_name": "Franke SilentPlus 1.5 Bowl Kitchen Sink — SS",
                            "item_type": "other", "dim_type": "fixed", "unit": "Nos", "qty": 24,
                            "specifications": "304 SS, sound deadening pads, drain basket included",
                            "material_preference": "Standard", "inferred_hsn": "7324",
                        },
                        {
                            "item_name": "Hettich Sensys Soft-Close Hinge",
                            "item_type": "hardware_hinges", "dim_type": "fixed", "unit": "Nos", "qty": 480,
                            "specifications": "clip-top, full overlay, 110° opening",
                            "material_preference": "Standard", "inferred_hsn": "8302",
                        },
                    ],
                },
            ],
        },
        "data_source": "demo",
    }


# ── Interior Brief Parser (AI NLP → structured quotation items) ───────────────

_INTERIOR_BRIEF_PROMPT = """You are an expert interior design estimator. Parse the user's brief into structured JSON.
Respond with valid JSON only — no preamble, no explanation.

FIXED DEFAULTS (always apply unless overridden by user):
- Kitchen base cabinets: carcass = "16mm BWP Board", depth = 600mm
- Kitchen wall cabinets: carcass = "16mm BWP Board", depth = 300mm
- Wardrobe: carcass = "18mm HDHMR", depth = 600mm
- TV Unit: carcass = "18mm MDF", depth = 400mm
- All hardware: soft-close unless "regular" or "standard" specified
- Gola / profile handle → item_type hardware = "gola_profile", calc rmt = shutter_width × shutter_count
- Tandem keyword → include tandem_box hardware

UNIT TYPE LOGIC:
- Kitchen cabinets → unit = "rft" (running feet, measure by width)
- Wardrobe, TV unit → unit = "rft"
- Wall paneling, dado tiles, flooring, false ceiling → unit = "sqft"
- Countertop → unit = "slab" (ceil to whole slabs; 1 slab = 27.5 sqft; calc: length×0.667ft / 27.5, ceil)
- Arch, cornice, beading, railing → unit = "rmt"
- Doors, chairs, loose furniture → unit = "nos"
- Civil, paint, electrical → unit = "ls" (lump sum)
- If user specifies a unit explicitly, use that instead

HARDWARE COUNT RULES:
- Shutter height < 600mm → 2 hinges per shutter
- Shutter height 600–900mm → 3 hinges per shutter
- Shutter height > 900mm → 4 hinges per shutter
- Every drawer → 1 set channels (soft-close by default)
- Tandem drawer → 1 tandem_box per drawer (replaces regular drawer)
- Gola profile → rmt = (shutter_width_ft × shutter_count) × 0.3048m

Return this schema:
{
  "items": [
    {
      "item_name": "string",
      "description_lines": ["16mm BWP Board — carcass", "Acrylic finish shutters", ...],
      "item_type": "cabinet|wardrobe|countertop|false_ceiling|flooring|dado_tiles|wall_panel|mandir|bed|tv_unit|door|hardware",
      "dim_type": "L|LxW|LxH|WxH|fixed",
      "length_ft": number_or_null,
      "width_ft": number_or_null,
      "height_ft": number_or_null,
      "unit": "rft|sqft|sqmt|rmt|nos|slab|ls|set",
      "quantity": number,
      "unit_price": number_or_null,
      "material": "string",
      "finish": "string",
      "hardware_list": [
        {"name": "Soft-close hinge", "qty": number, "unit": "nos", "note": "string"},
        ...
      ],
      "section_name": "Kitchen|Bedroom|Living Room|etc",
      "notes": "any assumptions"
    }
  ],
  "project_type": "Residential|Commercial|etc",
  "notes": "overall notes"
}

DESCRIPTION RULES:
- Always list fixed defaults first (carcass, board thickness)
- Then user-specified finishes, handles, accessories
- Never show null or N/A — omit fields not specified
- Countertop: show "X slabs required (ceil from Ysqft)"
"""


async def _ai_parse_interior_brief(text: str) -> dict:
    from app.core.config import get_settings
    cfg = get_settings()
    if not cfg.openai_api_key:
        return _demo_interior_parse(text)
    try:
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=cfg.openai_api_key, timeout=60.0)
        resp = await client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": _INTERIOR_BRIEF_PROMPT},
                {"role": "user", "content": text[:2000]},
            ],
            max_tokens=2000,
            response_format={"type": "json_object"},
        )
        parsed = json.loads(resp.choices[0].message.content)
        return {"parsed": parsed, "data_source": "ai"}
    except Exception as exc:
        logger.warning("design_quotes: interior brief parse failed — %s", exc)
        return _demo_interior_parse(text)


def _demo_interior_parse(text: str) -> dict:
    t = text.lower()
    items = []
    if "kitchen" in t:
        items.append({
            "item_name": "Kitchen Base Cabinets",
            "description_lines": ["16mm BWP Board — carcass and exposed panels", "Acrylic finish shutters", "Soft-close hinges", "Gold Gola profile handles"],
            "item_type": "cabinet", "dim_type": "L",
            "length_ft": 10, "width_ft": None, "height_ft": 2.5,
            "unit": "rft", "quantity": 10, "unit_price": 8500,
            "material": "16mm BWP Board", "finish": "Acrylic",
            "hardware_list": [
                {"name": "Soft-close hinge", "qty": 24, "unit": "nos", "note": "3 hinges × 8 shutters"},
                {"name": "Gola profile handle", "qty": 3.05, "unit": "rmt", "note": "10ft × 1 side"},
            ],
            "section_name": "Kitchen", "notes": "Demo — configure OPENAI_API_KEY for real parsing",
        })
        if "countertop" in t or "quartz" in t:
            items.append({
                "item_name": "Kitchen Countertop",
                "description_lines": ["Quartz countertop — imported", "2 slabs required (from 13.3 sqft; ceil)"],
                "item_type": "countertop", "dim_type": "LxW",
                "length_ft": 10, "width_ft": 2, "height_ft": None,
                "unit": "slab", "quantity": 2, "unit_price": 18000,
                "material": "Quartz (imported)", "finish": "",
                "hardware_list": [],
                "section_name": "Kitchen", "notes": "",
            })
    if "wardrobe" in t:
        items.append({
            "item_name": "Wardrobe",
            "description_lines": ["18mm HDHMR — carcass", "Acrylic finish shutters", "Soft-close hinges", "Soft-close channels (drawers)"],
            "item_type": "wardrobe", "dim_type": "L",
            "length_ft": 8, "width_ft": None, "height_ft": 8,
            "unit": "rft", "quantity": 8, "unit_price": 6800,
            "material": "18mm HDHMR", "finish": "Acrylic",
            "hardware_list": [
                {"name": "Soft-close hinge", "qty": 16, "unit": "nos", "note": "4 hinges × 4 shutters (H>900mm)"},
                {"name": "Soft-close channel", "qty": 2, "unit": "set", "note": "2 drawers"},
            ],
            "section_name": "Bedroom", "notes": "Demo result",
        })
    if not items:
        items.append({
            "item_name": "Interior Work",
            "description_lines": ["Custom interior work as per brief"],
            "item_type": "cabinet", "dim_type": "fixed",
            "length_ft": None, "width_ft": None, "height_ft": None,
            "unit": "ls", "quantity": 1, "unit_price": None,
            "material": "", "finish": "", "hardware_list": [],
            "section_name": "General", "notes": "Please provide more details for accurate parsing",
        })
    return {"parsed": {"items": items, "project_type": "Residential", "notes": "Demo parse result"}, "data_source": "demo"}


@router.post("/design-quotes/parse-interior-brief")
async def parse_interior_brief(payload: dict):
    text = (payload.get("brief_text") or "").strip()
    if not text:
        raise HTTPException(status_code=422, detail="brief_text is required")
    if len(text) > 2000:
        text = text[:2000]
    result = await _ai_parse_interior_brief(text)
    return result


# ── General Document Parser (any file type → quotation BOQ) ──────────────────

_GENERAL_DOCUMENT_PROMPT = """You are an expert at reading ANY Indian construction, interior design, or project specification document and converting it to a structured Bill of Quantities (BOQ).

Your task: read the complete document text and extract EVERYTHING that could be a line item in a quotation — products, materials, fittings, civil works, electrical, furniture, services, or any other deliverable.

Return ONLY valid JSON — no markdown, no explanation, no preamble.

Required JSON schema:
{
  "client_name": "full name of client/owner/buyer — check document header, salutation, or 'For:' / 'To:' fields",
  "client_phone": "phone number or null",
  "client_email": "email or null",
  "project_name": "project/building/flat name or null",
  "project_address": "site/flat/building address or null",
  "project_type": "Residential|Commercial|Hospitality|Office|Industrial|Other",
  "designer_name": "architect/designer name if mentioned or null",
  "notes": "any special requirements, payment terms, remarks found in the document",
  "rooms": [
    {
      "room_name": "Section / Room / Work Package / Floor name — use what the document uses",
      "items": [
        {
          "item_name": "Clear, complete item name — include brand, model, size if mentioned",
          "description": "Full specifications: material, finish, color, standard, model number, remarks",
          "item_type": "cp_fittings|sanitary_ware|tiles|hardware|flooring|false_ceiling|wall_panel|furniture|electrical|civil|plumbing|waterproofing|glass|aluminium|painting|other",
          "unit": "Nos|Set|Pair|SqFt|SqMtr|RFT|Mtr|Kg|Bag|Cft|Cum|LS|Lot",
          "qty": number,
          "unit_price": number_or_null,
          "length_ft": number_or_null,
          "width_ft": number_or_null,
          "height_ft": number_or_null,
          "inferred_hsn": "4-digit HSN code or null"
        }
      ]
    }
  ]
}

MANDATORY EXTRACTION RULES — follow every one:
1. READ THE ENTIRE TEXT. Every table row, bullet point, paragraph line, footnote.
2. Client name: look for "Mr/Mrs/M/s", "To:", "For:", "Customer:", "Owner:", "Name:" at start of document.
3. Project address: look for flat number, tower, building name, locality, city, pincode.
4. Group items by the document's own structure: room names, work packages, floors, area headings.
   If no headings exist, use ONE group named after the document subject.
5. Dimensions: convert everything to feet — mm÷304.8, cm÷30.48, m×3.281 (round to 2 dec).
6. Quantities: from "Qty", "Nos.", "Units", "Count" columns or inline numbers.
7. Unit prices: from "Rate", "MRP", "Price/Unit", "Unit Cost" columns. If NOT explicitly in the document,
   estimate Standard-grade South Indian market rate (2025) — do NOT leave null.
   Key benchmarks (INR): Curtain Rod/RFT 250 · UPVC Window/SqFt 520 · False Ceiling PVC/SqFt 100
   Wardrobe/SqFt 1500 · Kitchen/RFT 2800 · Vitrified Tile/SqFt 120 · Painting/SqFt 40
   Basin Mixer 7500 · EWC/WC 18000 · Soft-close Hinge 180 · Channel 18" pair 1200
   Waterproofing/SqFt 65 · Plumbing per bathroom 22000
8. HSN inference: cp_fittings→8481, sanitary_ware→6910, tiles→6907, hardware→8302,
   electrical→8536, plumbing→3917, civil/construction→6901, painting→3210, glass→7005.
9. NEVER return an empty rooms array. If the document has ANY readable text, extract at least one item.
10. If a row in a BOQ table is a heading/subtotal row, skip it (don't create an item for it).
11. For items without explicit qty, default qty = 1.
12. Translate non-English item names to English.

UNIVERSAL EXTRACTION RULES — CRITICAL — apply to ALL input types including plain text, chat messages, voice transcripts, and vague descriptions:
13. For ANY space/area/room mentioned (kitchen, bedroom, bathroom, living room, dining, hall, office, terrace, balcony, study, gym, utility, puja room, etc.) → create a separate room entry with that name
14. For ANY work, task, material, service, or scope described → create a line item under the relevant room:
    • "Modular kitchen / kitchen works" → item_name="Modular Kitchen Works", item_type="other", unit="LS"
    • "Bathroom renovation / bathroom fitout / complete bathroom" → item_name="Bathroom Renovation Works", item_type="other", unit="LS"
    • "Flooring / tiles / vitrified" → item_name="Flooring Works", item_type="flooring", unit="SqFt"
    • "False ceiling / POP / gypsum / grid ceiling" → item_name="False Ceiling Works", item_type="false_ceiling", unit="SqFt"
    • "Paint / painting / texture / wallpaper" → item_name="Painting Works", item_type="painting", unit="SqFt"
    • "Electrical / lighting / wiring / fixtures" → item_name="Electrical & Lighting Works", item_type="electrical", unit="LS"
    • "Wardrobe / storage / shelving" → item_name="Custom Wardrobe / Storage", item_type="furniture", unit="RFT"
    • "Renovation / fitout / interior" (vague) → item_name="General Interior Works", item_type="other", unit="LS"
    • "Door / window / glazing / shutter" → item_name="Door / Window Works", item_type="other", unit="Nos"
    • "Civil / demolition / structural" → item_name="Civil Works", item_type="civil", unit="LS"
    • "Plumbing / sanitary / drainage" → item_name="Plumbing & Sanitary Works", item_type="plumbing", unit="LS"
15. Single-sentence inputs: "renovate my kitchen and 2 bathrooms" → create rooms "Kitchen" and "Bathroom 1" + "Bathroom 2" each with a line item
16. Dimension extraction (mandatory): apply rule 5 to extract from ANY format — "10×12 kitchen" → length_ft=10, width_ft=12 on the Kitchen room's first area item
17. notes field: always populate with budget, timeline, brand preferences, special requirements, or any context not captured as a line item.

HANDWRITTEN / SCANNED DOCUMENT RULES — apply whenever reading photos, scans, or hand-noted requirement lists:
18. Indian shorthand room names — ALWAYS expand before creating sections:
    MBR / M.Bed = Master Bedroom | KBR / K.Bed = Kids Bedroom | Livi / Livi. = Living Room
    GBR = Guest Bedroom | MBT / M.Bath = Master Bathroom | CBT = Common Bathroom
    KT = Kitchen | Bal = Balcony | Terr / T = Terrace | DNG = Dining Room
    Study = Study Room | Util = Utility Room | PWD = Powder Room | Corr = Corridor
19. Grouped sub-item pattern: when a numbered/bulleted entry (e.g. "9) Curtain Rods all Windows") is followed by sub-entries that list room labels with dimensions (e.g. "6.72 · MBR", "6.92 · Livi", "4.355 · Balcony Door", "5' · KBR"), create ONE item per sub-entry placed under that room's own section. Use the dimension value as length_ft. Do NOT create a generic "Curtain Rods" section — put each curtain rod item under its specific room.
20. Feet/inch notation: number followed by ' or ft = feet (6.72' → length_ft=6.72); number followed by " or in = convert to feet (14" → 1.17 ft). Decimal feet are fine. Mixed "W×H" format → width_ft × height_ft.
21. Numbered list items — 1) 2) 9) 10) etc. are SEPARATE items. Do NOT skip any numbered entry, even if it looks like a heading.
22. Area section grouping: when a numbered entry is clearly an area/room name (e.g. "10) Balcony") and the lines below it are sub-items (UPVC Sliding, Fixed window, Mop Storage…), create a room section named after that area and place all sub-items inside it.
23. Item-type hints for common handwritten items:
    Curtain rod / curtain track → item_type="other", inferred_hsn="8302"
    UPVC Sliding / UPVC Fixed / UPVC window → item_type="aluminium", inferred_hsn="7610"
    False Ceiling PVC / PVC ceiling → item_type="false_ceiling", inferred_hsn="3925"
    Cloth hanger / ceiling hanger → item_type="other", inferred_hsn="7326"
    Mop storage / wall storage / storage unit → item_type="furniture", inferred_hsn="9403"
    Tulsi stand / garden stand → item_type="other", inferred_hsn="9403"
    Pipe panelling / wall panelling → item_type="wall_panel", inferred_hsn="3925"
24. Dimension items (e.g. "UPVC Sliding 8.22 x 7.45"): capture first number as length_ft, second as width_ft or height_ft. For "False Ceiling PVC 3.86 x 10 ft": length_ft=3.86, width_ft=10."""


async def _ai_parse_document(text: str) -> dict:
    """General-purpose document → BOQ extraction using a permissive prompt."""
    from app.core.config import get_settings
    cfg = get_settings()
    if not cfg.openai_api_key:
        return _demo_doc_parse_result()

    clean = _sanitize_doc_text(text)
    if not clean:
        return _demo_doc_parse_result()

    try:
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=cfg.openai_api_key, timeout=90.0)

        messages = [
            {"role": "system", "content": _GENERAL_DOCUMENT_PROMPT},
            {"role": "user",   "content": clean[:8000]},  # generous limit for large documents
        ]

        resp = await client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            max_tokens=4096,
            response_format={"type": "json_object"},
        )
        raw = resp.choices[0].message.content

        # Null-content guard — retry without strict JSON mode
        if not raw or not raw.strip():
            finish = resp.choices[0].finish_reason
            logger.warning("design_quotes: parse-document null content (finish=%s), retrying", finish)
            resp2 = await client.chat.completions.create(
                model="gpt-4o",
                messages=messages,
                max_tokens=4096,
            )
            raw = resp2.choices[0].message.content or ""
            raw = raw.strip()
            if raw.startswith("```"):
                raw = raw.split("```", 2)[1]
                if raw.startswith("json"):
                    raw = raw[4:]
                raw = raw.split("```")[0].strip()

        if not raw:
            raise ValueError("OpenAI returned empty response on both attempts — please retry.")

        extracted = json.loads(raw)

        # Ensure rooms is always a list
        if not isinstance(extracted.get("rooms"), list):
            extracted["rooms"] = []

        # If AI still returned 0 rooms but gave us text, create a fallback section
        if not extracted["rooms"] and clean.strip():
            extracted["rooms"] = [{
                "room_name": "Extracted Items",
                "items": [{
                    "item_name": "Document content extracted — please review",
                    "description": clean[:500],
                    "item_type": "other", "unit": "LS", "qty": 1,
                    "unit_price": None, "inferred_hsn": None,
                }],
            }]

        return {"extracted": extracted, "data_source": "ai"}

    except Exception as exc:
        logger.error("design_quotes: parse-document failed — %s", exc)
        demo = _demo_doc_parse_result()
        demo["data_source"] = "error"
        demo["scan_error"] = f"AI extraction failed: {str(exc)[:200]}"
        return demo


# ── OCR system prompt — only transcribes, does not structure ─────────────────

_OCR_SYSTEM_PROMPT = """You are an expert OCR system specialised in reading Indian interior design and architecture handwritten documents.
Transcribe ALL visible text from the document images EXACTLY as written.
Include: client names, project names, item numbers (1) 2) 9) 10)), item names, room labels, all dimensions, quantities, and notes.
Preserve the structure — numbered items, sub-entries, and indentation matter.
Output ONLY the transcribed text. No JSON, no explanations."""


# ── Vision structuring prompt — used AFTER OCR transcription ─────────────────

_VISION_STRUCTURE_PROMPT = """You are an expert at converting Indian interior design requirement lists into structured Bill of Quantities (BOQ) JSON.

You will receive a plain-text transcription of a handwritten architect/interior requirements document. Your job is to extract EVERY item into the JSON schema below.

Return ONLY valid JSON — no markdown, no explanation.

Required JSON schema:
{
  "client_name": "full name from document header (Mr./Mrs. prefix if present) or null",
  "project_name": "project / flat / building / society name or null",
  "project_address": "site address or null",
  "project_type": "Residential|Commercial|Hospitality|Office|Other",
  "notes": "budget, timeline, remarks, or any requirements not captured as line items",
  "rooms": [
    {
      "room_name": "Room or area name",
      "items": [
        {
          "item_name": "EXACT item name as written — never use generic names",
          "description": "dimensions, finish, material, brand, colour, spec",
          "item_type": "cp_fittings|sanitary_ware|tiles|hardware|flooring|false_ceiling|wall_panel|furniture|electrical|civil|plumbing|aluminium|other",
          "unit": "RFT|Nos|SqFt|SqMtr|Set|Pair|Mtr|LS|Lot",
          "qty": 1,
          "unit_price": 0,
          "length_ft": null,
          "width_ft": null,
          "height_ft": null,
          "inferred_hsn": "4-digit HSN code"
        }
      ]
    }
  ]
}

━━━ MANDATORY RULES ━━━

ROOM NAME EXPANSION — always expand Indian shorthand:
MBR = Master Bedroom | KBR = Kids Bedroom | Livi / Livi. = Living Room
GBR = Guest Bedroom | MBT = Master Bathroom | CBT = Common Bathroom
KT = Kitchen | Bal = Balcony | Terr = Terrace | Study = Study Room
DNG = Dining | Util = Utility | Corr = Corridor | PWD = Powder Room

GROUPED SUB-ITEMS (most important rule):
When a parent entry like "Curtain Rods all Windows" is followed by sub-entries with room labels and dimensions, create ONE item per sub-entry under its OWN ROOM section:
  "- 6.72 · MBR"     → room "Master Bedroom",  item "Curtain Rod",  length_ft=6.72,  unit="RFT"
  "- 6.92 · Livi"    → room "Living Room",      item "Curtain Rod",  length_ft=6.92,  unit="RFT"
  "- 4.355 · Balcony Door" → room "Balcony",    item "Curtain Rod — Balcony Door", length_ft=4.355, unit="RFT"
  "- 5' · KBR"       → room "Kids Bedroom",     item "Curtain Rod",  length_ft=5.0,   unit="RFT"

AREA SECTION ITEMS:
When a heading like "Balcony" is followed by a list, create a "Balcony" room and put each listed item as a separate items entry:
  "UPVC Sliding - 8.22 x 7.45"     → item_name="UPVC Sliding Window",    length_ft=8.22, width_ft=7.45, item_type="aluminium", hsn="7610"
  "Fixed - 4.55 x 7.45"            → item_name="UPVC Fixed Window",       length_ft=4.55, width_ft=7.45, item_type="aluminium", hsn="7610"
  "Mop Storage - 1 nos"             → item_name="Mop Storage Unit",        qty=1,          item_type="furniture", hsn="9403"
  "W-M Wall Storage - 1 nos - 3'"   → item_name="Washing Machine Wall Storage", qty=1, length_ft=3.0, item_type="furniture", hsn="9403"
  "Tulsi Stand Stone - 1 nos"       → item_name="Tulsi Stand Stone",       qty=1,          item_type="other", hsn="6802"
  "Ceiling Cloth Hanger - 9-5' max" → item_name="Ceiling Cloth Hanger",   length_ft=9.5,  item_type="other", hsn="7326"
  "Pipe Panelling 14\" x 8\""       → item_name="Pipe Panelling",          length_ft=1.17, width_ft=0.67, item_type="wall_panel", hsn="3925"
  "False Ceiling PVC - 3.86 x 10'"  → item_name="False Ceiling PVC",       length_ft=3.86, width_ft=10.0, item_type="false_ceiling", hsn="3925"

COMMON ARCHITECT VOCABULARY — always use the EXACT names below (never substitute with generics):
Storage items: Cloth Storage, Used Cloth Storage, Dressing / Dressing Unit, Mop Storage,
  W-M Wall Storage, Wall Storage, Shoe Rack, Loft Storage, Under-bed Storage
Windows / Doors: UPVC Sliding Window, UPVC Fixed Window, UPVC Casement Window,
  French Door, Sliding Door, Aluminium Window, Mosquito Net / Mesh, SS Railing
Ceilings: False Ceiling PVC, Gypsum False Ceiling, Grid Ceiling, Stretch Ceiling, POP Cornice
Décor / Fixtures: Curtain Rod, Curtain Track, Blind / Roller Blind, TV Unit,
  Mandir / Pooja Unit, Tulsi Stand Stone, Tulsi Stand, Ceiling Cloth Hanger
Panelling: Pipe Panelling, WPC Wall Panel, Fluted Panel, PVC Panel, Cladding
Civil / Structural: Demolition, Brick Work, Plastering, Waterproofing, Beam
Bathroom: Shower Partition, Shower Enclosure, Vanity Unit, Mirror Cabinet
Kitchen: Modular Kitchen, Chimney, Hob, Counter Top, Loft Cabinet, Tall Unit
Flooring: Vitrified Tile, Wooden Flooring, Epoxy Flooring, Carpet Tile, Anti-skid Tile

DIMENSION RULES:
  6.72 or 6.72' → length_ft = 6.72
  8.22 x 7.45   → length_ft = 8.22, width_ft = 7.45
  14\" (inches) → feet: 14/12 = 1.17 | 8\" → 0.67
  9-5' max      → length_ft = 9.5 (max span)

QUANTITY RULES:
  "1 nos" or "1 no" → qty = 1
  "2 nos" → qty = 2
  Dimensions without explicit qty → qty = 1

ITEM TYPE + HSN MAPPING:
  UPVC window/door → aluminium, 7610
  Curtain rod/track → other, 8302
  False ceiling / PVC panel / wall panel → false_ceiling / wall_panel, 3925
  Cloth hanger / ceiling hanger → other, 7326
  Storage / wardrobe / TV unit → furniture, 9403
  Tulsi stand / stone items → other, 6802
  CP fittings / taps → cp_fittings, 8481
  Sanitary ware / EWC → sanitary_ware, 6910
  Tiles / vitrified → tiles, 6907
  Civil / plastering → civil, 6901
  Electrical → electrical, 8536
  Flooring → flooring, 5702

NEVER use generic names like "General Interior Works", "Interior Requirements", or "Various Items".
Use the EXACT item name from the document. If unclear, use the nearest vocabulary match above.
NEVER return empty rooms or empty items arrays — every line in the transcription becomes at least one item.

RATE ESTIMATION — MANDATORY:
Set unit_price for EVERY item. Never leave 0 unless it is a genuinely custom item with no market precedent.
Use Standard-grade 2025 South India (Bengaluru / Hyderabad) market rates. Adjust ±25% for Premium/Economy.

INTERIOR FIT-OUT:
Curtain Rod / Track (RFT): 250 · UPVC Sliding Window (SqFt): 550 · UPVC Fixed Window (SqFt): 480
UPVC Casement Window (SqFt): 520 · SS Railing (RFT): 2200
False Ceiling PVC (SqFt): 100 · False Ceiling Gypsum (SqFt): 160 · Grid Ceiling (SqFt): 130
Wardrobe / Storage (SqFt shutters+carcass): 1500 · Modular Kitchen (RFT): 2800
Loft Cabinet (RFT): 1800 · Counter Top granite (SqFt): 850
TV Unit (Nos): 22000 · Shoe Rack (Nos): 8000 · Mirror plain (Nos): 4000 · LED Mirror (Nos): 8000
Mop Storage Unit (Nos): 4500 · W-M Wall Storage (Nos): 8500 · Ceiling Cloth Hanger (Nos): 2500
Tulsi Stand Stone (Nos): 3500 · Pipe Panelling / WPC Wall Panel (SqFt): 380
Fluted Panel (SqFt): 450 · Cladding (SqFt): 350
Vitrified Tile Flooring (SqFt supply+fix): 120 · Wooden Flooring (SqFt): 180
Painting (SqFt): 40 · Waterproofing (SqFt): 65
Shower Enclosure / Glass Partition (Nos): 22000 · Vanity Unit (Nos): 18000
For items not listed: estimate from similar items and your knowledge of South Indian market rates.
"""


async def _ai_ocr_transcribe(images: list, client) -> str:
    """
    Pass 1 of 2: Ask GPT-4o vision to transcribe ALL visible text from document images.
    This is a pure OCR step — no structuring, just reading every single line.
    Vision models are much more reliable at transcription than at combined read+structure.
    """
    MAX_B64_CHARS = 5_400_000
    image_blocks = []
    for img in images:
        if len(img.get("b64", "")) > MAX_B64_CHARS:
            continue
        image_blocks.append({
            "type": "image_url",
            "image_url": {"url": f"data:{img['ct']};base64,{img['b64']}", "detail": "high"},
        })

    if not image_blocks:
        return ""

    image_blocks.append({
        "type": "text",
        "text": (
            "Transcribe ALL text visible in these document images, line by line, exactly as written. "
            "Include: every numbered item (e.g. 9) Curtain Rods...), every sub-entry with dimensions "
            "(e.g. - 6.72 · MBR), every area heading (e.g. Balcony), every item listed below it "
            "(UPVC Sliding - 8.22 x 7.45, Mop Storage - 1 nos, etc.), all measurements, all quantities. "
            "Preserve indentation and numbering. Output plain text only."
        ),
    })

    try:
        resp = await client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": _OCR_SYSTEM_PROMPT},
                {"role": "user",   "content": image_blocks},
            ],
            max_tokens=2000,
            temperature=0,
        )
        transcription = resp.choices[0].message.content or ""
        logger.info("design_quotes: OCR pass yielded %d chars", len(transcription))
        return transcription
    except Exception as exc:
        logger.warning("design_quotes: OCR pass failed — %s", exc)
        return ""


async def _ai_structure_transcription(transcription: str, client, header_hint: str = "") -> dict:
    """
    Pass 2 of 2: Structure the OCR transcription into BOQ JSON using _VISION_STRUCTURE_PROMPT.
    This text-only pass is far more reliable than combined vision+structure.
    """
    combined = ""
    if header_hint.strip():
        combined = f"Document header context:\n{header_hint.strip()[:400]}\n\n"
    combined += f"Full document transcription:\n{transcription}"

    try:
        resp = await client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": _VISION_STRUCTURE_PROMPT},
                {"role": "user",   "content": combined[:8000]},
            ],
            max_tokens=4096,
            response_format={"type": "json_object"},
            temperature=0,
        )
        raw = resp.choices[0].message.content or ""
        if not raw.strip():
            return {}
        extracted = json.loads(raw)
        if not isinstance(extracted.get("rooms"), list):
            extracted["rooms"] = []
        return extracted
    except Exception as exc:
        logger.warning("design_quotes: structure pass failed — %s", exc)
        return {}


async def _ai_parse_document_vision(images: list, text_hint: str = "") -> dict:
    """
    Two-pass vision extraction for scanned/handwritten/hybrid PDFs:
      Pass 1 — OCR: GPT-4o vision reads and transcribes every visible line of text.
      Pass 2 — Structure: GPT-4o text converts the transcription into BOQ JSON.
    Separating concerns (reading vs structuring) yields specific item names, not generic fallbacks.
    Falls back to single-pass if OCR transcription is too short.
    """
    from app.core.config import get_settings
    cfg = get_settings()
    if not cfg.openai_api_key:
        return _demo_doc_parse_result()

    MAX_B64_CHARS = 5_400_000
    image_blocks: list = []
    for img in images:
        if len(img.get("b64", "")) > MAX_B64_CHARS:
            logger.warning("design_quotes: vision image too large (%d b64 chars), skipping", len(img["b64"]))
            continue
        image_blocks.append({
            "type": "image_url",
            "image_url": {"url": f"data:{img['ct']};base64,{img['b64']}", "detail": "high"},
        })

    if not image_blocks:
        logger.warning("design_quotes: no valid images after size filter")
        return _demo_doc_parse_result()

    try:
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=cfg.openai_api_key, timeout=90.0)

        # ── Pass 1: OCR transcription (pass raw images, not formatted blocks) ─
        transcription = await _ai_ocr_transcribe(images, client)

        extracted: dict = {}

        if transcription and len(transcription.strip()) > 30:
            # ── Pass 2: Structure the transcription ──────────────────────────
            logger.info("design_quotes: running structure pass on %d-char transcription", len(transcription))
            extracted = await _ai_structure_transcription(transcription, client, text_hint)
        else:
            logger.warning("design_quotes: OCR transcription too short (%d chars) — falling back to single-pass", len(transcription))

        # ── Single-pass fallback: send images directly to structure prompt ───
        if not extracted or not extracted.get("rooms") or all(
            len(r.get("items") or []) == 0 for r in extracted.get("rooms", [])
        ):
            logger.warning("design_quotes: two-pass returned 0 items — trying single-pass fallback")
            hint_text = (
                (f"Typed text from document:\n{text_hint.strip()[:400]}\n\n" if text_hint.strip() else "")
                + (f"OCR transcription:\n{transcription[:1500]}\n\n" if transcription.strip() else "")
                + "Extract ALL items from the images into the required JSON. Use EXACT item names from the document."
            )
            fallback_content = image_blocks + [{"type": "text", "text": hint_text}]
            resp_fb = await client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": _VISION_STRUCTURE_PROMPT},
                    {"role": "user",   "content": fallback_content},
                ],
                max_tokens=4096,
                response_format={"type": "json_object"},
            )
            raw_fb = resp_fb.choices[0].message.content or ""
            if raw_fb.strip():
                try:
                    extracted = json.loads(raw_fb)
                    if not isinstance(extracted.get("rooms"), list):
                        extracted["rooms"] = []
                except Exception:
                    pass

        # Final fallback: if still no items, keep at least client info with placeholder
        if not extracted.get("rooms") or all(len(r.get("items") or []) == 0 for r in extracted["rooms"]):
            logger.warning("design_quotes: vision extraction returned 0 items after retry")
            extracted.setdefault("rooms", [])
            if not extracted["rooms"]:
                extracted["rooms"] = [{
                    "room_name": "Document Requirements",
                    "items": [{
                        "item_name": "Could not read handwriting — please add items manually",
                        "description": "The document image quality may be too low for AI extraction.",
                        "item_type": "other", "unit": "LS", "qty": 1,
                        "unit_price": None, "inferred_hsn": None,
                    }],
                }]

        total = sum(len(r.get("items") or []) for r in extracted.get("rooms", []))
        logger.info("design_quotes: vision extraction complete — %d rooms, %d items",
                    len(extracted.get("rooms", [])), total)

        return {"extracted": extracted, "data_source": "ai"}

    except Exception as exc:
        logger.error("design_quotes: parse-document-vision failed — %s", exc)
        demo = _demo_doc_parse_result()
        demo["data_source"] = "error"
        demo["scan_error"] = f"Vision AI extraction failed: {str(exc)[:200]}"
        return demo


def _demo_doc_parse_result() -> dict:
    """Realistic demo result shown when OPENAI_API_KEY is not configured."""
    return {
        "demo_note": "Demo mode — configure OPENAI_API_KEY in backend/.env to extract from real documents.",
        "extracted": {
            "client_name": "Mr. Sandeep Kumar",
            "client_phone": "+91 98765 00000",
            "client_email": None,
            "project_name": "Pacific Bannerghatta — 3BHK Flat",
            "project_address": "Pacific Bannerghatta Road, Bangalore — 560 076",
            "project_type": "Residential",
            "designer_name": None,
            "notes": "Demo extraction — real document not parsed (no OpenAI API key).",
            "rooms": [
                {
                    "room_name": "Master Bathroom",
                    "items": [
                        {"item_name": "Basin Mixer — Single Lever Chrome", "description": "Hot & cold, wall-mounted, 35mm cartridge", "item_type": "cp_fittings", "unit": "Nos", "qty": 1, "unit_price": None, "length_ft": None, "width_ft": None, "height_ft": None, "inferred_hsn": "8481"},
                        {"item_name": "Overhead Shower Set", "description": "6\" round, SS arm 450mm", "item_type": "cp_fittings", "unit": "Set", "qty": 1, "unit_price": None, "length_ft": None, "width_ft": None, "height_ft": None, "inferred_hsn": "8481"},
                        {"item_name": "Wall-Hung EWC with Soft-Close Seat", "description": "Dual flush 3/6L, concealed cistern", "item_type": "sanitary_ware", "unit": "Set", "qty": 1, "unit_price": None, "length_ft": None, "width_ft": None, "height_ft": None, "inferred_hsn": "6910"},
                        {"item_name": "Bathroom Accessories 5-Piece Set", "description": "Towel bar 24\", ring, hook, soap dish, TP holder", "item_type": "sanitary_ware", "unit": "Set", "qty": 1, "unit_price": None, "length_ft": None, "width_ft": None, "height_ft": None, "inferred_hsn": "3922"},
                    ],
                },
                {
                    "room_name": "Kitchen",
                    "items": [
                        {"item_name": "SS Double Bowl Kitchen Sink", "description": "304 grade, sound deadening pads, drain basket", "item_type": "other", "unit": "Nos", "qty": 1, "unit_price": None, "length_ft": None, "width_ft": None, "height_ft": None, "inferred_hsn": "7324"},
                        {"item_name": "Kitchen Single Lever Mixer Tap", "description": "Chrome finish, swivel spout", "item_type": "cp_fittings", "unit": "Nos", "qty": 1, "unit_price": None, "length_ft": None, "width_ft": None, "height_ft": None, "inferred_hsn": "8481"},
                    ],
                },
            ],
        },
        "data_source": "demo",
    }


@router.post("/design-quotes/parse-document")
async def parse_document(
    file: List[UploadFile] = File(default=[]),
    text_input: Optional[str] = Form(default=None),
):
    """
    General-purpose document → quotation BOQ parser.
    Accepts PDF, DOCX, XLSX, CSV, images, or plain text.
    Scanned/handwritten PDFs are automatically rendered via PyMuPDF and sent to GPT-4o Vision
    using the general-purpose document prompt (not the hardware/sanitary scan prompt).
    Also extracts client name, project name, and address from document headers.
    """
    combined_text      = text_input or ""
    images: List[dict] = []
    is_scanned_doc     = False  # True when ≥1 PDF required vision fallback
    IMAGE_EXTS = (".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp")
    MAX_IMAGE_BYTES = 4 * 1024 * 1024

    for f in (file or []):
        if not f or not f.filename:
            continue
        raw = await f.read()
        ct  = (f.content_type or "").lower()
        fn  = (f.filename or "").lower()

        if ct.startswith("image/") or fn.endswith(IMAGE_EXTS):
            # Explicit image file — use as-is for vision
            if len(raw) > MAX_IMAGE_BYTES:
                raise HTTPException(
                    status_code=413,
                    detail=f"Image '{f.filename}' is {len(raw)//1024}KB — max 4 MB.",
                )
            b64 = base64.b64encode(raw).decode()
            images.append({"b64": b64, "ct": ct if ct.startswith("image/") else "image/jpeg"})

        elif fn.endswith(_DOC_EXTS_PDF) or "pdf" in ct:
            # PDFs ALWAYS go through the vision path (two-pass OCR + structure).
            # Reason: customer requirement PDFs are handwritten, scanned, or hybrid
            # (typed header + handwritten body). Many scanning apps (CamScanner, etc.)
            # embed a garbled OCR text layer that _text_has_item_content() incorrectly
            # passes, causing the text path to return "General Interior Works" fallbacks.
            # Vision with two-pass OCR is more reliable for ALL PDF types in this context.
            pdf_images = _pdf_to_vision_images(raw)
            if pdf_images:
                images.extend(pdf_images)
                is_scanned_doc = True
                # Extract any typed/digital text as a context hint for the vision prompt
                doc_text = _extract_document_text(f.filename, raw, f.content_type or "")
                if doc_text and doc_text.strip():
                    combined_text = (combined_text
                                     + f"\n\n[Typed text from {f.filename}]\n"
                                     + doc_text[:800])
                logger.info("design_quotes: '%s' → vision/OCR path (%d PDF pages rendered)", f.filename, len(pdf_images))
            else:
                # PyMuPDF not available — fall back to text extraction
                doc_text = _extract_document_text(f.filename, raw, f.content_type or "")
                if doc_text and doc_text.strip():
                    combined_text = combined_text + f"\n\n[Document: {f.filename}]\n" + doc_text
                    logger.warning("design_quotes: PyMuPDF unavailable for '%s' — using text fallback", f.filename)
                else:
                    logger.warning("design_quotes: could not extract any content from PDF '%s'", f.filename)

        else:
            # Other document types (DOCX, XLSX, CSV, plain text)
            doc_text = _extract_document_text(f.filename, raw, f.content_type or "")
            if doc_text and doc_text.strip():
                combined_text = combined_text + f"\n\n[Document: {f.filename}]\n" + doc_text
            else:
                logger.warning("design_quotes: parse-document could not extract text from '%s'", f.filename)

    if not combined_text.strip() and not images:
        raise HTTPException(status_code=400, detail="Provide at least one file or text_input")

    if images:
        if is_scanned_doc:
            # Scanned / handwritten document → use _GENERAL_DOCUMENT_PROMPT with GPT-4o vision
            return await _ai_parse_document_vision(images, combined_text)
        else:
            # Explicit product/room photos → use _INTERIOR_SCAN_PROMPT (hardware/sanitary focused)
            return await _ai_scan(combined_text, images)

    # Text / document path → use _GENERAL_DOCUMENT_PROMPT
    return await _ai_parse_document(combined_text)


# ── Routes: Templates ─────────────────────────────────────────────────────────

@router.get("/design-quotes/templates")
async def get_templates():
    return {
        "room_templates":  ROOM_TEMPLATES,
        "product_options": PRODUCT_OPTIONS,
        "data_source": "static",
    }


# ── Routes: Interior Quotations ───────────────────────────────────────────────

@router.get("/design-quotes")
async def list_quotes(status: Optional[str] = None, search: Optional[str] = None):
    _init_demo()
    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    where, params = [], []
                    if status:
                        where.append("status = %s"); params.append(status)
                    if search:
                        like = f"%{search}%"
                        where.append("(client_name LIKE %s OR project_name LIKE %s OR quote_number LIKE %s)")
                        params.extend([like, like, like])
                    q = "SELECT * FROM design_quotes"
                    if where:
                        q += " WHERE " + " AND ".join(where)
                    q += " ORDER BY created_at DESC LIMIT 200"
                    await cur.execute(q, params)
                    cols = [d[0] for d in cur.description]
                    rows = await cur.fetchall()
                    quotes = []
                    for row in rows:
                        d = dict(zip(cols, row))
                        d["sections"] = json.loads(d.get("sections_json") or "[]")
                        raw_created = d.get("created_at")
                        validity = int(d.get("validity_days") or 30)
                        d["valid_till"] = (
                            (raw_created + timedelta(days=validity)).strftime("%Y-%m-%d")
                            if hasattr(raw_created, "date") else _valid_till(validity)
                        )
                        d["created_at"] = str(raw_created)[:10] if raw_created else ""
                        quotes.append(d)
            return {"quotes": quotes, "data_source": "live"}
        except Exception as exc:
            logger.warning("design_quotes: list DB failed — %s", exc)

    # Demo fallback
    qs = list(_demo_quotes.values())
    if status:
        qs = [q for q in qs if q["status"] == status]
    if search:
        sl = search.lower()
        qs = [q for q in qs if sl in q.get("client_name","").lower()
              or sl in q.get("project_name","").lower()
              or sl in q.get("quote_number","").lower()]
    return {"quotes": sorted(qs, key=lambda x: x["created_at"], reverse=True), "data_source": "demo"}


@router.get("/design-quotes/{quote_id}")
async def get_quote(quote_id: int):
    _init_demo()
    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute("SELECT * FROM design_quotes WHERE id = %s", (quote_id,))
                    row = await cur.fetchone()
                    if not row:
                        raise HTTPException(status_code=404, detail="Quote not found")
                    cols = [d[0] for d in cur.description]
                    d = dict(zip(cols, row))
                    d["sections"] = json.loads(d.get("sections_json") or "[]")
                    raw_created = d.get("created_at")
                    validity = int(d.get("validity_days") or 30)
                    d["valid_till"] = (
                        (raw_created + timedelta(days=validity)).strftime("%Y-%m-%d")
                        if hasattr(raw_created, "date") else _valid_till(validity)
                    )
                    d["created_at"] = str(raw_created)[:10] if raw_created else ""
                    return d
        except HTTPException:
            raise
        except Exception as exc:
            logger.warning("design_quotes: get DB failed — %s", exc)

    q = _demo_quotes.get(quote_id)
    if not q:
        raise HTTPException(status_code=404, detail="Quote not found")
    return q


@router.post("/design-quotes")
async def create_quote(body: dict):
    _init_demo()
    sections = body.get("sections", [])
    qnum = body.get("quote_number") or _next_quote_number()
    validity_days = int(body.get("validity_days") or 30)

    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute("""
                        INSERT INTO design_quotes
                        (quote_number, client_name, client_phone, client_email, project_name,
                         project_address, project_type, designer_name, designer_company,
                         payment_terms, validity_days, gst_rate, include_gst, notes, terms,
                         status, subtotal, gst_amount, grand_total, total_area_sqft, sections_json,
                         margin_mode, overall_margin_pct)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                    """, (
                        qnum,
                        body.get("client_name", ""),   body.get("client_phone", ""),
                        body.get("client_email", ""),   body.get("project_name", ""),
                        body.get("project_address", ""),body.get("project_type", "Residential"),
                        body.get("designer_name", ""),  body.get("designer_company", ""),
                        body.get("payment_terms", ""),  validity_days,
                        float(body.get("gst_rate", 18)), int(bool(body.get("include_gst", True))),
                        body.get("notes", ""),          body.get("terms", ""),
                        body.get("status", "DRAFT"),
                        float(body.get("subtotal", 0)),  float(body.get("gst_amount", 0)),
                        float(body.get("grand_total", 0)),float(body.get("total_area_sqft", 0)),
                        json.dumps(sections),
                        body.get("margin_mode", "per_line"),
                        float(body.get("overall_margin_pct", 0)),
                    ))
                    new_id = cur.lastrowid
                    await conn.commit()
            return {"id": new_id, "quote_number": qnum, "data_source": "live"}
        except Exception as exc:
            logger.warning("design_quotes: create DB failed — %s", exc)

    new_id = max(_demo_quotes.keys(), default=0) + 1
    _demo_quotes[new_id] = {
        **body, "id": new_id, "quote_number": qnum,
        "sections": sections,
        "created_at": datetime.now().strftime("%Y-%m-%d"),
        "valid_till": _valid_till(validity_days),
    }
    return {"id": new_id, "quote_number": qnum, "data_source": "demo"}


@router.put("/design-quotes/{quote_id}")
async def update_quote(quote_id: int, body: dict):
    _init_demo()
    sections = body.get("sections", [])

    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute("""
                        UPDATE design_quotes SET
                          client_name=%s, client_phone=%s, client_email=%s, project_name=%s,
                          project_address=%s, project_type=%s, designer_name=%s, designer_company=%s,
                          payment_terms=%s, validity_days=%s, gst_rate=%s, include_gst=%s,
                          notes=%s, terms=%s, status=%s, subtotal=%s, gst_amount=%s, grand_total=%s,
                          total_area_sqft=%s, sections_json=%s,
                          margin_mode=%s, overall_margin_pct=%s
                        WHERE id=%s
                    """, (
                        body.get("client_name", ""),   body.get("client_phone", ""),
                        body.get("client_email", ""),   body.get("project_name", ""),
                        body.get("project_address", ""),body.get("project_type", "Residential"),
                        body.get("designer_name", ""),  body.get("designer_company", ""),
                        body.get("payment_terms", ""),  int(body.get("validity_days", 30)),
                        float(body.get("gst_rate", 18)), int(bool(body.get("include_gst", True))),
                        body.get("notes", ""),          body.get("terms", ""),
                        body.get("status", "DRAFT"),
                        float(body.get("subtotal", 0)),  float(body.get("gst_amount", 0)),
                        float(body.get("grand_total", 0)),float(body.get("total_area_sqft", 0)),
                        json.dumps(sections),
                        body.get("margin_mode", "per_line"),
                        float(body.get("overall_margin_pct", 0)),
                        quote_id,
                    ))
                    await conn.commit()
            return {"id": quote_id, "data_source": "live"}
        except Exception as exc:
            logger.warning("design_quotes: update DB failed — %s", exc)

    if quote_id in _demo_quotes:
        _demo_quotes[quote_id] = {**_demo_quotes[quote_id], **body, "sections": sections}
    return {"id": quote_id, "data_source": "demo"}


# ── Approval workflow constants ───────────────────────────────────────────────
# L1 = sales_manager, L2 = cfo, L3 = admin
# Transition table: (current_status, actor_role) → allowed_actions
_APPROVAL_TRANSITIONS = {
    # Any authenticated user may submit a DRAFT for internal approval
    ("DRAFT",      "architect"):     ["SUBMIT"],
    ("DRAFT",      "sales_manager"): ["SUBMIT"],
    ("DRAFT",      "cfo"):           ["SUBMIT"],
    ("DRAFT",      "admin"):         ["SUBMIT"],
    # L1 (sales_manager): approve or escalate to L2
    ("PENDING_L1", "sales_manager"): ["APPROVE", "ESCALATE_L2"],
    ("PENDING_L1", "admin"):         ["APPROVE", "ESCALATE_L2"],
    # L2 (cfo): approve or escalate to L3
    ("PENDING_L2", "cfo"):           ["APPROVE", "ESCALATE_L3"],
    ("PENDING_L2", "admin"):         ["APPROVE", "ESCALATE_L3"],
    # L3 (admin only): approve → returns to L1 for final sign-off; or reject
    ("PENDING_L3", "admin"):         ["APPROVE_RETURN_L1", "REJECT"],
    # Reject is available to any approver at any pending level
    ("PENDING_L1", "sales_manager"): ["APPROVE", "ESCALATE_L2", "REJECT"],
    ("PENDING_L1", "admin"):         ["APPROVE", "ESCALATE_L2", "REJECT"],
    ("PENDING_L2", "cfo"):           ["APPROVE", "ESCALATE_L3", "REJECT"],
    ("PENDING_L2", "admin"):         ["APPROVE", "ESCALATE_L3", "REJECT"],
}

_ACTION_TO_STATUS = {
    "SUBMIT":           "PENDING_L1",
    "APPROVE":          "APPROVED",
    "ESCALATE_L2":      "PENDING_L2",
    "ESCALATE_L3":      "PENDING_L3",
    "APPROVE_RETURN_L1":"PENDING_L1",
    "REJECT":           "DRAFT",
}

_LEVEL_FOR_STATUS = {"PENDING_L1": 1, "PENDING_L2": 2, "PENDING_L3": 3}

_ROLE_LABEL = {
    "architect": "Architect", "sales_manager": "Sales Manager",
    "cfo": "CFO", "admin": "Admin",
}


def _get_actor_from_request(request: Request) -> dict:
    """Extract actor role + display name from JWT in Authorization header. Returns defaults on failure."""
    try:
        from app.core.auth import decode_token
        auth_header = request.headers.get("authorization", "")
        token = auth_header.replace("Bearer ", "").strip()
        if not token:
            return {"role": "architect", "name": "Unknown User"}
        payload = decode_token(token)
        return {
            "role": payload.get("role", "architect"),
            "name": payload.get("display_name") or payload.get("sub", "Unknown"),
        }
    except Exception:
        return {"role": "architect", "name": "Unknown User"}


async def _record_approval_history(pool, quote_id: int, level: int, action: str,
                                   actor_role: str, actor_name: str,
                                   notes: str = "", ai_rec: str = "") -> None:
    if not pool:
        return
    try:
        async with pool.acquire() as conn:
            async with conn.cursor() as cur:
                await cur.execute(
                    """INSERT INTO quote_approval_history
                       (quote_id, level, action, actor_role, actor_name, notes, ai_rec)
                       VALUES (%s,%s,%s,%s,%s,%s,%s)""",
                    (quote_id, level, action, actor_role, actor_name, notes, ai_rec)
                )
                await conn.commit()
    except Exception as exc:
        logger.warning("design_quotes: approval history insert failed — %s", exc)


@router.put("/design-quotes/{quote_id}/status")
async def update_quote_status(quote_id: int, body: dict, request: Request):
    """
    Role-aware, 3-level internal approval engine for design quotations.
    Transition rules:
      DRAFT        → PENDING_L1  (any user: Submit for Approval)
      PENDING_L1   → APPROVED    (L1/admin: Approve)
                   → PENDING_L2  (L1/admin: Escalate to L2)
      PENDING_L2   → APPROVED    (L2/admin: Approve)
                   → PENDING_L3  (L2/admin: Escalate to L3)
      PENDING_L3   → PENDING_L1  (admin: L3 Approves → back to L1 for final)
      any PENDING  → DRAFT       (reject → back to drafter)
    Once L3 returns to L1 and L1 approves → APPROVED (cycle complete).
    """
    _init_demo()
    actor = _get_actor_from_request(request)
    actor_role = actor["role"]
    actor_name = actor["name"]
    action     = body.get("action", "")
    notes      = body.get("notes", "")

    # ── Fetch current status ───────────────────────────────────────────────────
    current_status = "DRAFT"
    approval_cycle = 0
    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute(
                        "SELECT status, COALESCE(approval_cycle,0) FROM design_quotes WHERE id=%s",
                        (quote_id,)
                    )
                    row = await cur.fetchone()
                    if row:
                        current_status, approval_cycle = row[0], row[1] or 0
        except Exception as exc:
            logger.warning("design_quotes: fetch status failed — %s", exc)
    else:
        q = _demo_quotes.get(quote_id, {})
        current_status = q.get("status", "DRAFT")
        approval_cycle = q.get("approval_cycle", 0)

    # ── Validate action is allowed for this role + current status ──────────────
    allowed = _APPROVAL_TRANSITIONS.get((current_status, actor_role), [])
    # admin override: always allowed to approve any pending quote
    if actor_role == "admin" and current_status.startswith("PENDING"):
        if "APPROVE" not in allowed:
            allowed = list(allowed) + ["APPROVE"]

    if action not in allowed:
        raise HTTPException(
            status_code=403,
            detail=f"Role '{actor_role}' cannot perform '{action}' on a quote with status '{current_status}'. "
                   f"Allowed actions: {allowed or ['none']}",
        )

    # ── Compute new status ─────────────────────────────────────────────────────
    new_status = _ACTION_TO_STATUS.get(action, current_status)

    # When L3 returns to L1, increment the approval cycle counter
    if action == "APPROVE_RETURN_L1":
        approval_cycle += 1

    level = _LEVEL_FOR_STATUS.get(current_status, 0)
    if action == "SUBMIT":
        level = 0  # submission recorded at level 0

    # ── Persist ────────────────────────────────────────────────────────────────
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute(
                        "UPDATE design_quotes SET status=%s, approval_cycle=%s, last_approver=%s WHERE id=%s",
                        (new_status, approval_cycle, actor_name, quote_id)
                    )
                    await conn.commit()
            await _record_approval_history(
                pool, quote_id, level, action, actor_role, actor_name, notes
            )
            return {"id": quote_id, "status": new_status, "action": action,
                    "actor": actor_name, "approval_cycle": approval_cycle,
                    "data_source": "live"}
        except Exception as exc:
            logger.warning("design_quotes: status update DB failed — %s", exc)

    # Demo fallback
    if quote_id in _demo_quotes:
        _demo_quotes[quote_id]["status"] = new_status
        _demo_quotes[quote_id]["approval_cycle"] = approval_cycle
        _demo_quotes[quote_id]["last_approver"] = actor_name
    return {"id": quote_id, "status": new_status, "action": action,
            "actor": actor_name, "approval_cycle": approval_cycle,
            "data_source": "demo"}


@router.get("/design-quotes/{quote_id}/approval-history")
async def get_approval_history(quote_id: int):
    """Return the full approval trail for a quote."""
    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute(
                        """SELECT level, action, actor_role, actor_name, notes, ai_rec, created_at
                           FROM quote_approval_history WHERE quote_id=%s ORDER BY created_at ASC""",
                        (quote_id,)
                    )
                    cols = ["level","action","actor_role","actor_name","notes","ai_rec","created_at"]
                    rows = await cur.fetchall()
                    history = [dict(zip(cols, r)) for r in rows]
                    for h in history:
                        if hasattr(h["created_at"], "isoformat"):
                            h["created_at"] = h["created_at"].isoformat()
            return {"quote_id": quote_id, "history": history, "data_source": "live"}
        except Exception as exc:
            logger.warning("design_quotes: approval history fetch failed — %s", exc)
    return {"quote_id": quote_id, "history": [], "data_source": "demo"}


@router.post("/design-quotes/{quote_id}/ai-approval-recommendation")
async def ai_approval_recommendation(quote_id: int, request: Request):
    """
    AI-powered approval recommendation for managers.
    Analyzes quote value, margin, complexity, and rates vs market benchmarks.
    Returns: recommendation (APPROVE/ESCALATE), confidence, reasoning, risk_factors.
    """
    pool = await _get_db_pool()
    quote = None

    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute("SELECT * FROM design_quotes WHERE id=%s", (quote_id,))
                    cols = [d[0] for d in cur.description]
                    row  = await cur.fetchone()
                    if row:
                        quote = dict(zip(cols, row))
                        quote["sections"] = json.loads(quote.get("sections_json") or "[]")
        except Exception as exc:
            logger.warning("design_quotes: AI rec fetch failed — %s", exc)

    if not quote:
        quote = _demo_quotes.get(quote_id) or {}

    if not quote:
        raise HTTPException(status_code=404, detail="Quote not found")

    api_key = os.getenv("OPENAI_API_KEY", "")
    if not api_key:
        return {
            "recommendation": "APPROVE",
            "confidence": 65,
            "reasoning": "AI analysis unavailable (no API key). Standard review recommended.",
            "risk_factors": [],
            "data_source": "demo",
        }

    try:
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=api_key, timeout=30.0)

        sections = quote.get("sections") or []
        num_sections = len(sections)
        num_items = sum(len(s.get("items", [])) for s in sections)
        grand_total = float(quote.get("grand_total") or 0)
        gst_rate    = float(quote.get("gst_rate") or 18)
        project_type = quote.get("project_type", "Residential")
        client_name  = quote.get("client_name", "Unknown")

        section_summary = "; ".join(
            f"{s.get('section_name','Section')} ({len(s.get('items',[]))} items, ₹{s.get('section_total',0):,.0f})"
            for s in sections[:8]
        )

        prompt = f"""You are a senior interior design studio manager reviewing a client quotation before approval.

Quote Summary:
- Client: {client_name}
- Project: {quote.get('project_name','—')} ({project_type})
- Sections: {num_sections} rooms / {num_items} items
- Grand Total: ₹{grand_total:,.0f} (incl. {gst_rate}% GST)
- Section breakdown: {section_summary or 'No sections'}
- Payment Terms: {quote.get('payment_terms','—')}

Evaluate this quotation and provide an approval recommendation. Consider:
1. Quote value appropriateness for the project type and scope
2. Number of rooms/sections vs typical {project_type.lower()} fit-out
3. Any unusual or missing items for this scope
4. Financial risk for the studio

Respond in this exact JSON format:
{{
  "recommendation": "APPROVE" or "ESCALATE",
  "confidence": 0-100,
  "reasoning": "2-3 sentence explanation",
  "risk_factors": ["factor1", "factor2"],
  "value_assessment": "LOW/STANDARD/HIGH",
  "notes_for_approver": "one actionable sentence"
}}"""

        resp = await client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=400,
            response_format={"type": "json_object"},
        )
        raw = resp.choices[0].message.content or "{}"
        result = json.loads(raw)
        result["data_source"] = "ai"
        result["quote_id"] = quote_id
        return result
    except Exception as exc:
        logger.warning("design_quotes: AI recommendation failed — %s", exc)
        return {
            "recommendation": "APPROVE",
            "confidence": 60,
            "reasoning": f"AI analysis encountered an error ({type(exc).__name__}). Please review manually.",
            "risk_factors": [],
            "data_source": "error",
        }


@router.get("/design-quotes/pending-approvals")
async def get_pending_approvals(request: Request):
    """Return quotes pending the calling user's approval level."""
    actor = _get_actor_from_request(request)
    role  = actor["role"]

    # Determine which statuses this role can act on
    pending_statuses = []
    if role in ("sales_manager", "admin"):
        pending_statuses.append("PENDING_L1")
    if role in ("cfo", "admin"):
        pending_statuses.append("PENDING_L2")
    if role == "admin":
        pending_statuses.append("PENDING_L3")

    if not pending_statuses:
        return {"quotes": [], "total": 0, "role": role, "data_source": "live"}

    placeholders = ",".join(["%s"] * len(pending_statuses))
    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute(
                        f"SELECT id, quote_number, client_name, project_name, status, "
                        f"grand_total, created_at, COALESCE(approval_cycle,0), COALESCE(last_approver,'') "
                        f"FROM design_quotes WHERE status IN ({placeholders}) ORDER BY created_at ASC",
                        pending_statuses
                    )
                    cols = ["id","quote_number","client_name","project_name","status",
                            "grand_total","created_at","approval_cycle","last_approver"]
                    rows = await cur.fetchall()
                    quotes = []
                    for r in rows:
                        d = dict(zip(cols, r))
                        if hasattr(d["created_at"], "isoformat"):
                            d["created_at"] = d["created_at"].isoformat()
                        d["grand_total"] = float(d["grand_total"] or 0)
                        quotes.append(d)
            return {"quotes": quotes, "total": len(quotes), "role": role, "data_source": "live"}
        except Exception as exc:
            logger.warning("design_quotes: pending approvals fetch failed — %s", exc)

    # Demo fallback
    _init_demo()
    pending = [
        {"id": k, **{f: v for f, v in q.items() if f in
                    ("quote_number","client_name","project_name","status","grand_total","created_at")},
         "approval_cycle": q.get("approval_cycle", 0), "last_approver": q.get("last_approver", "")}
        for k, q in _demo_quotes.items()
        if q.get("status") in pending_statuses
    ]
    return {"quotes": pending, "total": len(pending), "role": role, "data_source": "demo"}


@router.delete("/design-quotes/{quote_id}")
async def delete_quote(quote_id: int):
    _init_demo()
    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute("DELETE FROM design_quotes WHERE id=%s", (quote_id,))
                    await conn.commit()
            return {"deleted": True, "data_source": "live"}
        except Exception as exc:
            logger.warning("design_quotes: delete DB failed — %s", exc)

    _demo_quotes.pop(quote_id, None)
    return {"deleted": True, "data_source": "demo"}


# ── Route: Merge Quotes ───────────────────────────────────────────────────────

@router.post("/design-quotes/merge")
async def merge_quotes(body: dict):
    """
    Merge 2+ interior quotations into one new DRAFT quote.
    Client details from the first selected quote are used as the base.
    All rooms (sections) from every selected quote are combined.
    """
    _init_demo()
    quote_ids = [int(i) for i in (body.get("quote_ids") or []) if i]
    if len(quote_ids) < 2:
        raise HTTPException(status_code=400, detail="At least 2 quotes required to merge")

    selected: list = []
    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    for qid in quote_ids:
                        await cur.execute("SELECT * FROM design_quotes WHERE id=%s", (qid,))
                        cols = [d[0] for d in cur.description]
                        row = await cur.fetchone()
                        if row:
                            d = dict(zip(cols, row))
                            d["sections"] = json.loads(d.get("sections_json") or "[]")
                            selected.append(d)
        except Exception as exc:
            logger.warning("design_quotes: merge DB fetch failed — %s", exc)
            pool = None

    if not pool:
        for qid in quote_ids:
            q = _demo_quotes.get(qid)
            if q:
                selected.append(q)

    if not selected:
        raise HTTPException(status_code=404, detail="None of the selected quotes were found")

    base = selected[0]
    merged_sections: list = []
    for q in selected:
        for sec in (q.get("sections") or []):
            merged_sections.append({**sec, "section_order": len(merged_sections)})

    source_nums = ", ".join(q.get("quote_number", str(q.get("id", ""))) for q in selected)
    merged_notes = f"Merged from: {source_nums}"
    if base.get("notes"):
        merged_notes = base["notes"] + " · " + merged_notes

    payload = {
        "client_name":       base.get("client_name", ""),
        "client_phone":      base.get("client_phone", ""),
        "client_email":      base.get("client_email", ""),
        "project_name":      f"[Merged] {base.get('project_name', '')}",
        "project_address":   base.get("project_address", ""),
        "project_type":      base.get("project_type", "Residential"),
        "designer_name":     base.get("designer_name", ""),
        "designer_company":  base.get("designer_company", ""),
        "payment_terms":     base.get("payment_terms", ""),
        "validity_days":     int(base.get("validity_days") or 30),
        "gst_rate":          float(base.get("gst_rate") or 18),
        "include_gst":       bool(base.get("include_gst", True)),
        "notes":             merged_notes,
        "terms":             base.get("terms", ""),
        "status":            "DRAFT",
        "sections":          merged_sections,
        "subtotal":          sum(float(q.get("subtotal") or 0) for q in selected),
        "gst_amount":        sum(float(q.get("gst_amount") or 0) for q in selected),
        "grand_total":       sum(float(q.get("grand_total") or 0) for q in selected),
        "total_area_sqft":   sum(float(q.get("total_area_sqft") or 0) for q in selected),
        "margin_mode":       "per_line",
        "overall_margin_pct": 0,
    }

    # Delegate to create_quote logic
    qnum = _next_quote_number()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute("""
                        INSERT INTO design_quotes
                        (quote_number, client_name, client_phone, client_email, project_name,
                         project_address, project_type, designer_name, designer_company,
                         payment_terms, validity_days, gst_rate, include_gst, notes, terms,
                         status, subtotal, gst_amount, grand_total, total_area_sqft, sections_json,
                         margin_mode, overall_margin_pct)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                    """, (
                        qnum,
                        payload["client_name"], payload["client_phone"], payload["client_email"],
                        payload["project_name"], payload["project_address"], payload["project_type"],
                        payload["designer_name"], payload["designer_company"], payload["payment_terms"],
                        payload["validity_days"], payload["gst_rate"], int(payload["include_gst"]),
                        payload["notes"], payload["terms"], "DRAFT",
                        payload["subtotal"], payload["gst_amount"], payload["grand_total"],
                        payload["total_area_sqft"], json.dumps(merged_sections), "per_line", 0,
                    ))
                    new_id = cur.lastrowid
                    await conn.commit()
            return {**payload, "id": new_id, "quote_number": qnum, "created_at": datetime.now().strftime("%Y-%m-%d"), "valid_till": _valid_till(payload["validity_days"]), "data_source": "live"}
        except Exception as exc:
            logger.warning("design_quotes: merge DB save failed — %s", exc)

    # Demo fallback
    new_id = max(_demo_quotes.keys(), default=0) + 1
    merged_quote = {**payload, "id": new_id, "quote_number": qnum, "created_at": datetime.now().strftime("%Y-%m-%d"), "valid_till": _valid_till(payload["validity_days"])}
    _demo_quotes[new_id] = merged_quote
    return merged_quote


# ── Routes: AI Scan ───────────────────────────────────────────────────────────

@router.post("/design-quotes/scan")
async def scan_requirements(
    file: List[UploadFile] = File(default=[]),
    text_input: Optional[str] = Form(default=None),
):
    """Multi-file scanner: accepts any number of images/documents + optional text."""
    if text_input == "__demo__":
        return _demo_scan_result()

    combined_text = text_input or ""
    images: List[dict] = []
    IMAGE_EXTS = (".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp")

    MAX_IMAGE_BYTES = 4 * 1024 * 1024  # 4 MB per image
    doc_meta: list = []  # human-readable info about processed docs (for logging)

    for f in (file or []):
        if not f or not f.filename:
            continue
        raw = await f.read()
        ct  = (f.content_type or "").lower()
        fn  = (f.filename or "").lower()

        if ct.startswith("image/") or fn.endswith(IMAGE_EXTS):
            if len(raw) > MAX_IMAGE_BYTES:
                raise HTTPException(
                    status_code=413,
                    detail=f"Image '{f.filename}' is {len(raw)//1024}KB — maximum is 4 MB. "
                           "Please reduce the image resolution or compress it before uploading.",
                )
            b64 = base64.b64encode(raw).decode()
            images.append({"b64": b64, "ct": ct if ct.startswith("image/") else "image/jpeg"})
        else:
            # Use the structured extractor for PDF / DOCX / XLSX / CSV / text
            doc_text = _extract_document_text(f.filename, raw, f.content_type or "")
            if doc_text and doc_text.strip():
                header = f"\n\n[Document: {f.filename}]\n"
                combined_text = combined_text + header + doc_text
                doc_meta.append(f.filename)
            else:
                logger.warning("design_quotes: could not extract text from '%s'", f.filename)

    if not combined_text.strip() and not images:
        raise HTTPException(status_code=400, detail="Provide at least one file or text_input")

    return await _ai_scan(combined_text, images if images else None)


# ── Routes: Architect Proposals ───────────────────────────────────────────────

@router.get("/design-quotes/architect/proposals")
async def list_proposals(status: Optional[str] = None):
    _init_demo()
    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    q = "SELECT * FROM architect_proposals"
                    params = []
                    if status:
                        q += " WHERE status=%s"; params.append(status)
                    q += " ORDER BY created_at DESC LIMIT 200"
                    await cur.execute(q, params)
                    cols = [d[0] for d in cur.description]
                    rows = await cur.fetchall()
                    proposals = []
                    for row in rows:
                        d = dict(zip(cols, row))
                        d["phases"] = json.loads(d.get("phases_json") or "[]")
                        d["boq"]    = json.loads(d.get("boq_json") or "[]")
                        raw_created = d.get("created_at")
                        validity = int(d.get("validity_days") or 30)
                        d["valid_till"] = (
                            (raw_created + timedelta(days=validity)).strftime("%Y-%m-%d")
                            if hasattr(raw_created, "date") else _valid_till(validity)
                        )
                        d["created_at"] = str(raw_created)[:10] if raw_created else ""
                        proposals.append(d)
            return {"proposals": proposals, "data_source": "live"}
        except Exception as exc:
            logger.warning("design_quotes: proposals list DB failed — %s", exc)

    ps = list(_demo_proposals.values())
    if status:
        ps = [p for p in ps if p["status"] == status]
    return {"proposals": sorted(ps, key=lambda x: x["created_at"], reverse=True), "data_source": "demo"}


@router.get("/design-quotes/architect/proposals/{proposal_id}")
async def get_proposal(proposal_id: int):
    _init_demo()
    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute("SELECT * FROM architect_proposals WHERE id=%s", (proposal_id,))
                    row = await cur.fetchone()
                    if not row:
                        raise HTTPException(status_code=404, detail="Proposal not found")
                    cols = [d[0] for d in cur.description]
                    d = dict(zip(cols, row))
                    d["phases"] = json.loads(d.get("phases_json") or "[]")
                    d["boq"]    = json.loads(d.get("boq_json") or "[]")
                    raw_created = d.get("created_at")
                    validity = int(d.get("validity_days") or 30)
                    d["valid_till"] = (
                        (raw_created + timedelta(days=validity)).strftime("%Y-%m-%d")
                        if hasattr(raw_created, "date") else _valid_till(validity)
                    )
                    d["created_at"] = str(raw_created)[:10] if raw_created else ""
                    return d
        except HTTPException:
            raise
        except Exception as exc:
            logger.warning("design_quotes: get proposal DB failed — %s", exc)

    p = _demo_proposals.get(proposal_id)
    if not p:
        raise HTTPException(status_code=404, detail="Proposal not found")
    return p


@router.post("/design-quotes/architect/proposals")
async def create_proposal(body: dict):
    _init_demo()
    pnum = body.get("proposal_number") or _next_proposal_number()
    validity_days = int(body.get("validity_days") or 30)

    # Re-compute area + fee + phases if dimensions provided
    plot_l = float(body.get("plot_length") or 0)
    plot_w = float(body.get("plot_width") or 0)
    floors = int(body.get("floors") or 1)
    typology = body.get("typology", "villa")
    plot_unit = body.get("plot_unit", "feet")
    fee_model = body.get("fee_model", "percentage")
    fee_rate  = float(body.get("fee_rate") or 5.0)
    constr_cost = float(body.get("construction_cost") or 0)
    complexity = body.get("complexity", "medium")

    areas = _calc_areas(plot_l, plot_w, plot_unit, floors, typology) if plot_l and plot_w else {
        "site_area_sqft": float(body.get("site_area_sqft") or 0),
        "builtup_area_sqft": float(body.get("builtup_area_sqft") or 0),
        "carpet_area_sqft": float(body.get("carpet_area_sqft") or 0),
        "floor_plate_sqft": float(body.get("builtup_area_sqft") or 0) / max(floors, 1),
    }

    total_fee = _calc_fee(fee_model, fee_rate, constr_cost, areas["builtup_area_sqft"])
    phases = body.get("phases") or _default_phases(total_fee)
    boq    = body.get("boq") or (
        _generate_boq(areas, body.get("project_type", "residential"), complexity)
        if areas["builtup_area_sqft"] > 0 else []
    )

    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute("""
                        INSERT INTO architect_proposals
                        (proposal_number, client_name, client_phone, client_email, project_name,
                         project_type, typology, plot_length, plot_width, plot_unit,
                         site_area_sqft, floors, builtup_area_sqft, carpet_area_sqft,
                         fee_model, fee_rate, construction_cost, total_fee, gst_pct,
                         validity_days, notes, status, phases_json, boq_json)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                    """, (
                        pnum,
                        body.get("client_name", ""),   body.get("client_phone", ""),
                        body.get("client_email", ""),   body.get("project_name", ""),
                        body.get("project_type", "residential"), typology,
                        plot_l, plot_w, plot_unit,
                        areas["site_area_sqft"], floors,
                        areas["builtup_area_sqft"], areas["carpet_area_sqft"],
                        fee_model, fee_rate, constr_cost, total_fee,
                        float(body.get("gst_pct", 18)),
                        validity_days, body.get("notes", ""),
                        body.get("status", "DRAFT"),
                        json.dumps(phases), json.dumps(boq),
                    ))
                    new_id = cur.lastrowid
                    await conn.commit()
            return {"id": new_id, "proposal_number": pnum, "total_fee": total_fee, "data_source": "live"}
        except Exception as exc:
            logger.warning("design_quotes: create proposal DB failed — %s", exc)

    new_id = max(_demo_proposals.keys(), default=0) + 1
    _demo_proposals[new_id] = {
        **body, "id": new_id, "proposal_number": pnum,
        "phases": phases, "boq": boq,
        "total_fee": total_fee,
        "site_area_sqft": areas["site_area_sqft"],
        "builtup_area_sqft": areas["builtup_area_sqft"],
        "carpet_area_sqft": areas["carpet_area_sqft"],
        "created_at": datetime.now().strftime("%Y-%m-%d"),
        "valid_till": _valid_till(validity_days),
    }
    return {"id": new_id, "proposal_number": pnum, "total_fee": total_fee, "data_source": "demo"}


@router.put("/design-quotes/architect/proposals/{proposal_id}")
async def update_proposal(proposal_id: int, body: dict):
    _init_demo()
    phases = body.get("phases", [])
    boq    = body.get("boq", [])

    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute("""
                        UPDATE architect_proposals SET
                          client_name=%s, client_phone=%s, client_email=%s, project_name=%s,
                          project_type=%s, typology=%s, site_area_sqft=%s, floors=%s,
                          builtup_area_sqft=%s, carpet_area_sqft=%s, fee_model=%s, fee_rate=%s,
                          construction_cost=%s, total_fee=%s, gst_pct=%s, validity_days=%s,
                          notes=%s, status=%s, phases_json=%s, boq_json=%s
                        WHERE id=%s
                    """, (
                        body.get("client_name",""), body.get("client_phone",""),
                        body.get("client_email",""), body.get("project_name",""),
                        body.get("project_type","residential"), body.get("typology","villa"),
                        float(body.get("site_area_sqft",0)), int(body.get("floors",1)),
                        float(body.get("builtup_area_sqft",0)), float(body.get("carpet_area_sqft",0)),
                        body.get("fee_model","percentage"), float(body.get("fee_rate",5.0)),
                        float(body.get("construction_cost",0)), float(body.get("total_fee",0)),
                        float(body.get("gst_pct",18)), int(body.get("validity_days",30)),
                        body.get("notes",""), body.get("status","DRAFT"),
                        json.dumps(phases), json.dumps(boq),
                        proposal_id,
                    ))
                    await conn.commit()
            return {"id": proposal_id, "data_source": "live"}
        except Exception as exc:
            logger.warning("design_quotes: update proposal DB failed — %s", exc)

    if proposal_id in _demo_proposals:
        _demo_proposals[proposal_id] = {**_demo_proposals[proposal_id], **body, "phases": phases, "boq": boq}
    return {"id": proposal_id, "data_source": "demo"}


@router.put("/design-quotes/architect/proposals/{proposal_id}/status")
async def update_proposal_status(proposal_id: int, body: dict):
    _init_demo()
    status = body.get("status", "DRAFT")

    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute(
                        "UPDATE architect_proposals SET status=%s WHERE id=%s", (status, proposal_id)
                    )
                    await conn.commit()
            return {"id": proposal_id, "status": status, "data_source": "live"}
        except Exception as exc:
            logger.warning("design_quotes: proposal status DB failed — %s", exc)

    if proposal_id in _demo_proposals:
        _demo_proposals[proposal_id]["status"] = status
    return {"id": proposal_id, "status": status, "data_source": "demo"}


@router.delete("/design-quotes/architect/proposals/{proposal_id}")
async def delete_proposal(proposal_id: int):
    _init_demo()
    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute("DELETE FROM architect_proposals WHERE id=%s", (proposal_id,))
                    await conn.commit()
            return {"deleted": True, "data_source": "live"}
        except Exception as exc:
            logger.warning("design_quotes: delete proposal DB failed — %s", exc)

    _demo_proposals.pop(proposal_id, None)
    return {"deleted": True, "data_source": "demo"}


@router.post("/design-quotes/architect/parse-brief")
async def parse_architect_brief(body: dict):
    text = body.get("brief_text", "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="brief_text is required")
    return await _ai_parse_architect(text)


@router.post("/design-quotes/architect/calculate-areas")
async def calculate_areas(body: dict):
    plot_l   = float(body.get("plot_length") or 0)
    plot_w   = float(body.get("plot_width") or 0)
    floors   = int(body.get("floors") or 1)
    typology = body.get("typology", "villa")
    unit     = body.get("plot_unit", "feet")
    if not plot_l or not plot_w:
        raise HTTPException(status_code=400, detail="plot_length and plot_width required")
    return _calc_areas(plot_l, plot_w, unit, floors, typology)


@router.post("/design-quotes/architect/generate-boq")
async def generate_boq(body: dict):
    areas = {
        "site_area_sqft":    float(body.get("site_area_sqft") or 0),
        "builtup_area_sqft": float(body.get("builtup_area_sqft") or 0),
        "carpet_area_sqft":  float(body.get("carpet_area_sqft") or 0),
        "floor_plate_sqft":  float(body.get("floor_plate_sqft") or body.get("builtup_area_sqft") or 0),
    }
    project_type = body.get("project_type", "residential")
    complexity   = body.get("complexity", "medium")
    boq = _generate_boq(areas, project_type, complexity)
    total = sum(it["amount"] for it in boq)
    return {"boq": boq, "total": total}


# ── Route: Clone Design Quote ─────────────────────────────────────────────────

@router.post("/design-quotes/{quote_id}/clone")
async def clone_design_quote(quote_id: int):
    """
    Return a clone-ready payload of an existing design quote.
    Strips id, quote_number, created_at, valid_till.
    Status reset to DRAFT, project_name prefixed [Copy].
    No DB write — frontend opens QuoteFormModal pre-populated.
    """
    _init_demo()
    quote: dict = {}
    pool = await _get_db_pool()
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute("SELECT * FROM design_quotes WHERE id=%s", (quote_id,))
                    cols = [d[0] for d in cur.description]
                    row = await cur.fetchone()
                    if row:
                        d = dict(zip(cols, row))
                        d["sections"] = json.loads(d.get("sections_json") or "[]")
                        quote = d
        except Exception as exc:
            logger.warning("design_quotes: clone DB fetch failed — %s", exc)

    if not quote:
        quote = _demo_quotes.get(quote_id, {})

    if not quote:
        raise HTTPException(status_code=404, detail=f"Design quote {quote_id} not found")

    proj = str(quote.get("project_name") or "")
    prefix = "[Copy] "
    return {
        **{k: v for k, v in quote.items() if k not in ("id", "quote_number", "created_at", "valid_till", "updated_at", "sections_json")},
        "project_name": proj if proj.startswith(prefix) else prefix + proj,
        "status": "DRAFT",
        "sections": quote.get("sections") or [],
    }


# ── Route: Send Design Quote by Email ─────────────────────────────────────────

@router.post("/design-quotes/{quote_id}/send-email")
async def send_design_quote_email(quote_id: int, body: dict):
    recipient_email = (body.get("recipient_email") or "").strip()
    if not recipient_email:
        raise HTTPException(status_code=422, detail="recipient_email is required")

    _init_demo()

    # Fetch quote — DB first, demo fallback
    pool = await _get_db_pool()
    quote = None
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute("SELECT * FROM design_quotes WHERE id=%s", (quote_id,))
                    row = await cur.fetchone()
                    if row:
                        cols = [d[0] for d in cur.description]
                        quote = dict(zip(cols, row))
        except Exception as exc:
            logger.warning("design_quotes: email fetch failed — %s", exc)
    if not quote:
        quote = _demo_quotes.get(quote_id)
    if not quote:
        raise HTTPException(status_code=404, detail="Quote not found")

    recipient_name = (body.get("recipient_name") or quote.get("client_name") or "Client").strip()
    subject = (body.get("subject") or f"Interior Design Quotation — {quote.get('quote_number', '')}").strip()
    custom_msg = (body.get("message") or "").strip()
    quote_number = quote.get("quote_number", "")
    grand_total  = float(quote.get("grand_total") or 0)
    project_name = quote.get("project_name") or ""
    validity     = int(quote.get("validity_days") or 30)

    html_body = f"""
<div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;color:#1f2937">
  <div style="background:linear-gradient(135deg,#7c3aed,#a855f7);padding:24px;border-radius:8px 8px 0 0">
    <h2 style="color:#fff;margin:0;font-size:20px">Interior Design Quotation</h2>
    <p style="color:rgba(255,255,255,.85);margin:6px 0 0;font-size:13px">{quote_number}</p>
  </div>
  <div style="background:#fff;padding:24px;border:1px solid #e5e7eb;border-top:none;border-radius:0 0 8px 8px">
    <p style="margin:0 0 14px">Dear {recipient_name},</p>
    {f'<p style="margin:0 0 14px">{custom_msg}</p>' if custom_msg else '<p style="margin:0 0 14px">Please find your interior design quotation below. We hope it meets your requirements.</p>'}
    <div style="background:#f9fafb;border-radius:8px;padding:16px;margin:16px 0;border:1px solid #e5e7eb">
      <p style="margin:0;font-size:12px;color:#6b7280">Quote Number: <strong style="color:#1f2937">{quote_number}</strong></p>
      {"" if not project_name else f'<p style="margin:6px 0 0;font-size:12px;color:#6b7280">Project: <strong style="color:#1f2937">{project_name}</strong></p>'}
      <p style="margin:12px 0 0;font-size:24px;font-weight:800;color:#7c3aed">&#8377;{grand_total:,.0f}</p>
      <p style="margin:4px 0 0;font-size:11px;color:#9ca3af">Grand Total (incl. GST)</p>
    </div>
    <p style="font-size:12px;color:#9ca3af;margin:0">This quotation is valid for {validity} days. Please contact us if you have any questions or require adjustments.</p>
    <p style="font-size:12px;color:#9ca3af;margin:8px 0 0">Warm regards,<br/><strong style="color:#7c3aed">Design Quote Studio</strong></p>
  </div>
</div>""".strip()

    smtp_user = os.getenv("SMTP_USER", "")
    smtp_pass = os.getenv("SMTP_PASSWORD", "")

    if not smtp_user or not smtp_pass:
        return {
            "success": True, "simulated": True,
            "message": f"Email to {recipient_email} simulated (SMTP not configured). Add SMTP_USER / SMTP_PASSWORD to backend/.env to send real emails.",
            "quote_number": quote_number, "recipient": recipient_email,
        }

    def _send():
        msg = MIMEMultipart("alternative")
        msg["Subject"] = subject
        msg["From"]    = smtp_user
        msg["To"]      = recipient_email
        msg.attach(MIMEText(html_body, "html"))
        smtp_host = os.getenv("SMTP_HOST", "smtp.gmail.com")
        smtp_port = int(os.getenv("SMTP_PORT", "587"))
        with smtplib.SMTP(smtp_host, smtp_port, timeout=20) as s:
            s.ehlo(); s.starttls(); s.login(smtp_user, smtp_pass)
            s.sendmail(smtp_user, recipient_email, msg.as_string())

    try:
        await asyncio.get_event_loop().run_in_executor(None, _send)
        return {"success": True, "simulated": False, "message": f"Email sent to {recipient_email}.", "quote_number": quote_number, "recipient": recipient_email}
    except Exception as exc:
        logger.warning("design_quotes: email send failed — %s", exc)
        raise HTTPException(status_code=502, detail=f"Email delivery failed: {exc}")


# ── Route: Send Architect Fee Proposal by Email ───────────────────────────────

@router.post("/design-quotes/architect/proposals/{proposal_id}/send-email")
async def send_proposal_email(proposal_id: int, body: dict):
    recipient_email = (body.get("recipient_email") or "").strip()
    if not recipient_email:
        raise HTTPException(status_code=422, detail="recipient_email is required")

    _init_demo()

    pool = await _get_db_pool()
    proposal = None
    if pool:
        try:
            async with pool.acquire() as conn:
                async with conn.cursor() as cur:
                    await cur.execute("SELECT * FROM architect_proposals WHERE id=%s", (proposal_id,))
                    row = await cur.fetchone()
                    if row:
                        cols = [d[0] for d in cur.description]
                        proposal = dict(zip(cols, row))
                        proposal["phases"] = json.loads(proposal.get("phases_json") or "[]")
        except Exception as exc:
            logger.warning("design_quotes: proposal email fetch failed — %s", exc)
    if not proposal:
        proposal = _demo_proposals.get(proposal_id)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    recipient_name  = (body.get("recipient_name") or proposal.get("client_name") or "Client").strip()
    subject         = (body.get("subject") or f"Architect Fee Proposal — {proposal.get('proposal_number', '')}").strip()
    custom_msg      = (body.get("message") or "").strip()
    proposal_number = proposal.get("proposal_number", "")
    total_fee       = float(proposal.get("total_fee") or 0)
    project_name    = proposal.get("project_name") or ""
    fee_model       = proposal.get("fee_model", "percentage")
    fee_rate        = float(proposal.get("fee_rate") or 0)
    gst_pct         = float(proposal.get("gst_pct") or 18)
    validity        = int(proposal.get("validity_days") or 30)
    phases          = proposal.get("phases") or []

    fee_desc = (f"{fee_rate}% of construction cost" if fee_model == "percentage"
                else f"₹{fee_rate:,.0f}/sqft built-up" if fee_model == "per_sqft"
                else f"Lump sum ₹{fee_rate:,.0f}")

    phase_rows = "".join(
        f"""<tr>
          <td style="padding:7px 10px;border-bottom:1px solid #e5e7eb">{p.get('phase_name','')}</td>
          <td style="padding:7px 10px;border-bottom:1px solid #e5e7eb;text-align:center">{p.get('pct_of_total',0)}%</td>
          <td style="padding:7px 10px;border-bottom:1px solid #e5e7eb;text-align:right;font-weight:600">
            &#8377;{float(p.get('fee_amount',0)):,.0f}
          </td>
          <td style="padding:7px 10px;border-bottom:1px solid #e5e7eb;text-align:center">
            {p.get('due_date','—')}
          </td>
          <td style="padding:7px 10px;border-bottom:1px solid #e5e7eb;text-align:center">
            {'<span style=\"color:#16a34a;font-weight:700\">✓ Paid</span>' if p.get('is_paid') else '<span style=\"color:#d97706\">Pending</span>'}
          </td>
        </tr>"""
        for p in phases
    )

    html_body = f"""
<div style="font-family:Arial,sans-serif;max-width:620px;margin:0 auto;color:#1f2937">
  <div style="background:linear-gradient(135deg,#6366f1,#818cf8);padding:24px;border-radius:8px 8px 0 0">
    <div style="display:flex;align-items:center;gap:12px">
      <div style="width:40px;height:40px;background:rgba(255,255,255,.25);border-radius:8px;display:flex;align-items:center;justify-content:center;color:#fff;font-weight:900;font-size:18px">IQ</div>
      <div>
        <div style="color:#fff;font-weight:800;font-size:16px">Architect Fee Proposal</div>
        <div style="color:rgba(255,255,255,.8);font-size:12px">{proposal_number}</div>
      </div>
    </div>
  </div>
  <div style="background:#fff;padding:24px;border:1px solid #e5e7eb;border-top:none;border-radius:0 0 8px 8px">
    <p style="margin:0 0 14px">Dear {recipient_name},</p>
    {f'<p style="margin:0 0 14px">{custom_msg}</p>' if custom_msg else
     '<p style="margin:0 0 14px">Please find our architect fee proposal for your project below.</p>'}
    <div style="background:#f9fafb;border-radius:8px;padding:16px;margin:16px 0;border:1px solid #e5e7eb">
      <p style="margin:0;font-size:12px;color:#6b7280">Proposal: <strong style="color:#1f2937">{proposal_number}</strong></p>
      {f'<p style="margin:6px 0 0;font-size:12px;color:#6b7280">Project: <strong style="color:#1f2937">{project_name}</strong></p>' if project_name else ''}
      <p style="margin:8px 0 0;font-size:12px;color:#6b7280">Fee Basis: <strong style="color:#1f2937">{fee_desc}</strong></p>
      <p style="margin:12px 0 0;font-size:26px;font-weight:900;color:#6366f1">&#8377;{total_fee:,.0f}</p>
      <p style="margin:2px 0 0;font-size:11px;color:#9ca3af">Total Architect Fee (excl. GST {gst_pct:.0f}%)</p>
    </div>
    {f'''<table style="width:100%;border-collapse:collapse;font-size:12px;margin:16px 0">
      <thead>
        <tr style="background:#6366f1;color:#fff">
          <th style="padding:8px 10px;text-align:left;font-weight:700">Phase</th>
          <th style="padding:8px 10px;text-align:center;font-weight:700">%</th>
          <th style="padding:8px 10px;text-align:right;font-weight:700">Fee (₹)</th>
          <th style="padding:8px 10px;text-align:center;font-weight:700">Due Date</th>
          <th style="padding:8px 10px;text-align:center;font-weight:700">Status</th>
        </tr>
      </thead>
      <tbody>{phase_rows}</tbody>
    </table>''' if phase_rows else ''}
    <p style="font-size:12px;color:#9ca3af;margin:16px 0 0">
      This proposal is valid for {validity} days. GST @ {gst_pct:.0f}% applicable additionally.
      Please contact us if you have any questions.
    </p>
    <p style="font-size:12px;color:#9ca3af;margin:8px 0 0">
      Warm regards,<br/><strong style="color:#6366f1">Design Quote Studio · InvenIQ</strong>
    </p>
  </div>
</div>""".strip()

    smtp_user = os.getenv("SMTP_USER", "")
    smtp_pass = os.getenv("SMTP_PASSWORD", "")

    if not smtp_user or not smtp_pass:
        return {
            "success": True, "simulated": True,
            "message": f"Email to {recipient_email} simulated (SMTP not configured). Add SMTP_USER / SMTP_PASSWORD to backend/.env.",
            "proposal_number": proposal_number, "recipient": recipient_email,
        }

    def _send():
        msg = MIMEMultipart("alternative")
        msg["Subject"] = subject
        msg["From"]    = smtp_user
        msg["To"]      = recipient_email
        msg.attach(MIMEText(html_body, "html"))
        smtp_host = os.getenv("SMTP_HOST", "smtp.gmail.com")
        smtp_port = int(os.getenv("SMTP_PORT", "587"))
        with smtplib.SMTP(smtp_host, smtp_port, timeout=20) as s:
            s.ehlo(); s.starttls(); s.login(smtp_user, smtp_pass)
            s.sendmail(smtp_user, recipient_email, msg.as_string())

    try:
        await asyncio.get_event_loop().run_in_executor(None, _send)
        return {"success": True, "simulated": False, "message": f"Email sent to {recipient_email}.", "proposal_number": proposal_number, "recipient": recipient_email}
    except Exception as exc:
        logger.warning("design_quotes: proposal email failed — %s", exc)
        raise HTTPException(status_code=502, detail=f"Email delivery failed: {exc}")
