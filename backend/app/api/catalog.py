"""
Product Catalog API — InvenIQ
Comprehensive product database covering Louvers, Laminates, PVC, Building Materials,
and Furniture Hardware (Ebco, Hafele, Hettich and more).
Supports AI-powered product extraction from images, PDFs, spreadsheets, or typed text.
New products can be scanned and added to the live catalog without restart.
"""
import base64
import io as _io
import json
import logging
import os
from typing import List, Optional

from fastapi import APIRouter, File, Form, HTTPException, UploadFile

logger = logging.getLogger(__name__)
router = APIRouter(tags=["Product Catalog"])

# ── Runtime product additions (session-persistent; DB-persistent with MySQL) ──
_RUNTIME_PRODUCTS: List[dict] = []

try:
    from app.api.ebco_catalog import EBCO_CATALOG as _EBCO_CATALOG
except ImportError:
    _EBCO_CATALOG = []


def _get_all_products() -> List[dict]:
    return CATALOG + _EBCO_CATALOG + _RUNTIME_PRODUCTS


def _next_id() -> int:
    ids = [p["product_id"] for p in _get_all_products()]
    return max(ids, default=199) + 1


# ── Full product catalog ───────────────────────────────────────────────────────

CATALOG = [
    # ── HIGH PRESSURE LAMINATES ────────────────────────────────────────────────
    {
        "product_id": 101, "sku_code": "HPL-1MM-MATTE",
        "name": "HPL 1mm Matte / Suede",
        "brand": "Greenlam / Merino / Century",
        "category": "High Pressure Laminate", "sub_category": "Standard HPL",
        "unit": "sheet", "size": "8×4 ft (2440×1220 mm)", "thickness": "1.0 mm",
        "finish": "Matte, Suede, Smooth", "weight_kg": 3.2,
        "colors": "200+ shades — solid colours, wood grain, stone, metallic",
        "buy_price": 950, "sell_price": 1300, "margin_pct": 26.9,
        "gst_rate": 18.0, "hsn_code": "4814",
        "applications": [
            "Kitchen cabinet shutters", "Wardrobe panels & sliding doors",
            "Office furniture & workstations", "Wall cladding (interior)",
            "Reception counters", "Retail display fixtures",
        ],
        "certifications": ["IS 2046:2019", "E1 Formaldehyde emission", "Fire retardant grade (FR) available", "CARB Phase 2 compliant"],
        "features": [
            "Abrasion resistant — EN 438 Grade P", "Moisture & steam resistant",
            "Anti-fingerprint surface option", "25-year product life", "Easy to clean & maintain",
            "Impact & scratch resistant",
        ],
        "installation_tips": "Use white PVC edgeband to match. Recommended adhesive: contact cement or PU adhesive.",
        "lead_time": "5–7 days", "moq": 20, "stock_status": "in_stock",
        "tags": ["popular", "furniture", "interior", "kitchen"],
        "competitors": ["Formica", "Durian", "Stylam", "Action TESA"],
    },
    {
        "product_id": 102, "sku_code": "HPL-1.5MM-MATTE",
        "name": "HPL 1.5mm Matte (Post-form)",
        "brand": "Greenlam / Merino",
        "category": "High Pressure Laminate", "sub_category": "Post-Form HPL",
        "unit": "sheet", "size": "8×4 ft", "thickness": "1.5 mm",
        "finish": "Matte, Gloss", "weight_kg": 4.8,
        "colors": "100+ shades",
        "buy_price": 1250, "sell_price": 1680, "margin_pct": 25.6,
        "gst_rate": 18.0, "hsn_code": "4814",
        "applications": [
            "Post-formed countertops", "Round-edge furniture", "Curved doors",
            "Lab furniture", "Hospital trolleys",
        ],
        "certifications": ["IS 2046:2019", "Post-formable grade certified"],
        "features": [
            "Can be heat-bent to radius ≥ 40mm", "Higher bending strength",
            "Moisture resistant", "Hygienic surface",
        ],
        "installation_tips": "Heat to 160°C, bend over jig while warm. Cool under clamps.",
        "lead_time": "7–10 days", "moq": 15, "stock_status": "in_stock",
        "tags": ["post-form", "countertop", "curved"],
        "competitors": ["Formica", "Wilsonart"],
    },
    {
        "product_id": 103, "sku_code": "HPL-COMPACT-6MM",
        "name": "HPL Compact Laminate 6mm",
        "brand": "Greenlam / Stylam",
        "category": "Compact Laminate", "sub_category": "Compact HPL",
        "unit": "sheet", "size": "8×4 ft (2440×1220 mm)", "thickness": "6 mm",
        "finish": "Matte both sides", "weight_kg": 19.2,
        "colors": "30+ solid shades, wood grain",
        "buy_price": 2980, "sell_price": 3600, "margin_pct": 17.2,
        "gst_rate": 18.0, "hsn_code": "4814",
        "applications": [
            "Toilet cubicle partitions & doors", "Exterior cladding & facades",
            "Wet area wall panels", "Swimming pool surrounds",
            "Laboratory benchtops", "Hospital & cleanroom panels",
            "Locker room dividers",
        ],
        "certifications": [
            "IS 2046:2019", "Moisture & fungal resistant", "Fire retardant (EN 13501-1 Class B)",
            "UV resistant", "Anti-graffiti surface available",
        ],
        "features": [
            "Phenolic core — no substrate needed", "Both faces finished",
            "Waterproof & humid area suitable", "Impact & vandal resistant",
            "No swelling, no delamination", "10-year warranty",
        ],
        "installation_tips": "Use stainless steel or aluminium fittings. 3mm expansion gap on all edges. Silicone seal perimeter.",
        "lead_time": "7–10 days", "moq": 10, "stock_status": "in_stock",
        "tags": ["compact", "toilet-cubicle", "exterior", "wet-area"],
        "competitors": ["Pfleiderer", "Kronospan", "ABET Laminati"],
    },
    {
        "product_id": 104, "sku_code": "HPL-COMPACT-12MM",
        "name": "HPL Compact Laminate 12mm",
        "brand": "Greenlam / Stylam",
        "category": "Compact Laminate", "sub_category": "Compact HPL",
        "unit": "sheet", "size": "8×4 ft", "thickness": "12 mm",
        "finish": "Matte both sides", "weight_kg": 38.4,
        "colors": "20+ solid shades",
        "buy_price": 5800, "sell_price": 7200, "margin_pct": 19.4,
        "gst_rate": 18.0, "hsn_code": "4814",
        "applications": [
            "Full-height toilet cubicle doors (standalone, no frame)",
            "Structural exterior panels", "Benchtop slabs",
            "Interior wall systems", "High-traffic partitions",
        ],
        "certifications": ["IS 2046:2019", "Structural grade", "Fire Class B-s1-d0"],
        "features": [
            "Self-supporting structural panels", "Both faces finished HPL",
            "Maximum impact resistance", "Hygienic — bacteria inhibiting surface",
        ],
        "installation_tips": "Bottom gap 20mm. Use stainless M6 bolts. No silicone on top.",
        "lead_time": "10–14 days", "moq": 5, "stock_status": "in_stock",
        "tags": ["compact", "structural", "toilet-cubicle"],
        "competitors": ["Pfleiderer", "ABET Laminati"],
    },
    {
        "product_id": 105, "sku_code": "ACRYLIC-LAM-84",
        "name": "Acrylic Laminate High-Gloss",
        "brand": "Durian / Generic",
        "category": "Acrylic Laminate", "sub_category": "High-Gloss",
        "unit": "sheet", "size": "8×4 ft (2440×1220 mm)", "thickness": "1.0 mm",
        "finish": "High-Gloss, Super-Gloss, Pearl, Metallic", "weight_kg": 2.8,
        "colors": "80+ shades — solid, metallic, pearlescent",
        "buy_price": 1720, "sell_price": 2100, "margin_pct": 18.1,
        "gst_rate": 18.0, "hsn_code": "3921",
        "applications": [
            "Modular kitchen shutter fronts", "Wardrobe mirror-finish shutters",
            "Retail display counters", "Luxury furniture", "Show flat interiors",
            "Hospitality & hotel furniture",
        ],
        "certifications": ["Anti-scratch coating", "UV resistant", "Food-safe surface"],
        "features": [
            "Piano-gloss mirror finish", "Anti-fingerprint coating available",
            "Scratch & impact resistant acrylic layer", "No fading under UV",
            "Easy wipe-clean maintenance",
        ],
        "installation_tips": "Apply with PU glue only. Avoid contact cement. Use edge profile to finish. Handle with cotton gloves.",
        "lead_time": "5–7 days", "moq": 20, "stock_status": "in_stock",
        "tags": ["acrylic", "high-gloss", "kitchen", "luxury"],
        "competitors": ["Merino Acrylic", "Action TESA", "Greenlam Acrylic"],
    },
    {
        "product_id": 106, "sku_code": "PVC-LAM-GLOSS",
        "name": "PVC Laminate (Solid Colour)",
        "brand": "Various",
        "category": "PVC Laminate", "sub_category": "Solid PVC",
        "unit": "sheet", "size": "8×4 ft", "thickness": "0.8 mm",
        "finish": "Gloss, Matte, Texture", "weight_kg": 1.9,
        "colors": "150+ colours",
        "buy_price": 380, "sell_price": 580, "margin_pct": 34.5,
        "gst_rate": 18.0, "hsn_code": "3921",
        "applications": [
            "Budget furniture overlays", "Modular kitchen (economy range)",
            "Office partitions", "Retail shelving",
        ],
        "certifications": ["Moisture resistant", "IS 12823 compliant"],
        "features": [
            "Economy alternative to HPL", "Easy to cut & apply",
            "Wide colour range", "Gloss or matte variants",
        ],
        "installation_tips": "Use contact cement (fevicol or equivalent). Press with roller. Trim with router.",
        "lead_time": "2–3 days", "moq": 50, "stock_status": "in_stock",
        "tags": ["pvc", "economy", "laminate", "furniture"],
        "competitors": ["Nilkamal", "Supreme"],
    },
    {
        "product_id": 107, "sku_code": "PVC-WOOD-GRAIN",
        "name": "PVC Wood-Grain Laminate",
        "brand": "Various",
        "category": "PVC Laminate", "sub_category": "Wood-Grain PVC",
        "unit": "sheet", "size": "8×4 ft", "thickness": "0.8 mm",
        "finish": "Wood grain texture", "weight_kg": 1.9,
        "colors": "50+ wood species — teak, wenge, walnut, oak, maple",
        "buy_price": 420, "sell_price": 640, "margin_pct": 34.4,
        "gst_rate": 18.0, "hsn_code": "3921",
        "applications": ["Furniture overlays", "Door skins", "Cabinet interiors"],
        "certifications": ["Moisture resistant"],
        "features": ["Embossed wood texture", "Realistic wood grain", "Cost-effective wood look"],
        "installation_tips": "Same as solid PVC. Apply uniform pressure with J-roller.",
        "lead_time": "2–3 days", "moq": 50, "stock_status": "in_stock",
        "tags": ["pvc", "wood-grain", "furniture"],
        "competitors": ["Nilkamal"],
    },
    # ── ALUMINIUM LOUVERS ──────────────────────────────────────────────────────
    {
        "product_id": 108, "sku_code": "LOUV-ALU-Z100-ANOD",
        "name": "Aluminium Z-Profile 100mm Anodized",
        "brand": "Supreme Profile / Jindal Aluminium",
        "category": "Aluminium Louvers", "sub_category": "Fixed Blade",
        "unit": "RM", "size": "Width: 100mm | Thickness: 2.0mm",
        "finish": "Anodized — Silver, Bronze, Gold, Black",
        "weight_kg": 1.8,
        "colors": "Anodized: Natural Silver, Bronze, Champagne, Dark Anodized",
        "buy_price": 1720, "sell_price": 2100, "margin_pct": 18.1,
        "gst_rate": 18.0, "hsn_code": "7604",
        "applications": [
            "Building facade sun shading", "Ventilation screens & louvre walls",
            "Car park screening (all levels)", "Plant room enclosures",
            "Industrial screening", "Commercial & residential balcony privacy",
        ],
        "certifications": [
            "AA-25 Anodizing (25 micron) — IS 1868", "QUALICOAT certification",
            "Wind load tested to IS 875 Part 3", "Aluminium Alloy 6063-T5 / 6061-T6",
        ],
        "features": [
            "Maintenance-free for 20+ years", "No painting required",
            "Corrosion resistant — coastal environments",
            "Custom blade angle (0°–90°)", "Concealed fixing system",
            "Structural aluminium — self-supporting",
        ],
        "installation_tips": "Fix to structural sub-frame at 600mm centres. Use stainless M8 bolts. Thermal break required for A/C-rated facades.",
        "lead_time": "8–12 days", "moq": 100, "stock_status": "in_stock",
        "tags": ["aluminium", "louver", "facade", "anodized", "commercial"],
        "competitors": ["Technal", "Alufit", "YKK AP"],
    },
    {
        "product_id": 109, "sku_code": "LOUV-ALU-Z80-PC",
        "name": "Aluminium Z-Profile 80mm Powder Coated",
        "brand": "Aluline / Alufit",
        "category": "Aluminium Louvers", "sub_category": "Fixed Blade",
        "unit": "RM", "size": "Width: 80mm | Thickness: 1.8mm",
        "finish": "Powder Coated — any RAL colour",
        "weight_kg": 1.4,
        "colors": "Full RAL range (3000+ colours) — custom colour on request",
        "buy_price": 1350, "sell_price": 1680, "margin_pct": 19.6,
        "gst_rate": 18.0, "hsn_code": "7604",
        "applications": [
            "Office building facades", "Interior partitions & screens",
            "Staircase balustrades", "Hotel balcony privacy screens",
            "Retail facade feature elements", "Residential gate screens",
        ],
        "certifications": [
            "PVDF coating — EN 13523", "RAL colour certified",
            "Aluminium Alloy 6063-T5", "IS 875 wind load compliance",
        ],
        "features": [
            "PVDF powder coat — 15-year colour guarantee", "Any RAL colour",
            "Lightweight — easier installation", "Powder coat over anodized base",
            "Fire rated option (15/30 min) available",
        ],
        "installation_tips": "Pre-drill holes at 800mm centres. Use neoprene washers to prevent galvanic corrosion with steel.",
        "lead_time": "10–14 days", "moq": 80, "stock_status": "in_stock",
        "tags": ["aluminium", "louver", "powder-coat", "RAL", "commercial"],
        "competitors": ["Technal", "Alufit"],
    },
    {
        "product_id": 110, "sku_code": "LOUV-ALU-C150-ANOD",
        "name": "Aluminium C-Profile 150mm (Aerofoil Blade)",
        "brand": "Jindal / Supreme Profile",
        "category": "Aluminium Louvers", "sub_category": "Aerofoil Blade",
        "unit": "RM", "size": "Width: 150mm | Chord: 2.5mm",
        "finish": "Anodized or Powder Coated",
        "weight_kg": 2.4,
        "colors": "Anodized or RAL powder coat",
        "buy_price": 2800, "sell_price": 3500, "margin_pct": 20.0,
        "gst_rate": 18.0, "hsn_code": "7604",
        "applications": [
            "Premium facade systems", "High-wind zone buildings",
            "Airports & transportation hubs", "Hospitals & institutions",
            "Luxury residential towers",
        ],
        "certifications": ["CWCT tested", "Wind zone 4 (IS 875)", "Structural aerofoil grade"],
        "features": [
            "Aerofoil cross-section — superior strength", "Wider blade — more privacy + shading",
            "Low wind resistance design", "Suitable for high-rise buildings",
            "Concealed rod-and-bracket system",
        ],
        "installation_tips": "Structural engineer sign-off required for buildings >15m. Stainless M10 fixings.",
        "lead_time": "14–21 days", "moq": 50, "stock_status": "on_order",
        "tags": ["aluminium", "louver", "aerofoil", "premium", "facade"],
        "competitors": ["Technal", "YKK AP"],
    },
    # ── PVC LOUVERS ───────────────────────────────────────────────────────────
    {
        "product_id": 111, "sku_code": "LOUV-PVC-100",
        "name": "PVC Louver Blades 100mm",
        "brand": "Polycab / Supreme / Coltors",
        "category": "PVC Louvers", "sub_category": "Fixed Blade",
        "unit": "RM", "size": "Width: 100mm | Thickness: 3.0mm",
        "finish": "Smooth — White, Ivory, Grey, Black",
        "weight_kg": 0.38,
        "colors": "White, Ivory, Cream, Light Grey, Dark Grey, Black — custom colour available",
        "buy_price": 390, "sell_price": 580, "margin_pct": 32.8,
        "gst_rate": 18.0, "hsn_code": "3925",
        "applications": [
            "Residential window & door screens", "Car park screening (budget)",
            "Building ventilation grilles", "Garden & compound walls",
            "HVAC equipment screening", "Residential privacy screens",
        ],
        "certifications": [
            "UV stabilised — EN 13049", "10-year UV warranty",
            "Self-extinguishing (V-2 rating)", "IS 12823 compliant",
        ],
        "features": [
            "UV stabilised — no yellowing for 10+ years", "Lightweight — easy DIY installation",
            "Termite & rust proof", "Maintenance free", "Low cost solution",
            "Cut on site with standard hacksaw",
        ],
        "installation_tips": "Fix to aluminium or GI channel tracks. Use UV-resistant screws. Expansion gap 3mm per metre.",
        "lead_time": "3–5 days", "moq": 150, "stock_status": "in_stock",
        "tags": ["pvc", "louver", "budget", "residential", "car-park"],
        "competitors": ["Hindustan Profiles", "Sintex"],
    },
    {
        "product_id": 112, "sku_code": "LOUV-PVC-75-WIDE",
        "name": "PVC Louver Blades 75mm (Slim)",
        "brand": "Polycab / Coltors",
        "category": "PVC Louvers", "sub_category": "Slim Blade",
        "unit": "RM", "size": "Width: 75mm | Thickness: 2.5mm",
        "finish": "Smooth — White, Grey", "weight_kg": 0.28,
        "colors": "White, Light Grey",
        "buy_price": 280, "sell_price": 430, "margin_pct": 34.9,
        "gst_rate": 18.0, "hsn_code": "3925",
        "applications": ["Window louvres", "Ventilation panels", "Residential gates"],
        "certifications": ["UV stabilised", "Self-extinguishing"],
        "features": ["Slimmer profile for smaller openings", "Lightweight", "Easy cut & fit"],
        "installation_tips": "Same as 100mm. Use narrower tracks to match 75mm width.",
        "lead_time": "3–5 days", "moq": 200, "stock_status": "in_stock",
        "tags": ["pvc", "louver", "slim", "residential"],
        "competitors": ["Hindustan Profiles"],
    },
    # ── OPERABLE LOUVRE SYSTEMS ────────────────────────────────────────────────
    {
        "product_id": 113, "sku_code": "LOUV-OPS-MTR",
        "name": "Operable Louvre System — Motorised",
        "brand": "Technal / Somfy / YKK AP",
        "category": "Operable Louvre System", "sub_category": "Motorised",
        "unit": "SQM", "size": "Custom — up to 8m span",
        "finish": "Powder coated — any RAL colour",
        "weight_kg": 28,
        "colors": "Full RAL range",
        "buy_price": 9200, "sell_price": 12000, "margin_pct": 23.3,
        "gst_rate": 18.0, "hsn_code": "8302",
        "applications": [
            "Rooftop pergolas & terrace covers", "Commercial atrium roofs",
            "Restaurant alfresco dining roofs", "Hotel poolside covers",
            "Residential roof gardens", "Event venue retractable roofs",
        ],
        "certifications": [
            "Somfy motor — CE certified, IP54", "Wind speed rated to 120 km/h",
            "Rain sensor auto-close option", "5-year system warranty",
            "IS 875 Part 3 wind compliance",
        ],
        "features": [
            "Remote control / smartphone app (SOMFY TaHoma)",
            "Auto-close on rain / wind sensor",
            "0°–90° blade rotation", "Concealed rainwater drainage",
            "LED strip lighting integration option",
            "Full open = 80% ventilation; full close = watertight",
        ],
        "installation_tips": "Structural engineer assessment required. Minimum beam depth 200mm. Allow 3 weeks for fabrication + 3 days install.",
        "lead_time": "21–30 days", "moq": 8, "stock_status": "on_order",
        "tags": ["operable", "motorised", "pergola", "premium", "somfy"],
        "competitors": ["Vergola", "Louvretec", "Eco Awnings"],
    },
    {
        "product_id": 114, "sku_code": "LOUV-OPS-MAN",
        "name": "Operable Louvre System — Manual Crank",
        "brand": "Generic / Alufit",
        "category": "Operable Louvre System", "sub_category": "Manual",
        "unit": "SQM", "size": "Custom — up to 5m span",
        "finish": "Powder coated",
        "weight_kg": 22,
        "colors": "RAL colour on request",
        "buy_price": 5800, "sell_price": 8000, "margin_pct": 27.5,
        "gst_rate": 18.0, "hsn_code": "8302",
        "applications": [
            "Residential pergolas & gazebos", "Small commercial covered areas",
            "Budget hospitality projects",
        ],
        "certifications": ["Wind rated to 80 km/h", "3-year system warranty"],
        "features": [
            "Manual hand-crank operation", "0°–90° blade rotation",
            "Rainwater drainage channels", "No electricity required",
        ],
        "installation_tips": "Suitable for spans up to 5m. Upgrade to motorised for larger projects.",
        "lead_time": "14–18 days", "moq": 10, "stock_status": "in_stock",
        "tags": ["operable", "manual", "pergola", "residential"],
        "competitors": ["Vergola"],
    },
    # ── EXTERIOR CLADDING ──────────────────────────────────────────────────────
    {
        "product_id": 115, "sku_code": "HPL-EXT-CLADDING-6MM",
        "name": "HPL Exterior Cladding 6mm (Facade)",
        "brand": "Greenlam Clad / Trespa",
        "category": "Exterior Cladding", "sub_category": "HPL Facade",
        "unit": "sheet", "size": "8×4 ft or custom", "thickness": "6 mm",
        "finish": "Weather-proof matte / wood grain",
        "weight_kg": 19.2,
        "colors": "50+ exterior-grade colours, wood grains, stone",
        "buy_price": 3800, "sell_price": 4800, "margin_pct": 20.8,
        "gst_rate": 18.0, "hsn_code": "4814",
        "applications": [
            "Building exterior facade cladding", "Balcony & terrace panels",
            "Column cladding", "Signage backing panels",
            "Wet area exterior walls",
        ],
        "certifications": [
            "Weather resistant — EN 438-6", "UV & rain tested",
            "Fire Class A2 option available", "15-year exterior warranty",
        ],
        "features": [
            "100% waterproof phenolic core", "Zero maintenance",
            "Will not rot, rust or corrode", "Fade-resistant pigments",
            "Custom cut sizes available",
        ],
        "installation_tips": "Ventilated facade system — 25mm cavity minimum. Stainless T-shaped clips. Expansion gap 5mm per sheet.",
        "lead_time": "10–14 days", "moq": 10, "stock_status": "in_stock",
        "tags": ["exterior", "cladding", "facade", "waterproof"],
        "competitors": ["Trespa", "Panolam", "Vivalda"],
    },
    {
        "product_id": 116, "sku_code": "ACP-3MM-PE",
        "name": "Aluminium Composite Panel 3mm (PE Core)",
        "brand": "Alucobond / Alumax / Alucoil",
        "category": "Aluminium Composite Panel", "sub_category": "Standard ACP",
        "unit": "sheet", "size": "8×4 ft (2440×1220 mm)", "thickness": "3 mm",
        "finish": "PVDF / PE coated — smooth",
        "weight_kg": 6.8,
        "colors": "Full RAL + special metallic & mirror finishes",
        "buy_price": 1200, "sell_price": 1600, "margin_pct": 25.0,
        "gst_rate": 18.0, "hsn_code": "7606",
        "applications": [
            "Shop front signage & fascias", "Interior wall cladding",
            "Partition panels", "Lift interiors",
            "Furniture & display stands",
        ],
        "certifications": ["IS 2553", "PE/PVDF coating standard"],
        "features": [
            "Lightweight — easy handling & cutting", "Flatness — no warping",
            "Wide colour range", "Easy to fabricate (CNC routing, bending)",
        ],
        "installation_tips": "Use router to groove and fold. Rivets or screws for mounting. Use compatible sealant.",
        "lead_time": "5–7 days", "moq": 20, "stock_status": "in_stock",
        "tags": ["acp", "aluminium-composite", "signage", "cladding"],
        "competitors": ["Alucobond", "Alucoil", "Viva"],
    },
    {
        "product_id": 117, "sku_code": "ACP-4MM-FR",
        "name": "Aluminium Composite Panel 4mm (FR Core)",
        "brand": "Alucobond / Alumax",
        "category": "Aluminium Composite Panel", "sub_category": "Fire Rated ACP",
        "unit": "sheet", "size": "8×4 ft", "thickness": "4 mm",
        "finish": "PVDF coated", "weight_kg": 8.2,
        "colors": "Full RAL range",
        "buy_price": 1800, "sell_price": 2400, "margin_pct": 25.0,
        "gst_rate": 18.0, "hsn_code": "7606",
        "applications": [
            "Exterior building cladding (mandatory fire rated)", "High-rise facades",
            "Airport & hospital cladding",
        ],
        "certifications": ["EN 13501-1 Class B-s1-d0", "Fire rated core", "BIS approved"],
        "features": [
            "Fire resistant mineral core", "Meets NBC 2016 fire code",
            "PVDF 70% coating — 10-year colour warranty",
        ],
        "installation_tips": "Compulsory for buildings > G+4 as per NBC 2016 clause 4.2.",
        "lead_time": "7–10 days", "moq": 15, "stock_status": "in_stock",
        "tags": ["acp", "fire-rated", "exterior", "high-rise"],
        "competitors": ["Alucobond FR", "Alpolic FR"],
    },
    # ── FURNITURE HARDWARE — EBCO ─────────────────────────────────────────────
    {
        "product_id": 118, "sku_code": "EBCO-DS-350",
        "name": "Ebco Telescopic Drawer Slide 350mm",
        "brand": "Ebco",
        "category": "Drawer Slides", "sub_category": "Telescopic",
        "unit": "pair", "size": "350mm (14 inch) | Channel height: 17mm",
        "thickness": "17mm channel height",
        "finish": "Zinc plated / White epoxy",
        "weight_kg": 0.38,
        "colors": "Zinc plated, White",
        "buy_price": 110, "sell_price": 165, "margin_pct": 33.3,
        "gst_rate": 18.0, "hsn_code": "8302",
        "applications": [
            "Kitchen drawer boxes", "Office furniture drawers",
            "Wardrobe pull-out trays", "TV unit drawers", "Study furniture drawers",
        ],
        "certifications": ["Load rated 30 kg per pair", "IS 4992 compliant"],
        "features": [
            "3/4 extension — smooth ball-bearing operation",
            "Load capacity 30 kg per pair", "Easy bottom-mount installation",
            "Self-closing feature available",
        ],
        "installation_tips": "Fix inner runner to drawer box bottom. Fix outer runner to cabinet side. Ensure level alignment.",
        "lead_time": "2–3 days", "moq": 10, "stock_status": "in_stock",
        "tags": ["ebco", "drawer-slide", "hardware", "furniture", "kitchen"],
        "competitors": ["Hafele", "Hettich", "Blum"],
    },
    {
        "product_id": 119, "sku_code": "EBCO-DS-SC400",
        "name": "Ebco Soft-Close Full Extension Drawer Slide 400mm",
        "brand": "Ebco",
        "category": "Drawer Slides", "sub_category": "Soft-Close Full Extension",
        "unit": "pair", "size": "400mm (16 inch) | Full extension",
        "thickness": "17mm channel height",
        "finish": "Zinc plated / Black",
        "weight_kg": 0.52,
        "colors": "Zinc plated, Black",
        "buy_price": 265, "sell_price": 400, "margin_pct": 33.8,
        "gst_rate": 18.0, "hsn_code": "8302",
        "applications": [
            "Premium modular kitchen drawers", "Bedroom wardrobe pull-outs",
            "Office pedestal drawers", "Study table drawers",
        ],
        "certifications": ["Load rated 45 kg", "50,000 cycle tested"],
        "features": [
            "100% full extension — complete drawer access",
            "Integrated soft-close hydraulic damper",
            "Silent & smooth operation", "Under-mount option available",
        ],
        "installation_tips": "Fix inner runner to drawer box side. Outer runner to cabinet side. Align level before tightening.",
        "lead_time": "2–3 days", "moq": 5, "stock_status": "in_stock",
        "tags": ["ebco", "drawer-slide", "soft-close", "full-extension", "premium", "kitchen"],
        "competitors": ["Blum Tandem", "Hettich Arena Plus", "Hafele"],
    },
    {
        "product_id": 120, "sku_code": "EBCO-HNG-35STD",
        "name": "Ebco Concealed Hinge 35mm Standard (Pack of 10)",
        "brand": "Ebco",
        "category": "Hinges", "sub_category": "Concealed Hinge",
        "unit": "pack", "size": "35mm cup diameter | Full overlay",
        "thickness": "For doors 16–19mm thick",
        "finish": "Zinc alloy / Nickel plated",
        "weight_kg": 0.32,
        "colors": "Zinc / Nickel silver",
        "buy_price": 165, "sell_price": 245, "margin_pct": 32.7,
        "gst_rate": 18.0, "hsn_code": "8302",
        "applications": [
            "Kitchen cabinet doors", "Wardrobe shutters",
            "Bookshelf doors", "TV unit shutters", "Bathroom cabinet doors",
        ],
        "certifications": ["80,000 cycle tested", "Load rated 40 kg per pair"],
        "features": [
            "3-way adjustable (vertical, horizontal, depth)",
            "Quick snap-on/off mounting plate",
            "Full overlay, half overlay, inset available",
            "Soft spring action",
        ],
        "installation_tips": "Drill 35mm Forstner bit cup hole 13mm deep. Maintain 3-4mm edge distance from door edge.",
        "lead_time": "2–3 days", "moq": 5, "stock_status": "in_stock",
        "tags": ["ebco", "hinge", "concealed-hinge", "cabinet", "wardrobe", "hardware"],
        "competitors": ["Hafele", "Hettich Sensys", "Blum Clip Top"],
    },
    {
        "product_id": 121, "sku_code": "EBCO-HNG-35SC",
        "name": "Ebco Concealed Hinge 35mm Soft-Close (Pack of 10)",
        "brand": "Ebco",
        "category": "Hinges", "sub_category": "Soft-Close Concealed",
        "unit": "pack", "size": "35mm cup | Full overlay",
        "thickness": "For doors 16–19mm thick",
        "finish": "Nickel plated",
        "weight_kg": 0.38,
        "colors": "Nickel / Silver",
        "buy_price": 320, "sell_price": 480, "margin_pct": 33.3,
        "gst_rate": 18.0, "hsn_code": "8302",
        "applications": [
            "Premium modular kitchen shutters", "Wardrobe soft-close doors",
            "High-end furniture shutters",
        ],
        "certifications": ["100,000 cycle tested", "Integrated damper certified"],
        "features": [
            "Integrated hydraulic soft-close damper",
            "Silent door closing — prevents slamming",
            "Adjustable closing speed & force",
            "3-way position adjustment",
        ],
        "installation_tips": "Same as standard hinge. Adjust damper strength via screw on hinge body.",
        "lead_time": "2–3 days", "moq": 5, "stock_status": "in_stock",
        "tags": ["ebco", "hinge", "soft-close", "concealed", "premium", "modular-kitchen"],
        "competitors": ["Blum Clip Top Blumotion", "Hettich Sensys", "Hafele"],
    },
    {
        "product_id": 122, "sku_code": "EBCO-HDL-ALU480",
        "name": "Ebco Aluminium Profile Handle 480mm",
        "brand": "Ebco",
        "category": "Handles & Knobs", "sub_category": "Aluminium Profile Handle",
        "unit": "Nos", "size": "480mm overall (CC: 448mm) | Bar dia: 10mm",
        "thickness": "Projection: 32mm",
        "finish": "Anodized — Matt Silver, Matt Black, Champagne Gold",
        "weight_kg": 0.18,
        "colors": "Matt Silver, Matt Black, Champagne Gold, Rose Gold, Gunmetal",
        "buy_price": 78, "sell_price": 118, "margin_pct": 33.9,
        "gst_rate": 18.0, "hsn_code": "8302",
        "applications": [
            "Kitchen shutter handles", "Wardrobe door pulls",
            "Vanity unit drawer handles", "TV unit handles", "Office furniture handles",
        ],
        "certifications": ["ISO 9001 manufactured", "Salt spray test 96 hrs"],
        "features": [
            "Aluminium alloy — lightweight yet rigid",
            "Anti-corrosion anodized finish",
            "Premium matte finish — fingerprint resistant",
            "M4 screw installation — easy & secure",
        ],
        "installation_tips": "Centre-to-centre: 448mm. Use M4 × 35mm screws. Mark carefully before drilling.",
        "lead_time": "2–3 days", "moq": 20, "stock_status": "in_stock",
        "tags": ["ebco", "handle", "aluminium", "kitchen", "wardrobe", "furniture"],
        "competitors": ["Hafele", "Hettich", "Ozone", "Dorset"],
    },
    {
        "product_id": 123, "sku_code": "EBCO-LCK-CAM16",
        "name": "Ebco Cam Lock 16mm (Pack of 50)",
        "brand": "Ebco",
        "category": "Furniture Locks", "sub_category": "Cam Lock",
        "unit": "pack", "size": "16mm diameter barrel | Depth: 16mm",
        "thickness": "For panels 15–19mm",
        "finish": "Zinc die-cast, Nickel plated",
        "weight_kg": 0.48,
        "colors": "Nickel silver",
        "buy_price": 185, "sell_price": 280, "margin_pct": 33.9,
        "gst_rate": 18.0, "hsn_code": "8301",
        "applications": [
            "Cabinet flat-pack assembly lock", "Wardrobe panel connection",
            "Modular furniture securing", "Office furniture assembly",
        ],
        "certifications": ["IS 7542 compliant", "Pull-out force 1.2 kN tested"],
        "features": [
            "Quarter-turn locking mechanism",
            "RTA (ready-to-assemble) furniture standard",
            "Concealed — invisible when assembled",
            "Reusable — disassemble and reassemble",
        ],
        "installation_tips": "Drill 16mm hole × 13mm deep for cam. 8mm hole in mating panel for bolt. Quarter turn to lock.",
        "lead_time": "2–3 days", "moq": 5, "stock_status": "in_stock",
        "tags": ["ebco", "cam-lock", "furniture-lock", "connector", "flat-pack"],
        "competitors": ["Hafele", "Hettich", "Fischer"],
    },
    {
        "product_id": 124, "sku_code": "EBCO-KB-CORNER900",
        "name": "Ebco Corner Magic Pull-Out Basket (900mm Cabinet)",
        "brand": "Ebco",
        "category": "Kitchen Systems", "sub_category": "Corner Pull-Out",
        "unit": "set", "size": "For 900×900mm corner cabinet",
        "thickness": "Full extension swing-out",
        "finish": "Chrome plated wire / White epoxy coated",
        "weight_kg": 4.2,
        "colors": "Chrome, White",
        "buy_price": 1850, "sell_price": 2800, "margin_pct": 33.9,
        "gst_rate": 18.0, "hsn_code": "8302",
        "applications": [
            "Modular kitchen L-shaped corner cabinets",
            "Kitchen dead-corner utilisation",
            "Corner storage organisation",
        ],
        "certifications": ["Load rated 15 kg per shelf", "50,000 cycle tested"],
        "features": [
            "Full swing-out — 100% access to corner space",
            "2-tier baskets included",
            "Soft-close damper on return",
            "Adjustable basket height positions",
        ],
        "installation_tips": "For 900×900mm corner. Hinge baskets to door so they pull out together. Use provided mounting template.",
        "lead_time": "3–5 days", "moq": 2, "stock_status": "in_stock",
        "tags": ["ebco", "kitchen", "corner-basket", "pull-out", "modular-kitchen"],
        "competitors": ["Hafele", "Hettich", "Blum Space Corner"],
    },
    {
        "product_id": 125, "sku_code": "EBCO-TANDEM-400",
        "name": "Ebco Tandem Box Drawer System 400mm",
        "brand": "Ebco",
        "category": "Kitchen Systems", "sub_category": "Tandem Drawer Box",
        "unit": "set", "size": "400mm depth | Height: 86mm/136mm/186mm",
        "thickness": "Side panel 12mm steel",
        "finish": "White powder coated / Grey",
        "weight_kg": 2.8,
        "colors": "White, Grey",
        "buy_price": 1180, "sell_price": 1780, "margin_pct": 33.7,
        "gst_rate": 18.0, "hsn_code": "8302",
        "applications": [
            "Premium modular kitchen drawer boxes",
            "Bedroom wardrobe drawer system",
            "Office pedestal drawers",
        ],
        "certifications": ["Load rated 30 kg", "40,000 cycle tested", "Integrated soft-close"],
        "features": [
            "Steel drawer box — no separate bottom needed",
            "Integrated soft-close runners",
            "Push-to-open option available",
            "3 height variants for different storage needs",
        ],
        "installation_tips": "Fix runners to cabinet sides first. Click-in drawer box onto runners. Adjust level with runner clip.",
        "lead_time": "3–5 days", "moq": 5, "stock_status": "in_stock",
        "tags": ["ebco", "tandem", "drawer-box", "kitchen", "premium", "modular"],
        "competitors": ["Blum Tandem", "Hettich InnoTech", "Hafele Matrix"],
    },
    {
        "product_id": 126, "sku_code": "EBCO-GH-GLASS35",
        "name": "Ebco Glass Door Hinge 35mm Hydraulic (Pair)",
        "brand": "Ebco",
        "category": "Glass Hardware", "sub_category": "Glass Door Hinge",
        "unit": "pair", "size": "For 4–6mm glass doors",
        "thickness": "Fits 4mm, 5mm, 6mm glass",
        "finish": "Chrome plated / Matt Black",
        "weight_kg": 0.28,
        "colors": "Chrome, Matt Black, Brushed Nickel",
        "buy_price": 340, "sell_price": 520, "margin_pct": 34.6,
        "gst_rate": 18.0, "hsn_code": "8302",
        "applications": [
            "Glass cabinet doors", "Display unit glass shutters",
            "Crockery unit glass doors", "Retail display cabinets",
            "Bookshelf glass doors",
        ],
        "certifications": ["Load rated 8 kg glass door", "Hydraulic damper cycle tested"],
        "features": [
            "Frameless glass door — no drilling of glass required",
            "Hydraulic soft-close built in",
            "Adjustable tension for closing speed",
            "Spring-loaded auto-close",
        ],
        "installation_tips": "Drill 8mm hole in cabinet side panel. Clamp jaw to glass edge. Adjust torque tension screw.",
        "lead_time": "3–5 days", "moq": 5, "stock_status": "in_stock",
        "tags": ["ebco", "glass-hinge", "glass-door", "display-cabinet", "hardware"],
        "competitors": ["Dorma", "Hafele", "Hettich"],
    },
    {
        "product_id": 127, "sku_code": "EBCO-LED-STRIP1M",
        "name": "Ebco LED Strip Light Warm White 1000mm",
        "brand": "Ebco",
        "category": "Furniture LED Lights", "sub_category": "LED Strip",
        "unit": "Nos", "size": "1000mm length | Strip width: 8mm",
        "thickness": "3mm strip",
        "finish": "Aluminium channel + diffuser cover",
        "weight_kg": 0.12,
        "colors": "Warm White (3000K) / Cool White (6000K) / Neutral (4000K)",
        "buy_price": 165, "sell_price": 250, "margin_pct": 34.0,
        "gst_rate": 18.0, "hsn_code": "9405",
        "applications": [
            "Under-cabinet kitchen lighting", "Wardrobe interior lighting",
            "Display shelf accent lighting", "Furniture showcase lighting",
            "Bookshelf illumination",
        ],
        "certifications": ["IP20 indoor rated", "BIS approved", "12V DC safe"],
        "features": [
            "12V DC LED strip in anodized aluminium channel",
            "Diffuser cover for even, glare-free light",
            "Motion sensor & dimmer variants available",
            "Link multiple strips with connector cables",
        ],
        "installation_tips": "Connect to 12V adapter (sold separately). Peel & stick with double-side tape. Join strips with link cables.",
        "lead_time": "2–3 days", "moq": 10, "stock_status": "in_stock",
        "tags": ["ebco", "led", "lighting", "furniture-light", "kitchen", "wardrobe"],
        "competitors": ["Hafele Loox", "Hettich LED", "Wipro Interior"],
    },
    {
        "product_id": 128, "sku_code": "EBCO-WRD-FLAP450",
        "name": "Ebco Wardrobe Flap Stay (Lift-Up) 450mm",
        "brand": "Ebco",
        "category": "Bed & Wardrobe Fittings", "sub_category": "Flap Stay / Lift",
        "unit": "pair", "size": "Arm length 450mm | For doors 400–800mm wide",
        "thickness": "Zinc alloy arm",
        "finish": "Nickel plated / Grey",
        "weight_kg": 0.42,
        "colors": "Nickel, Grey",
        "buy_price": 285, "sell_price": 430, "margin_pct": 33.7,
        "gst_rate": 18.0, "hsn_code": "8302",
        "applications": [
            "Overhead wardrobe flap doors", "Bed box lift-up lids",
            "Upper kitchen cabinet lift doors", "TV lift mechanisms",
        ],
        "certifications": ["20,000 cycle tested", "Load rated 4–6 kg door"],
        "features": [
            "Free-stop mechanism — stays open at any position",
            "Adjustable spring tension via hex key",
            "Compatible with push-to-open systems",
            "Prevents sudden door drop",
        ],
        "installation_tips": "Mount one arm to door inner face, other to cabinet inner side. Adjust spring tension for door weight.",
        "lead_time": "3–5 days", "moq": 5, "stock_status": "in_stock",
        "tags": ["ebco", "flap-stay", "lift-up", "wardrobe", "furniture"],
        "competitors": ["Hafele Free Flap", "Hettich", "Blum Aventos"],
    },
    {
        "product_id": 129, "sku_code": "EBCO-ALU-TRIM2M",
        "name": "Ebco Aluminium Furniture Trim Profile 2000mm",
        "brand": "Ebco",
        "category": "Aluminium Profiles & Handles", "sub_category": "Furniture Trim",
        "unit": "Nos", "size": "2000mm length | Section: 19×8mm",
        "thickness": "8mm height",
        "finish": "Anodized — Silver, Champagne Gold, Matt Black",
        "weight_kg": 0.22,
        "colors": "Silver Anodized, Champagne Gold, Matt Black, Rose Gold",
        "buy_price": 62, "sell_price": 95, "margin_pct": 34.7,
        "gst_rate": 18.0, "hsn_code": "7604",
        "applications": [
            "Furniture edge T-trim finishing",
            "Glass door aluminium frame profiles",
            "Wardrobe edge profile alternative to PVC banding",
            "Decorative furniture corner trim",
        ],
        "certifications": ["Anodized 15 micron — IS 1868"],
        "features": [
            "Aluminium 6063 alloy — lightweight & rigid",
            "Anodized for long-lasting corrosion resistance",
            "Easy cut with hacksaw or mitre saw",
            "Multiple section profiles in the range",
        ],
        "installation_tips": "Apply with contact adhesive. Mitre corners at 45°. Press firmly and clamp for 30 minutes.",
        "lead_time": "3–5 days", "moq": 20, "stock_status": "in_stock",
        "tags": ["ebco", "aluminium-profile", "trim", "furniture-edge", "profile"],
        "competitors": ["Hafele", "Dorset", "Ozone"],
    },
    {
        "product_id": 130, "sku_code": "EBCO-MINIFIX-50",
        "name": "Ebco Minifix Connecting Cam Set (Pack of 50 sets)",
        "brand": "Ebco",
        "category": "Joinery Fittings & Screws", "sub_category": "Cam Connector",
        "unit": "pack", "size": "15mm cam dia | M6 × 45mm bolt",
        "thickness": "For panels 15–19mm",
        "finish": "Zinc die-cast, bright finish",
        "weight_kg": 0.65,
        "colors": "Zinc silver",
        "buy_price": 195, "sell_price": 295, "margin_pct": 33.9,
        "gst_rate": 18.0, "hsn_code": "7318",
        "applications": [
            "Flat-pack furniture panel joining",
            "Cabinet carcase assembly",
            "Modular wardrobe & shelf joining",
            "RTA furniture manufacture",
        ],
        "certifications": ["Pull-out force 1.5 kN tested"],
        "features": [
            "3-part system: cam + pre-inserted bolt + dowel",
            "Quarter-turn cam locking — invisible connection",
            "Fully concealed — clean appearance",
            "Reusable — disassemble and reassemble multiple times",
        ],
        "installation_tips": "15mm hole × 13mm deep for cam in panel face. 8mm × 30mm in mating panel edge for bolt. Quarter turn with screwdriver to lock.",
        "lead_time": "2–3 days", "moq": 5, "stock_status": "in_stock",
        "tags": ["ebco", "minifix", "cam-connector", "joinery", "flat-pack", "fittings"],
        "competitors": ["Hafele", "Fischer", "Hettich"],
    },
    {
        "product_id": 131, "sku_code": "EBCO-HDL-SS128",
        "name": "Ebco Stainless Steel Bar Handle 128mm CC",
        "brand": "Ebco",
        "category": "Handles & Knobs", "sub_category": "SS Bar Handle",
        "unit": "Nos", "size": "128mm CC | Overall: 168mm | Bar dia: 12mm",
        "thickness": "Projection: 32mm",
        "finish": "Brushed SS / Mirror Polish SS / PVD Gold / PVD Black",
        "weight_kg": 0.14,
        "colors": "Brushed SS, Mirror SS, PVD Gold, PVD Black, PVD Rose Gold",
        "buy_price": 92, "sell_price": 140, "margin_pct": 34.3,
        "gst_rate": 18.0, "hsn_code": "8302",
        "applications": [
            "Kitchen drawer pulls", "Modular furniture knobs",
            "Bathroom cabinet handles", "Wardrobe handles",
        ],
        "certifications": ["304 grade stainless steel", "Salt spray 200 hrs"],
        "features": [
            "Solid 304 SS — rust-free for life",
            "PVD coating option — scratch resistant",
            "Standard 128mm CC fits most cabinets",
            "Heavy duty — suits commercial & residential",
        ],
        "installation_tips": "Standard M4 screws included. CC: 128mm. Can be used on 16–25mm door/drawer thickness.",
        "lead_time": "2–3 days", "moq": 20, "stock_status": "in_stock",
        "tags": ["ebco", "ss-handle", "stainless-steel", "kitchen", "pvd", "hardware"],
        "competitors": ["Hafele", "Hettich", "Ozone", "Starbucks Hardware"],
    },
    {
        "product_id": 132, "sku_code": "EBCO-OFFC-GASLIFT",
        "name": "Ebco Office Chair Gas Lift Cylinder 100mm",
        "brand": "Ebco",
        "category": "Office Furniture Fittings", "sub_category": "Gas Lift Cylinder",
        "unit": "Nos", "size": "Stroke: 100mm | Shaft dia: 28mm | Base dia: 50mm",
        "thickness": "Total height: 290–390mm",
        "finish": "Chrome plated / Black",
        "weight_kg": 1.2,
        "colors": "Chrome, Black",
        "buy_price": 280, "sell_price": 420, "margin_pct": 33.3,
        "gst_rate": 18.0, "hsn_code": "8483",
        "applications": [
            "Office swivel chair height adjustment",
            "Ergonomic workstation chairs",
            "Computer chairs", "Executive chairs",
        ],
        "certifications": ["BIFMA X5.1 compliant", "Load rated 150 kg", "ANSI/BIFMA certified"],
        "features": [
            "Pneumatic height adjustment 100mm stroke",
            "Class 4 gas cylinder — heavy duty",
            "Universal fit — standard chair base",
            "Smooth single-hand height adjustment",
        ],
        "installation_tips": "Press into chair base and seat plate — tapered fit, no tools needed. Replace by pulling straight up firmly.",
        "lead_time": "3–5 days", "moq": 5, "stock_status": "in_stock",
        "tags": ["ebco", "office-chair", "gas-lift", "office-furniture", "ergonomic"],
        "competitors": ["Hafele", "Wipro Furniture", "KK Nagar"],
    },
]

# ── Routes ────────────────────────────────────────────────────────────────────

@router.get("/catalog")
async def get_catalog(
    category: Optional[str] = None,
    search:   Optional[str] = None,
    in_stock: Optional[bool] = None,
):
    all_products = _get_all_products()
    products = all_products
    if category:
        products = [p for p in products if p["category"].lower() == category.lower()]
    if in_stock is not None:
        products = [p for p in products if (p["stock_status"] == "in_stock") == in_stock]
    if search:
        q = search.lower()
        products = [
            p for p in products
            if q in p["name"].lower() or q in p["category"].lower()
            or q in p.get("brand", "").lower()
            or q in " ".join(p.get("tags", [])).lower()
            or any(q in app.lower() for app in p.get("applications", []))
        ]
    cats = sorted(set(p["category"] for p in all_products))
    return {
        "products":    products,
        "total":       len(products),
        "categories":  cats,
        "data_source": "catalog" if not _RUNTIME_PRODUCTS else "catalog+added",
    }


@router.get("/catalog/categories")
async def get_categories():
    cats = {}
    for p in _get_all_products():
        cats.setdefault(p["category"], []).append(p["sub_category"])
    return {"categories": {k: list(set(v)) for k, v in cats.items()}}


@router.get("/catalog/{product_id}")
async def get_product(product_id: int):
    p = next((p for p in _get_all_products() if p["product_id"] == product_id), None)
    if not p:
        raise HTTPException(404, f"Product {product_id} not found in catalog")
    return p


# ── AI catalog scan helpers ────────────────────────────────────────────────────

_CATALOG_SCAN_SYSTEM = """You are a product catalog extraction specialist for Indian building materials, furniture hardware, and interior products.

Extract ALL products from the input — catalog images, MRP price lists, spec sheets, product photos, or typed descriptions.

PRICING (Indian market context):
- MRP list "incl. GST 18%": sell_price = round(MRP / 1.18), buy_price = round(sell_price * 0.72) for standard 28% margin
- If only net price given: sell_price = that price, buy_price = round(sell_price * 0.75)
- Always set margin_pct = round((sell_price - buy_price) / sell_price * 100, 1)

HSN CODES (use these exact codes):
- Hinges, drawer slides, handles, stays, shelf supports, furniture fittings: 8302
- Locks, padlocks: 8301
- LED lights, luminaires: 9405
- Aluminium extrusions/profiles: 7604
- Screws, bolts, nuts, cam connectors, dowels: 7318
- HPL / laminates: 4814
- PVC sheets/profiles: 3921
- ACP aluminium composite: 7606
- Kitchen wire baskets, steel articles: 7323
- Gas lifts, pneumatic cylinders: 8483
- Wardrobe fittings, bed mechanisms: 8302

EXACT CATEGORY NAMES (must match exactly):
"Drawer Slides" | "Hinges" | "Handles & Knobs" | "Furniture Locks" | "Joinery Fittings & Screws"
"Bed & Wardrobe Fittings" | "Office Furniture Fittings" | "Kitchen Systems" | "Glass Hardware"
"Furniture LED Lights" | "Aluminium Profiles & Handles"
"High Pressure Laminate" | "Compact Laminate" | "Acrylic Laminate" | "PVC Laminate"
"Aluminium Louvers" | "PVC Louvers" | "Operable Louvre System" | "Exterior Cladding" | "Aluminium Composite Panel"

Return ONLY this JSON (no markdown, no extra text):
{
  "brand": "primary brand name from the document",
  "document_type": "catalog/price_list/spec_sheet/photo/description",
  "total_found": 0,
  "products": [
    {
      "name": "full descriptive product name including key spec (e.g. 'Ebco Soft-Close Drawer Slide 400mm')",
      "sku_code": "item/model code from document, else derive from name",
      "category": "use EXACT category from list above",
      "sub_category": "sub-type within category",
      "brand": "brand name",
      "unit": "Nos/pair/set/pack/sheet/RM/SQM/kg",
      "size": "dimensions as stated in document",
      "thickness": "thickness if applicable, else empty string",
      "finish": "available finish options",
      "weight_kg": 0.0,
      "colors": "available colors/finishes as comma-separated string",
      "sell_price": 0,
      "buy_price": 0,
      "margin_pct": 0.0,
      "gst_rate": 18.0,
      "hsn_code": "4-digit HSN from guide above",
      "applications": ["specific application 1", "specific application 2", "specific application 3"],
      "certifications": [],
      "features": ["key feature with value/spec", "feature 2"],
      "installation_tips": "one practical tip for installation",
      "lead_time": "2-3 days",
      "moq": 1,
      "stock_status": "in_stock",
      "tags": ["brand-lowercase", "category-keyword", "application-keyword"],
      "competitors": []
    }
  ]
}"""


def _extract_for_catalog(file_bytes: bytes, content_type: str, filename: str):
    """Returns (text, is_image, image_b64, image_ct)."""
    IMAGE_EXTS = (".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".tiff", ".heic")
    is_image = content_type.startswith("image/") or any(filename.endswith(e) for e in IMAGE_EXTS)
    if is_image:
        b64 = base64.b64encode(file_bytes).decode()
        ct  = content_type if content_type.startswith("image/") else "image/jpeg"
        return "", True, b64, ct

    if content_type == "application/pdf" or filename.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(_io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    pages_text.append(t)
            text = "\n".join(pages_text)
            if text.strip():
                # 32k chars ~ 8k tokens, enough for large MRP lists
                return text[:32000], False, "", ""
            # No text layer — try rendering first page as image via pdf2image if available
            try:
                from pdf2image import convert_from_bytes
                images = convert_from_bytes(file_bytes, first_page=1, last_page=1, dpi=200)
                if images:
                    img_buf = _io.BytesIO()
                    images[0].save(img_buf, format="JPEG", quality=85)
                    b64 = base64.b64encode(img_buf.getvalue()).decode()
                    return "", True, b64, "image/jpeg"
            except Exception:
                pass
            return "__scanned_pdf__", False, "", ""
        except Exception as e:
            logger.warning("pypdf failed: %s", e)
            return "", False, "", ""

    if filename.endswith((".docx", ".doc")) or "wordprocessingml" in content_type:
        try:
            from docx import Document
            doc = Document(_io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    parts.append("\t".join(c.text.strip() for c in row.cells if c.text.strip()))
            return "\n".join(parts)[:16000], False, "", ""
        except Exception as e:
            logger.warning("python-docx failed: %s", e)

    if filename.endswith((".xlsx", ".xls")) or "spreadsheetml" in content_type:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(_io.BytesIO(file_bytes), read_only=True, data_only=True)
            lines = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    row_str = "\t".join(str(c) for c in row if c is not None)
                    if row_str.strip():
                        lines.append(row_str)
            return "\n".join(lines)[:16000], False, "", ""
        except Exception as e:
            logger.warning("openpyxl failed: %s", e)

    if filename.endswith(".csv") or content_type in ("text/csv", "application/csv"):
        try:
            import csv
            decoded = file_bytes.decode("utf-8", errors="ignore")
            import io as io_
            reader = csv.reader(io_.StringIO(decoded))
            rows = [", ".join(r) for r in reader if any(c.strip() for c in r)]
            return "\n".join(rows)[:16000], False, "", ""
        except Exception as e:
            logger.warning("csv parse failed: %s", e)

    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            text = file_bytes.decode(enc, errors="ignore")
            if text.strip():
                return text[:16000], False, "", ""
        except Exception:
            continue
    return "", False, "", ""


# ── Catalog scan endpoint ──────────────────────────────────────────────────────

@router.post("/catalog/scan-image")
async def scan_catalog_image(
    file:       Optional[UploadFile] = File(None),
    text_input: str = Form(""),
):
    """AI-powered product extraction from any file or typed text. Returns extracted products for review before adding."""
    api_key = os.getenv("OPENAI_API_KEY", "")
    if not api_key:
        return {
            "brand": "Demo Brand",
            "document_type": "demo",
            "total_found": 2,
            "products": [
                {
                    "name": "Demo Concealed Hinge 35mm Soft-Close (Pack 10)",
                    "sku_code": "DEMO-HNG-35SC", "category": "Hinges",
                    "sub_category": "Soft-Close Concealed", "brand": "Demo Brand",
                    "unit": "pack", "size": "35mm cup", "thickness": "For 16-18mm doors",
                    "finish": "Nickel plated", "weight_kg": 0.38,
                    "colors": "Nickel Silver",
                    "sell_price": 480, "buy_price": 320, "margin_pct": 33.3,
                    "gst_rate": 18.0, "hsn_code": "8302",
                    "applications": ["Kitchen cabinet doors", "Wardrobe shutters", "Furniture doors"],
                    "certifications": ["100,000 cycle tested"],
                    "features": ["Hydraulic soft-close", "3-way adjustable", "Silent closing"],
                    "installation_tips": "Drill 35mm Forstner bit. 3mm edge distance.",
                    "lead_time": "2-3 days", "moq": 5, "stock_status": "in_stock",
                    "tags": ["hinge", "soft-close", "hardware"], "competitors": [],
                },
                {
                    "name": "Demo Telescopic Drawer Slide 400mm",
                    "sku_code": "DEMO-DS-400", "category": "Drawer Slides",
                    "sub_category": "Telescopic", "brand": "Demo Brand",
                    "unit": "pair", "size": "400mm", "thickness": "17mm channel",
                    "finish": "Zinc plated", "weight_kg": 0.42,
                    "colors": "Zinc plated",
                    "sell_price": 195, "buy_price": 128, "margin_pct": 34.4,
                    "gst_rate": 18.0, "hsn_code": "8302",
                    "applications": ["Kitchen drawers", "Wardrobe pull-outs", "Furniture drawers"],
                    "certifications": ["Load 30 kg"],
                    "features": ["3/4 extension", "Ball bearing", "30 kg load"],
                    "installation_tips": "Fix runners level on both sides.",
                    "lead_time": "2-3 days", "moq": 10, "stock_status": "in_stock",
                    "tags": ["drawer-slide", "hardware"], "competitors": [],
                },
            ],
            "demo_note": "No OPENAI_API_KEY set — showing demo extraction. Add your key to see real AI extraction.",
        }

    has_file = file is not None and file.filename
    has_text = bool(text_input.strip())
    if not has_file and not has_text:
        return {"brand": "", "document_type": "empty", "total_found": 0, "products": [], "error": "No input provided."}

    file_bytes, content_type, filename = b"", "", ""
    if has_file:
        file_bytes   = await file.read()
        content_type = file.content_type or ""
        filename     = (file.filename or "").lower()

    text_content, is_image, image_b64, image_ct = _extract_for_catalog(
        file_bytes, content_type, filename
    ) if has_file else ("", False, "", "")

    if text_content == "__scanned_pdf__":
        return {
            "brand": "", "document_type": "scanned_pdf", "total_found": 0, "products": [],
            "error": "This PDF has no text layer (scanned/image-only). Take a screenshot of the price list page and upload as JPG or PNG — AI Vision will read it directly.",
        }

    combined = "\n\n".join(filter(None, [text_input.strip(), text_content]))

    try:
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=api_key, timeout=60.0)

        if is_image and not has_text:
            user_content = [
                {"type": "image_url", "image_url": {"url": f"data:{image_ct};base64,{image_b64}", "detail": "high"}},
                {"type": "text", "text": "Extract ALL products from this catalog image. Return ONLY valid JSON as specified."},
            ]
        elif is_image and has_text:
            user_content = [
                {"type": "image_url", "image_url": {"url": f"data:{image_ct};base64,{image_b64}", "detail": "high"}},
                {"type": "text", "text": f"Extract ALL products. Additional context:\n{text_input.strip()}\n\nReturn ONLY valid JSON."},
            ]
        else:
            user_content = (
                f"Extract ALL products from this input. It may be a price list, catalog, spec sheet, or description.\n\n"
                f"---\n{combined}\n---\n\nReturn ONLY valid JSON as specified."
            )

        resp = await client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": _CATALOG_SCAN_SYSTEM},
                {"role": "user",   "content": user_content},
            ],
            max_tokens=8000,
            response_format={"type": "json_object"},
        )
        raw = resp.choices[0].message.content or ""
        # Strip markdown fences GPT sometimes adds despite json_object mode
        clean = raw.strip()
        if clean.startswith("```"):
            clean = clean.split("```", 2)[-1] if clean.count("```") >= 2 else clean
            if clean.startswith("json"):
                clean = clean[4:]
            end = clean.rfind("```")
            if end != -1:
                clean = clean[:end]
            clean = clean.strip()
        # Handle truncated JSON (max_tokens hit mid-object)
        if not clean.endswith("}"):
            last_close = max(clean.rfind("}"), 0)
            if last_close:
                clean = clean[:last_close + 1]
                # Repair the outer products array if truncated inside it
                opens = clean.count("[") - clean.count("]")
                closes = clean.count("{") - clean.count("}")
                clean += "]" * max(opens, 0) + "}" * max(closes, 0)
        try:
            result = json.loads(clean)
        except json.JSONDecodeError:
            # Last resort: try to salvage partial products array
            import re as _re
            products_match = _re.search(r'"products"\s*:\s*(\[.*)', clean, _re.DOTALL)
            if products_match:
                arr_str = products_match.group(1)
                # Close any unclosed brackets
                arr_str += "]" * (arr_str.count("[") - arr_str.count("]"))
                try:
                    products = json.loads(arr_str)
                    result = {"brand": "", "document_type": "partial", "products": products}
                except Exception:
                    raise
            else:
                raise
        result.setdefault("products", [])
        result["total_found"] = len(result["products"])
        return result

    except Exception as exc:
        logger.exception("catalog scan error: %s", exc)
        err_type = type(exc).__name__
        if "JSONDecodeError" in err_type:
            msg = "AI returned an incomplete response (catalog may be very large). Try uploading fewer pages at a time, or paste the product list as text."
        elif "AuthenticationError" in err_type or "Unauthorized" in str(exc):
            msg = "Invalid OPENAI_API_KEY. Check your .env file and restart the backend."
        elif "RateLimitError" in err_type:
            msg = "OpenAI rate limit reached. Wait 60 seconds and try again."
        else:
            msg = f"AI extraction failed ({err_type}). Check your OPENAI_API_KEY and try again."
        return {
            "brand": "", "document_type": "error", "total_found": 0, "products": [],
            "error": msg,
        }


# ── Add product endpoints ──────────────────────────────────────────────────────

@router.post("/catalog/add-product", status_code=201)
async def add_catalog_product(product: dict):
    """Add a single product to the runtime catalog. Assigns a new product_id."""
    try:
        from app.db.connection import get_pool
        pool = await get_pool()
        if pool:
            logger.info("DB mode: catalog product would be persisted — table not yet created")
    except Exception:
        pass

    new_id = _next_id()
    product["product_id"] = new_id
    product.setdefault("stock_status", "in_stock")
    product.setdefault("gst_rate", 18.0)
    product.setdefault("certifications", [])
    product.setdefault("features", [])
    product.setdefault("applications", [])
    product.setdefault("tags", [])
    product.setdefault("competitors", [])
    if not product.get("margin_pct") and product.get("sell_price") and product.get("buy_price"):
        sp = float(product["sell_price"])
        bp = float(product["buy_price"])
        product["margin_pct"] = round((sp - bp) / sp * 100, 1) if sp > 0 else 0.0
    _RUNTIME_PRODUCTS.append(product)
    logger.info("Catalog: added product id=%s name=%s", new_id, product.get("name", ""))
    return {"product": product, "message": f"Product added to catalog (ID: {new_id})", "data_source": "runtime"}


@router.post("/catalog/bulk-add", status_code=201)
async def bulk_add_catalog_products(payload: dict):
    """Add multiple products to the runtime catalog at once."""
    products = payload.get("products", [])
    added, skipped = [], []
    for p in products:
        if not p.get("name"):
            skipped.append(p)
            continue
        new_id = _next_id()
        p["product_id"] = new_id
        p.setdefault("stock_status", "in_stock")
        p.setdefault("gst_rate", 18.0)
        p.setdefault("certifications", [])
        p.setdefault("features", [])
        p.setdefault("applications", [])
        p.setdefault("tags", [])
        p.setdefault("competitors", [])
        if not p.get("margin_pct") and p.get("sell_price") and p.get("buy_price"):
            sp = float(p["sell_price"])
            bp = float(p["buy_price"])
            p["margin_pct"] = round((sp - bp) / sp * 100, 1) if sp > 0 else 0.0
        _RUNTIME_PRODUCTS.append(p)
        added.append(p)
        logger.info("Catalog bulk-add: id=%s name=%s", new_id, p.get("name", ""))

    return {
        "added":    len(added),
        "skipped":  len(skipped),
        "products": added,
        "message":  f"{len(added)} product(s) added to catalog successfully.",
        "data_source": "runtime",
    }


@router.post("/catalog/parse-import")
async def parse_import_file(file: UploadFile = File(...)):
    """Parse a CSV or Excel file and return raw tabular data (headers + rows) for column mapping."""
    import csv
    import io as _sio

    file_bytes   = await file.read()
    filename     = (file.filename or "").lower()
    content_type = file.content_type or ""

    headers: list = []
    rows: list    = []

    if filename.endswith(".csv") or content_type in ("text/csv", "application/csv"):
        try:
            text     = file_bytes.decode("utf-8", errors="replace")
            reader   = csv.reader(_sio.StringIO(text))
            all_rows = list(reader)
            if all_rows:
                headers = [str(h).strip() for h in all_rows[0]]
                rows    = [
                    [str(c).strip() for c in r]
                    for r in all_rows[1:]
                    if any(str(c).strip() for c in r)
                ]
        except Exception as exc:
            raise HTTPException(400, f"CSV parse error: {exc}") from exc

    elif filename.endswith((".xlsx", ".xls")) or "spreadsheetml" in content_type:
        try:
            from openpyxl import load_workbook
            wb  = load_workbook(_io.BytesIO(file_bytes), read_only=True, data_only=True)
            ws  = wb.active
            all_rows = []
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(c).strip() if c is not None else "" for c in row]
                if any(v for v in row_vals):
                    all_rows.append(row_vals)
            wb.close()
            if all_rows:
                headers = all_rows[0]
                rows    = [r for r in all_rows[1:] if any(v for v in r)]
        except Exception as exc:
            raise HTTPException(400, f"Excel parse error: {exc}") from exc

    else:
        raise HTTPException(
            400,
            "Unsupported file type. Upload a .csv or .xlsx / .xls file.",
        )

    return {"headers": headers, "rows": rows[:500], "count": len(rows)}
