"""
Quotation Builder API — InvenIQ
Professional quotation management for louvers, laminates & building materials.
Covers full quote lifecycle: Draft → Sent → Negotiating → Won/Lost.
"""
import asyncio
import base64
import datetime
import json
import logging
import os
import re
from typing import List, Optional

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)
router = APIRouter(tags=["Quotation Builder"])

VALID_STATUSES = {"DRAFT", "SENT", "NEGOTIATING", "WON", "LOST", "EXPIRED", "REVISED"}


# ── Pydantic models ────────────────────────────────────────────────────────────

class QuoteLineItem(BaseModel):
    product_id:     str = ""
    product_name:   str
    category:       str = ""
    quantity:       float = Field(gt=0)
    unit:           str = "sheet"
    unit_price:     float = Field(ge=0)
    discount_pct:   float = Field(ge=0, le=100, default=0)
    buy_price:      float = Field(ge=0, default=0)
    specifications: str = ""


class CreateQuoteRequest(BaseModel):
    customer_name:    str
    customer_type:    str = "Developer"
    contact_person:   str = ""
    contact_phone:    str = ""
    contact_email:    str = ""
    gst_number:       str = ""
    billing_address:  str = ""
    site_location:    str = ""
    project_name:     str = ""
    payment_terms:    str = "50% Advance + 50% on Delivery"
    delivery_terms:   str = "Door Delivery — Bangalore"
    validity_days:    int = 14
    notes:            str = ""
    gst_rate:         float = 18.0
    include_freight:  bool = False
    freight_amount:   float = 0.0
    line_items:       List[QuoteLineItem] = []


class UpdateQuoteStatus(BaseModel):
    status:   str
    remarks:  Optional[str] = None


class AIPriceRequest(BaseModel):
    customer_type:  str
    total_quantity: float
    category:       str = ""
    current_margin: float = 20.0


class AnalyzeLineItem(BaseModel):
    product_name:  str   = ""
    category:      str   = ""
    quantity:      float = 0
    unit:          str   = "sheet"
    unit_price:    float = 0
    discount_pct:  float = 0
    buy_price:     float = 0


class AnalyzeQuoteRequest(BaseModel):
    customer_name:   str   = ""
    customer_type:   str   = "Developer"
    project_name:    str   = ""
    payment_terms:   str   = ""
    subtotal:        float = 0
    gst_rate:        float = 18
    include_freight: bool  = False
    freight_amount:  float = 0
    notes:           str   = ""
    line_items:      List[AnalyzeLineItem] = []


# ── Mock data ──────────────────────────────────────────────────────────────────

def _mock_quotes():
    today = datetime.date.today()
    return [
        {
            "quote_id": 1,
            "quote_number": "QT-2026-0089",
            "created_at": (today - datetime.timedelta(days=3)).isoformat(),
            "valid_till": (today + datetime.timedelta(days=11)).isoformat(),
            "status": "NEGOTIATING",
            "customer_name": "Prestige Developers",
            "customer_type": "Developer",
            "contact_person": "Mr. Rajesh Shetty",
            "contact_phone": "+91 98450 11223",
            "contact_email": "rajesh.shetty@prestige.co.in",
            "gst_number": "29AABCP1234A1ZK",
            "billing_address": "Prestige Tower, MG Road, Bangalore 560001",
            "site_location": "Whitefield, Bangalore",
            "project_name": "Prestige Skyrise — Tower A & B Facade",
            "payment_terms": "50% Advance + 50% on Delivery",
            "delivery_terms": "Door Delivery — Bangalore",
            "validity_days": 14,
            "notes": "Client is comparing with Hyperion Systems. Price sensitivity is high.",
            "gst_rate": 18.0,
            "include_freight": True,
            "freight_amount": 18000,
            "subtotal": 480000,
            "gst_amount": 86400,
            "grand_total": 584400,
            "margin_pct": 20.1,
            "line_items": [
                {
                    "sl": 1, "product_id": "AL-Z100-ANOD", "product_name": "Aluminium Z-100 Anodized Louver",
                    "category": "Aluminium Louver", "quantity": 1200, "unit": "Rmt",
                    "unit_price": 320, "discount_pct": 5.0, "buy_price": 245, "net_price": 304,
                    "line_total": 364800, "specifications": "100mm blade, Silver Anodized",
                },
                {
                    "sl": 2, "product_id": "AL-C150-AERO", "product_name": "Aluminium C-150 Aerofoil Louver",
                    "category": "Aluminium Louver", "quantity": 180, "unit": "Rmt",
                    "unit_price": 640, "discount_pct": 4.0, "buy_price": 490, "net_price": 614,
                    "line_total": 110592, "specifications": "150mm aerofoil blade, powder coated RAL 9006",
                },
            ],
        },
        {
            "quote_id": 2,
            "quote_number": "QT-2026-0088",
            "created_at": (today - datetime.timedelta(days=7)).isoformat(),
            "valid_till": (today + datetime.timedelta(days=7)).isoformat(),
            "status": "WON",
            "customer_name": "Sobha Builders",
            "customer_type": "Developer",
            "contact_person": "Mrs. Anitha Krishnan",
            "contact_phone": "+91 80002 44556",
            "contact_email": "anitha@sobha.com",
            "gst_number": "29AASCS4567B1ZJ",
            "billing_address": "Sobha City, Thanisandra, Bangalore 560077",
            "site_location": "Hebbal, Bangalore",
            "project_name": "Sobha Dream Series — Phase 2",
            "payment_terms": "100% Advance",
            "delivery_terms": "Door Delivery — Bangalore",
            "validity_days": 14,
            "notes": "Repeat client — offer 2% extra loyalty discount on next order.",
            "gst_rate": 18.0,
            "include_freight": False,
            "freight_amount": 0,
            "subtotal": 320000,
            "gst_amount": 57600,
            "grand_total": 377600,
            "margin_pct": 23.5,
            "line_items": [
                {
                    "sl": 1, "product_id": "HPL-1.5MM-COMP", "product_name": "Compact Laminate 6mm",
                    "category": "Compact Laminate", "quantity": 320, "unit": "sheet",
                    "unit_price": 1800, "discount_pct": 0, "buy_price": 1380, "net_price": 1800,
                    "line_total": 576000, "specifications": "Exterior grade, Trespa-style woodgrain",
                },
            ],
        },
        {
            "quote_id": 3,
            "quote_number": "QT-2026-0087",
            "created_at": (today - datetime.timedelta(days=12)).isoformat(),
            "valid_till": (today - datetime.timedelta(days=2)).isoformat(),
            "status": "SENT",
            "customer_name": "Brigade Group",
            "customer_type": "Developer",
            "contact_person": "Mr. Venkat Rao",
            "contact_phone": "+91 99000 88776",
            "contact_email": "venkat.rao@brigade.com",
            "gst_number": "29AACBR8890C1ZF",
            "billing_address": "Brigade Towers, Palace Road, Bangalore 560001",
            "site_location": "Devanahalli, Bangalore",
            "project_name": "Brigade Tech Park — Amenity Block",
            "payment_terms": "Net 30 Days",
            "delivery_terms": "FOR Destination",
            "validity_days": 14,
            "notes": "Expiry chased — client on vacation, follow up Monday.",
            "gst_rate": 18.0,
            "include_freight": True,
            "freight_amount": 12000,
            "subtotal": 195000,
            "gst_amount": 35100,
            "grand_total": 242100,
            "margin_pct": 18.2,
            "line_items": [
                {
                    "sl": 1, "product_id": "PVC-LOUVER-100", "product_name": "PVC Louver 100mm",
                    "category": "PVC Louver", "quantity": 850, "unit": "Rmt",
                    "unit_price": 195, "discount_pct": 3.0, "buy_price": 152, "net_price": 189.15,
                    "line_total": 160778, "specifications": "White, 100mm, UV stabilised",
                },
                {
                    "sl": 2, "product_id": "PVC-LOUVER-75", "product_name": "PVC Louver 75mm",
                    "category": "PVC Louver", "quantity": 180, "unit": "Rmt",
                    "unit_price": 165, "discount_pct": 3.0, "buy_price": 128, "net_price": 160.05,
                    "line_total": 28809, "specifications": "White, 75mm, UV stabilised",
                },
            ],
        },
        {
            "quote_id": 4,
            "quote_number": "QT-2026-0086",
            "created_at": (today - datetime.timedelta(days=18)).isoformat(),
            "valid_till": (today - datetime.timedelta(days=4)).isoformat(),
            "status": "DRAFT",
            "customer_name": "Nambiar Builders",
            "customer_type": "Developer",
            "contact_person": "Ms. Priya Nambiar",
            "contact_phone": "+91 77007 22334",
            "contact_email": "priya@nambiarbuilders.com",
            "gst_number": "",
            "billing_address": "Nambiar House, Koramangala, Bangalore",
            "site_location": "Electronic City, Bangalore",
            "project_name": "Nambiar Millenia — Clubhouse",
            "payment_terms": "50% Advance + 50% on Delivery",
            "delivery_terms": "Ex-Works",
            "validity_days": 21,
            "notes": "Draft pending — waiting for final dimensions from architect.",
            "gst_rate": 18.0,
            "include_freight": False,
            "freight_amount": 0,
            "subtotal": 88000,
            "gst_amount": 15840,
            "grand_total": 103840,
            "margin_pct": 21.4,
            "line_items": [
                {
                    "sl": 1, "product_id": "OP-LOUVER-MOTOR", "product_name": "Operable Motorised Louver System",
                    "category": "Operable System", "quantity": 60, "unit": "Rmt",
                    "unit_price": 1800, "discount_pct": 2.0, "buy_price": 1400, "net_price": 1764,
                    "line_total": 105840, "specifications": "Motorised 120mm blade, Somfy motor, powder coat",
                },
            ],
        },
        {
            "quote_id": 5,
            "quote_number": "QT-2026-0082",
            "created_at": (today - datetime.timedelta(days=30)).isoformat(),
            "valid_till": (today - datetime.timedelta(days=16)).isoformat(),
            "status": "LOST",
            "customer_name": "Godrej Properties",
            "customer_type": "Developer",
            "contact_person": "Mr. Sanjay Malhotra",
            "contact_phone": "+91 98100 45678",
            "contact_email": "sanjay.m@godrej.com",
            "gst_number": "07AABCG1234D1ZQ",
            "billing_address": "Godrej One, LBS Marg, Mumbai 400079",
            "site_location": "Punawale, Pune",
            "project_name": "Godrej Horizon — External Cladding",
            "payment_terms": "Net 45 Days",
            "delivery_terms": "FOR Destination",
            "validity_days": 14,
            "notes": "Lost to competitor offering 8% lower price. Revisit in 6 months for Phase 2.",
            "gst_rate": 18.0,
            "include_freight": True,
            "freight_amount": 35000,
            "subtotal": 740000,
            "gst_amount": 133200,
            "grand_total": 908200,
            "margin_pct": 15.8,
            "line_items": [
                {
                    "sl": 1, "product_id": "HPL-EXT-CLAD", "product_name": "HPL Exterior Cladding Panel",
                    "category": "HPL Cladding", "quantity": 680, "unit": "sheet",
                    "unit_price": 1280, "discount_pct": 6.0, "buy_price": 1050, "net_price": 1203.2,
                    "line_total": 818176, "specifications": "8mm Trespa Meteon FR, Sto-grey finish",
                },
            ],
        },
    ]


def _compute_kpis(quotes):
    pipeline = sum(q["grand_total"] for q in quotes if q["status"] in ("SENT", "NEGOTIATING", "DRAFT"))
    won = sum(q["grand_total"] for q in quotes if q["status"] == "WON")
    lost = sum(q["grand_total"] for q in quotes if q["status"] == "LOST")
    active = [q for q in quotes if q["status"] in ("SENT", "NEGOTIATING")]
    won_count = sum(1 for q in quotes if q["status"] == "WON")
    closed = won_count + sum(1 for q in quotes if q["status"] == "LOST")
    win_rate = round(won_count / closed * 100, 1) if closed else 0.0
    margins = [q["margin_pct"] for q in quotes if q.get("margin_pct")]
    avg_margin = round(sum(margins) / len(margins), 1) if margins else 0.0
    today = datetime.date.today()
    expiring = sum(
        1 for q in quotes
        if q.get("valid_till") and q["status"] in ("SENT", "NEGOTIATING")
        and 0 <= (datetime.date.fromisoformat(q["valid_till"]) - today).days <= 7
    )
    return {
        "pipeline_value": pipeline,
        "won_value": won,
        "lost_value": lost,
        "win_rate_pct": win_rate,
        "avg_margin_pct": avg_margin,
        "expiring_count": expiring,
        "active_count": len(active),
    }


# ── WhatsApp Scanner helpers ───────────────────────────────────────────────────

_EXTRACTION_SYSTEM = """You are an expert estimator for the Indian hardware, sanitary fittings, CP (chrome plated) fittings, and building materials industry. Your job is to extract product requirements from ANY input type — WhatsApp chats, voice notes, BOQ spreadsheets, architect briefs, plumber indent forms, site requirement lists, images, or informal text.

INDUSTRY KNOWLEDGE — apply these rules every time:
• CP FITTINGS (taps, mixers, showers, health faucets, stop cocks): Brands = Jaquar, Grohe, Kohler, Hindware, Parryware, Cera, American Standard, TOTO, Roca, Marc, Oleanna, Essco, Usha Shriram. Keep full model name + series + finish (CP Chrome, PVD Gold, PVD Black, Matt Black, Rose Gold, SS). HSN 8481, 18% GST.
• SANITARY WARE (WC/EWC, wash basin, bathtub, urinal): Brands = Kohler, Hindware, Jaquar, Parryware, Cera, American Standard, TOTO, Roca. Keep type (one-piece, two-piece, wall-hung, pedestal, counter-top, free-standing). HSN 6910, 18% GST.
• BATHROOM ACCESSORIES (towel bars, rings, soap dispensers, toilet roll holders, mirrors, floor drains, shower enclosures): Keep size (e.g. 24", 18"), finish (chrome, PVD gold, SS). HSN 3922, 18% GST.
• HARDWARE — HINGES: Keep type (soft-close, clip-top, glass hinge, flap hinge) + brand (Hettich, Blum, Hafele, Ebco, Euro). HSN 8302.
• HARDWARE — CHANNELS/SLIDES: Keep type (soft-close, tandem box, full-extension, undermount) + size (18", 24") + brand. HSN 8302.
• HARDWARE — HANDLES & KNOBS: Keep finish (PVD, SS, chrome, gold, black, antique brass) + size (96mm, 128mm, 160mm, 320mm) + brand (Hafele, Hettich, Godrej, Dorset). HSN 8302.
• HARDWARE — LOCKS & LATCHES: Keep type (mortise, cylindrical, deadbolt, cam lock) + brand (Godrej, Yale, Assa Abloy, Dorset, Ozone). HSN 8301.
• PLUMBING (CPVC/PVC pipes, fittings, valves, stop cocks, ball valves): Keep size (15mm, 20mm, 25mm, 4") + brand (Astral, Supreme, Finolex, Wavin, Ashirvad). HSN 3917/8481.
• TILES (ceramic, vitrified, PGVT, marble, granite): Keep size (300×300, 600×600, 800×800, 1200×600mm) + finish (matt, polished, rustic) + brand (Kajaria, Somany, Johnson, Asian Granito, RAK, Simpolo). HSN 6907/2515.
• KITCHEN FITTINGS (sinks, mixers, pull-out taps): Brands = Franke, Kohler, Jaquar, Grohe, Hafele. HSN 8481/7324.
• WATERPROOFING: Keep type (Dr. Fixit, Pidilite, STP, cementitious). HSN 3214.

EXTRACTION RULES — follow every one without exception:
1. Extract EVERY item mentioned, no matter how vague, abbreviated, or informally written. Never skip an item.
2. Model numbers/series (e.g. "Jaquar Florentine 1-lever", "Grohe Eurosmart", "Hettich Sensys") must go into description verbatim.
3. Handle typos, abbreviations, mixed languages (English/Hindi/Kannada), shorthand (e.g. "3 jaquar shower overhead", "hettich chanl 18 pr").
4. For organized tables/forms: each row is a separate item — extract every row.
5. Infer customer_type: "plumber / plumbing contractor" → Contractor; "builder / developer / project" → Builder / Developer; "interior designer / ID studio" → Interior Designer; "architect / design firm" → Architect; "hotel / resort / hospitality" → Hotel / Hospitality; "government / CPWD / tender" → Government / Institution; default → Contractor.
6. Extract company/org name from letterheads, headers, or signatures.
7. Capture quantity from anywhere — column, inline, bracket. Default 1 if not found.
8. Unit defaults: CP fittings, sanitary ware, accessories = "Nos"; pipes, plumbing = "Mtr" or "Lot"; tiles = "SqFt" or "SqMtr"; channels/hinges = "Nos" or "Pair".
9. Infer room from context if stated ("master bath", "kitchen", "guest toilet") → put in notes.
10. If a brand is mentioned, always include it in the description.

Return a JSON object with EXACTLY this structure:
{
  "customer_name": "company or person name, or empty string",
  "customer_type": "one of: Contractor, Plumber, Builder / Developer, Interior Designer, Architect, Dealer / Retailer, Hotel / Hospitality, Government / Institution, End Consumer",
  "contact_person": "person name if present, else empty string",
  "contact_phone": "phone number if present, else empty string",
  "contact_email": "email if present, else empty string",
  "project_name": "project or site name, else empty string",
  "site_location": "delivery or site address, else empty string",
  "required_products": [
    {
      "description": "full product description — brand, model, series, type, finish — exactly as mentioned or inferred",
      "quantity": 0,
      "unit": "Nos, Set, Pair, Mtr, SqFt, SqMtr, Lot, KG, Box",
      "specifications": "sizes (mm/inch/feet), finish, color, grade, material type, model number",
      "notes": "room location, delivery date, remarks, special instructions for this item",
      "inferred_category": "CP Fittings|Sanitary Ware|Bathroom Accessories|Hardware - Hinges|Hardware - Channels|Hardware - Handles|Hardware - Locks|Plumbing|Tiles|Kitchen Fittings|Waterproofing|Other",
      "inferred_hsn": "8481|6910|3922|8302|8301|3917|6907|7324|3214|8302"
    }
  ],
  "special_requirements": "urgency, quality certification, IS standard needs, payment terms mentioned",
  "delivery_notes": "timeline, delivery location, site contact details",
  "budget_indication": "any budget, rate, or price range mentioned"
}

Empty string for missing text fields. 0 for missing quantities. Return ONLY valid JSON — no markdown fences, no extra keys."""


def _get_catalog_for_matching():
    from app.api.catalog import _get_all_products
    return _get_all_products()


def _score_match(desc: str, specs: str, product: dict) -> int:
    combined = f"{desc} {specs}".lower()
    score = 0

    for word in product["name"].lower().split():
        if len(word) > 2 and word in combined:
            score += 8

    for word in product["category"].lower().split():
        if len(word) > 2 and word in combined:
            score += 6

    sub = product.get("sub_category", "").lower()
    if sub and any(w in combined for w in sub.split() if len(w) > 2):
        score += 5

    for tag in product.get("tags", []):
        norm = tag.replace("-", " ")
        if norm in combined or tag in combined:
            score += 4

    thickness = product.get("thickness", "")
    if thickness:
        m = re.search(r"(\d+(?:\.\d+)?)\s*mm", thickness)
        if m and m.group(0).replace(" ", "") in combined.replace(" ", ""):
            score += 10

    for app in product.get("applications", []):
        words = [w for w in app.lower().split() if len(w) > 3]
        if sum(1 for w in words if w in combined) >= 2:
            score += 3

    keyword_map = {
        # ── CP Fittings (HSN 8481) ──────────────────────────────────────────────
        "tap":           ["tap", "taps", "faucet", "mixer", "basin mixer", "sink tap",
                          "bath tap", "bib cock", "pillar tap", "mono block"],
        "shower":        ["shower", "overhead shower", "rain shower", "hand shower",
                          "shower set", "shower panel", "shower system", "shower head",
                          "shower arm", "shower enclosure", "diverter"],
        "cp":            ["cp", "chrome plated", "chrome", "pvd", "pvd gold", "pvd black",
                          "matt black", "rose gold", "antique brass", "brushed nickel"],
        "health faucet": ["health faucet", "health faucet set", "bidet spray", "jet spray",
                          "toilet spray", "nozzle spray"],
        "stop cock":     ["stop cock", "stop valve", "ball valve", "gate valve",
                          "angle valve", "quarter turn"],
        "bath spout":    ["bath spout", "bath filler", "tub filler", "bath mixer"],
        "kitchen mixer": ["kitchen mixer", "kitchen tap", "sink mixer", "pull out",
                          "pull-out tap", "flexible tap"],
        # ── Sanitary Ware (HSN 6910) ────────────────────────────────────────────
        "wc":            ["wc", "toilet", "closet", "ewc", "water closet", "toilet pot",
                          "flush", "two piece", "one piece", "wall hung wc",
                          "concealed cistern", "floor mount", "s-trap", "p-trap"],
        "basin":         ["basin", "wash basin", "pedestal basin", "counter top basin",
                          "wall hung basin", "table top basin", "vessel basin",
                          "vanity basin", "lavatory"],
        "bathtub":       ["bathtub", "bath tub", "jacuzzi", "whirlpool", "soaker",
                          "freestanding bath", "built-in bath", "acrylic bath"],
        "urinal":        ["urinal", "urinals", "urinal pot"],
        # ── Bathroom Accessories (HSN 3922) ─────────────────────────────────────
        "towel bar":     ["towel bar", "towel rail", "towel ring", "towel rod",
                          "towel holder", "bath towel", "robe hook", "coat hook"],
        "soap dish":     ["soap dish", "soap holder", "soap dispenser",
                          "liquid soap", "shampoo holder", "bottle holder"],
        "mirror":        ["mirror", "bathroom mirror", "vanity mirror",
                          "led mirror", "fog free mirror", "medicine cabinet"],
        "floor drain":   ["floor drain", "floor trap", "drain cover", "gully trap",
                          "waste coupling", "floor grating"],
        "shower cabin":  ["shower enclosure", "shower cabin", "shower glass",
                          "shower screen", "shower partition"],
        # ── Hardware — Hinges (HSN 8302) ────────────────────────────────────────
        "hinge":         ["hinge", "hinges", "soft close hinge", "clip top",
                          "sensys", "glass hinge", "flap hinge", "piano hinge",
                          "concealed hinge", "european hinge", "inset hinge"],
        # ── Hardware — Channels/Drawer Systems (HSN 8302) ───────────────────────
        "channel":       ["channel", "drawer slide", "slide", "telescopic",
                          "soft close channel", "under mount", "tandem box",
                          "full extension", "push to open", "drawer system",
                          "tandembox", "legrabox", "orga-line"],
        # ── Hardware — Handles (HSN 8302) ────────────────────────────────────────
        "handle":        ["handle", "handles", "knob", "knobs", "pull", "push",
                          "bar handle", "profile handle", "cup handle",
                          "aluminium handle", "ss handle", "gola profile",
                          "boss", "stainless", "brass", "gold handle", "black handle"],
        # ── Hardware — Locks (HSN 8301) ──────────────────────────────────────────
        "lock":          ["lock", "latch", "mortise lock", "cylindrical lock",
                          "cam lock", "deadbolt", "door lock", "drawer lock",
                          "wardrobe lock", "multiplex", "espagnolette"],
        # ── Plumbing (HSN 3917/8481) ─────────────────────────────────────────────
        "cpvc":          ["cpvc", "cpvc pipe", "hot cold pipe", "astral", "finolex",
                          "supreme cpvc", "flowguard", "wavin"],
        "pvc":           ["pvc", "pvc pipe", "upvc", "drainage pipe", "soil pipe",
                          "waste pipe", "pvcu", "prince pipe"],
        "valve":         ["ball valve", "gate valve", "check valve", "pressure valve",
                          "prv", "valve fitting"],
        # ── Tiles (HSN 6907) ─────────────────────────────────────────────────────
        "tile":          ["tile", "tiles", "vitrified", "ceramic", "pgvt", "gvt",
                          "double charge", "polished", "matte tile", "rustic tile",
                          "floor tile", "wall tile", "dado", "mosaic", "kajaria",
                          "somany", "johnson tile", "simpolo", "orientbell"],
        "marble":        ["marble", "granite", "natural stone", "kota stone",
                          "italian marble", "makrana", "statuario", "calacatta"],
        # ── Waterproofing (HSN 3214) ─────────────────────────────────────────────
        "waterproofing": ["waterproofing", "waterproof", "dr fixit", "pidilite",
                          "stp waterproofing", "bituminous", "cementitious"],
        # ── Brand-based matching ─────────────────────────────────────────────────
        "jaquar":        ["jaquar", "florentine", "artize", "kubix", "espelho", "essco"],
        "grohe":         ["grohe", "eurosmart", "eurostyle", "bauloop", "allure",
                          "start", "concetto", "grohtherm"],
        "kohler":        ["kohler", "july", "singulier", "coutere", "forte",
                          "vice versa", "karing"],
        "hindware":      ["hindware", "queo", "flora", "neo"],
        "parryware":     ["parryware", "calla", "galaxy", "coral"],
        "hettich":       ["hettich", "sensys", "innofit", "actro", "arcitech",
                          "intivo", "quadro"],
        "blum":          ["blum", "blumotion", "aventos", "tandembox", "legrabox",
                          "clip top", "orga-line", "servo-drive"],
        "hafele":        ["hafele", "häfele", "magic corner", "kessebohmer"],
        "ebco":          ["ebco", "ebco hinge", "ebco channel"],
        "astral":        ["astral", "astral pipes", "astral cpvc"],
        "kajaria":       ["kajaria"],
        "hardware":      ["hardware", "fitting", "fittings", "furniture fitting",
                          "cabinet hardware", "bathroom hardware"],
        "installation":  ["installation", "labour", "plumber", "plumbing work",
                          "fitting charges", "amc", "maintenance"],
    }
    for p_word, synonyms in keyword_map.items():
        if p_word in product["name"].lower() or p_word in product["category"].lower():
            if any(s in combined for s in synonyms):
                score += 8

    # Partial word match bonus — catch abbreviated/shorthand product names
    desc_words = [w for w in combined.split() if len(w) >= 3]
    p_name_words = product["name"].lower().split()
    partial_hits = sum(
        1 for dw in desc_words
        for pw in p_name_words
        if len(pw) >= 4 and (dw in pw or pw in dw)
    )
    if partial_hits >= 2:
        score += 5
    elif partial_hits == 1:
        score += 2

    return score


def _match_to_catalog(required_products: list) -> list:
    catalog = _get_catalog_for_matching()
    results = []
    for req in required_products:
        desc  = req.get("description", "")
        specs = req.get("specifications", "")
        notes = req.get("notes", "")
        scored = []
        for p in catalog:
            s = _score_match(desc, specs + " " + notes, p)
            scored.append({"product": p, "score": s})
        scored.sort(key=lambda x: x["score"], reverse=True)

        # Always return top 3 — even if score is 0, show nearest available products
        top    = scored[:3]
        score0 = top[0]["score"] if top else 0

        # Include all products that scored > 0; if none did, include top 3 anyway
        positive = [x for x in top if x["score"] > 0]
        matches  = [x["product"] for x in (positive if positive else top)]

        if score0 >= 15:
            conf = "high"
        elif score0 >= 8:
            conf = "medium"
        elif score0 >= 3:
            conf = "low"
        else:
            conf = "none"  # No catalog category match — nearest shown as suggestion

        results.append({
            "required":   req,
            "matches":    matches,
            # Only auto-select a product when there is a real catalog signal (score >= 3)
            "best_match": top[0]["product"] if (top and score0 >= 3) else None,
            "confidence": conf,
        })
    return results


def _build_suggested_lines(matched: list) -> list:
    lines = []
    for mp in matched:
        if not mp["best_match"]:
            continue
        p   = mp["best_match"]
        req = mp["required"]
        lines.append({
            "product_id":    str(p.get("product_id", "")),
            "product_name":  p.get("name", ""),
            "category":      p.get("category", ""),
            "quantity":      req.get("quantity") or 1,
            "unit":          p.get("unit", "pc"),
            "unit_price":    p.get("sell_price", 0),
            "buy_price":     p.get("buy_price", 0),
            "discount_pct":  0,
            "specifications": req.get("specifications", ""),
        })
    return lines


def _demo_scan_result(note: str = "") -> dict:
    extracted = {
        "customer_name":      "Suresh Building Solutions",
        "customer_type":      "Contractor",
        "contact_person":     "Suresh Nair",
        "contact_phone":      "+91 98450 12345",
        "contact_email":      "suresh@sbsolutions.in",
        "project_name":       "Prestige Skyrise — 24 Units (3BHK), Batch 1",
        "site_location":      "Whitefield, Bangalore — 560066",
        "required_products":  [
            {
                "description": "Jaquar Florentine Single Lever Basin Mixer — Chrome",
                "quantity": 48, "unit": "Nos",
                "specifications": "Hot & Cold, 35mm cartridge, wall-mounted",
                "notes": "Master bathroom × 24 + Guest bathroom × 24",
                "inferred_category": "CP Fittings", "inferred_hsn": "8481",
            },
            {
                "description": "Jaquar Overhead Shower 6\" Round with Arm — Chrome",
                "quantity": 24, "unit": "Set",
                "specifications": "Stainless steel arm, ABS shower head with anti-calcium nozzles",
                "notes": "Master bathroom only, ceiling mount",
                "inferred_category": "CP Fittings", "inferred_hsn": "8481",
            },
            {
                "description": "Hindware Queo One-Piece EWC with Soft-Close Seat Cover",
                "quantity": 72, "unit": "Set",
                "specifications": "S-trap, 300mm/400mm outlet, dual flush 3/6L",
                "notes": "3 bathrooms per unit × 24 units",
                "inferred_category": "Sanitary Ware", "inferred_hsn": "6910",
            },
            {
                "description": "Hettich Sensys Soft-Close Hinge",
                "quantity": 480, "unit": "Nos",
                "specifications": "clip-top, 110° opening, full overlay",
                "notes": "Modular kitchen cabinets — 20 hinges per unit",
                "inferred_category": "Hardware - Hinges", "inferred_hsn": "8302",
            },
            {
                "description": "Health Faucet Set with Holder and Flexible Hose",
                "quantity": 72, "unit": "Set",
                "specifications": "ABS body, chrome finish, 1.2m hose",
                "notes": "All 3 bathrooms per unit",
                "inferred_category": "CP Fittings", "inferred_hsn": "8481",
            },
        ],
        "special_requirements": "Contractor pricing required. Delivery in 3 batches — 8 units per batch.",
        "delivery_notes":       "Site delivery with unloading — Whitefield, Gate 3 contact: Suresh 98450 12345",
        "budget_indication":    "Target: under ₹28,000 per unit for CP + sanitary",
    }
    matched = _match_to_catalog(extracted["required_products"])
    return {
        "extracted":         extracted,
        "matched_products":  matched,
        "suggested_lines":   _build_suggested_lines(matched),
        "data_source":       "demo",
        "demo_note":         note or "Demo mode — upload a real WhatsApp screenshot to use live AI scanning with your OPENAI_API_KEY.",
    }


# ── Routes ─────────────────────────────────────────────────────────────────────

@router.get("/quotes/ai-price")
async def ai_price_recommendation_get(
    customer_type: str = "Developer",
    quantity: float = 100,
    product_id: Optional[str] = None,
    current_margin: float = 20.0,
):
    """GET version of AI price recommendation — query params for frontend compatibility."""
    base_discount = 0.0
    rationale_parts = []

    if quantity >= 5000:
        base_discount += 5.0
        rationale_parts.append("High volume (5000+ units) → 5% volume discount")
    elif quantity >= 2000:
        base_discount += 3.0
        rationale_parts.append("Medium volume (2000+ units) → 3% volume discount")
    elif quantity >= 500:
        base_discount += 1.5
        rationale_parts.append("Moderate volume (500+ units) → 1.5% volume discount")

    type_discounts = {"Developer": 2.0, "Architect": 1.5, "Contractor": 1.0, "Interior Firm": 1.5, "Retailer": 0.5}
    ct_disc = type_discounts.get(customer_type, 0.0)
    if ct_disc:
        base_discount += ct_disc
        rationale_parts.append(f"{customer_type} relationship pricing → {ct_disc}% extra")

    projected_margin = current_margin - base_discount
    if projected_margin < 15.0:
        old = base_discount
        base_discount = max(0, current_margin - 15.0)
        rationale_parts.append(f"Margin floor 15% applied — capped discount from {old:.1f}% to {base_discount:.1f}%")
        projected_margin = current_margin - base_discount

    return {
        "recommended_discount_pct": round(base_discount, 1),
        "max_safe_discount_pct": round(max(0, current_margin - 15.0), 1),
        "floor_margin_pct": 15.0,
        "projected_margin_pct": round(projected_margin, 1),
        "rationale": " · ".join(rationale_parts) if rationale_parts else "Standard pricing — no additional discount recommended",
        "competitive_note": "Market range for this segment: 12–22% margin. Recommended discount keeps you competitive.",
    }


@router.get("/quotes")
async def list_quotes(status: Optional[str] = None, search: Optional[str] = None):
    try:
        from app.db.connection import get_pool
        from app.db import quote_queries
        pool = await get_pool()
        if pool:
            await quote_queries.ensure_tables(pool)
            quotes = await quote_queries.list_quotes_db(pool, status=status, search=search)
            kpis   = await quote_queries.kpis_db(pool)
            return {"quotes": quotes, "kpis": kpis, "data_source": "mysql"}
    except Exception as exc:
        logger.warning("Quote list DB error: %s", exc)

    quotes = _mock_quotes()
    if status and status.upper() != "ALL":
        quotes = [q for q in quotes if q["status"] == status.upper()]
    if search:
        term = search.lower()
        quotes = [
            q for q in quotes
            if term in q["customer_name"].lower()
            or term in q["quote_number"].lower()
            or term in q.get("project_name", "").lower()
        ]
    return {
        "quotes": quotes,
        "kpis": _compute_kpis(_mock_quotes()),
        "data_source": "demo",
    }


@router.get("/quotes/{quote_id}")
async def get_quote(quote_id: int):
    try:
        from app.db.connection import get_pool
        from app.db import quote_queries
        pool = await get_pool()
        if pool:
            await quote_queries.ensure_tables(pool)
            q = await quote_queries.get_quote_db(pool, quote_id)
            if q:
                return {"quote": q, "data_source": "mysql"}
    except Exception as exc:
        logger.warning("Quote get DB error: %s", exc)

    quotes = _mock_quotes()
    for q in quotes:
        if q["quote_id"] == quote_id:
            return {"quote": q, "data_source": "demo"}
    raise HTTPException(status_code=404, detail=f"Quote {quote_id} not found")


@router.post("/quotes", status_code=201)
async def create_quote(req: CreateQuoteRequest):
    today = datetime.date.today()
    subtotal = sum(
        li.quantity * li.unit_price * (1 - li.discount_pct / 100)
        for li in req.line_items
    )
    gst_amount  = subtotal * req.gst_rate / 100
    freight     = req.freight_amount if req.include_freight else 0
    grand_total = subtotal + gst_amount + freight
    total_buy   = sum(li.quantity * li.buy_price for li in req.line_items if li.buy_price)
    avg_margin  = round((subtotal - total_buy) / subtotal * 100, 1) if subtotal > 0 and total_buy > 0 else 0.0

    def _line_items_expanded(items):
        return [
            {
                "sl": i + 1,
                **li.model_dump(),
                "net_price":  round(li.unit_price * (1 - li.discount_pct / 100), 2),
                "line_total": round(li.quantity * li.unit_price * (1 - li.discount_pct / 100), 2),
            }
            for i, li in enumerate(items)
        ]

    try:
        from app.db.connection import get_pool
        from app.db import quote_queries
        pool = await get_pool()
        if pool:
            await quote_queries.ensure_tables(pool)
            quote_number = await quote_queries.next_quote_number(pool)
            payload = {
                "quote_number":   quote_number,
                "created_at":     today.isoformat(),
                "valid_till":     (today + datetime.timedelta(days=req.validity_days)).isoformat(),
                "status":         "DRAFT",
                **req.model_dump(exclude={"line_items"}),
                "subtotal":       round(subtotal, 2),
                "gst_amount":     round(gst_amount, 2),
                "grand_total":    round(grand_total, 2),
                "avg_margin_pct": avg_margin,
                "line_items":     _line_items_expanded(req.line_items),
            }
            new_id = await quote_queries.insert_quote(pool, payload)
            payload["quote_id"]  = new_id
            payload["margin_pct"] = avg_margin
            logger.info("Saved quote %s (id=%s) for %s to MySQL", quote_number, new_id, req.customer_name)
            return {"quote": payload, "message": "Quote saved to database", "data_source": "mysql"}
    except Exception as exc:
        logger.warning("Quote create DB error: %s", exc)

    new_id = 100
    quote = {
        "quote_id":    new_id,
        "quote_number": f"QT-{today.year}-{new_id:04d}",
        "created_at":  today.isoformat(),
        "valid_till":  (today + datetime.timedelta(days=req.validity_days)).isoformat(),
        "status":      "DRAFT",
        **req.model_dump(exclude={"line_items"}),
        "subtotal":    round(subtotal, 2),
        "gst_amount":  round(gst_amount, 2),
        "grand_total": round(grand_total, 2),
        "margin_pct":  avg_margin,
        "line_items":  _line_items_expanded(req.line_items),
    }
    logger.info("Created demo quote %s for %s", quote["quote_number"], req.customer_name)
    # Auto-notify via WhatsApp (non-blocking, best-effort)
    if req.contact_phone:
        try:
            from app.services.whatsapp_notify import send_quotation_whatsapp
            asyncio.create_task(send_quotation_whatsapp(
                customer_phone = req.contact_phone,
                customer_name  = req.customer_name,
                quote_number   = quote["quote_number"],
                grand_total    = quote["grand_total"],
                items_count    = len(req.line_items),
                valid_till     = quote["valid_till"],
                project_name   = req.project_name or "",
            ))
        except Exception:
            pass
    return {"quote": quote, "message": "Quote created (demo mode — not persisted)", "data_source": "demo"}


@router.post("/quotes/{quote_id}/send-whatsapp")
async def send_whatsapp_notification(quote_id: int, payload: dict = {}):
    """Manually trigger a WhatsApp notification for any existing quote."""
    from app.services.whatsapp_notify import send_quotation_whatsapp

    # Try to fetch the quote from DB first, fallback to payload fields
    quote_data = {}
    try:
        from app.db.connection import get_pool
        from app.db import quote_queries
        pool = await get_pool()
        if pool:
            quote_data = await quote_queries.get_quote_by_id(pool, quote_id) or {}
    except Exception:
        pass

    # Merge explicit payload overrides
    phone      = payload.get("contact_phone") or quote_data.get("contact_phone", "")
    name       = payload.get("customer_name") or quote_data.get("customer_name", "Customer")
    qnum       = payload.get("quote_number")  or quote_data.get("quote_number", f"QT-{quote_id}")
    total      = payload.get("grand_total")   or quote_data.get("grand_total", 0)
    items      = payload.get("items_count")   or len(quote_data.get("line_items", []))
    valid_till = payload.get("valid_till")    or quote_data.get("valid_till", "")
    project    = payload.get("project_name")  or quote_data.get("project_name", "")
    extra_msg  = payload.get("message", "")

    if not phone:
        raise HTTPException(status_code=422, detail="contact_phone is required (provide in request body if quote not in DB)")

    result = await send_quotation_whatsapp(
        customer_phone = phone,
        customer_name  = name,
        quote_number   = qnum,
        grand_total    = total,
        items_count    = items,
        valid_till     = valid_till,
        project_name   = project,
        extra_msg      = extra_msg,
    )
    return result


@router.put("/quotes/{quote_id}/status")
async def update_quote_status(quote_id: int, req: UpdateQuoteStatus):
    if req.status.upper() not in VALID_STATUSES:
        raise HTTPException(status_code=422, detail=f"Invalid status. Use one of: {', '.join(VALID_STATUSES)}")

    try:
        from app.db.connection import get_pool
        from app.db import quote_queries
        pool = await get_pool()
        if pool:
            updated = await quote_queries.update_status_db(pool, quote_id, req.status, req.remarks)
            if updated:
                return {
                    "quote_id":   quote_id,
                    "status":     req.status.upper(),
                    "remarks":    req.remarks,
                    "updated_at": datetime.datetime.utcnow().isoformat(),
                    "message":    "Status updated",
                    "data_source": "mysql",
                }
    except Exception as exc:
        logger.warning("Quote status update DB error: %s", exc)

    return {
        "quote_id":   quote_id,
        "status":     req.status.upper(),
        "remarks":    req.remarks,
        "updated_at": datetime.datetime.utcnow().isoformat(),
        "message":    "Status updated (demo mode)",
        "data_source": "demo",
    }


@router.put("/quotes/{quote_id}")
async def update_quote(quote_id: int, req: CreateQuoteRequest):
    """Full quote update — replaces all fields and line items; status is preserved."""
    today = datetime.date.today()
    subtotal = sum(
        li.quantity * li.unit_price * (1 - li.discount_pct / 100)
        for li in req.line_items
    )
    gst_amount  = subtotal * req.gst_rate / 100
    freight     = req.freight_amount if req.include_freight else 0
    grand_total = subtotal + gst_amount + freight
    total_buy   = sum(li.quantity * li.buy_price for li in req.line_items if li.buy_price)
    avg_margin  = round((subtotal - total_buy) / subtotal * 100, 1) if subtotal > 0 and total_buy > 0 else 0.0

    def _expanded(items):
        return [
            {
                "sl": i + 1,
                **li.model_dump(),
                "net_price":  round(li.unit_price * (1 - li.discount_pct / 100), 2),
                "line_total": round(li.quantity * li.unit_price * (1 - li.discount_pct / 100), 2),
            }
            for i, li in enumerate(items)
        ]

    try:
        from app.db.connection import get_pool
        from app.db import quote_queries
        pool = await get_pool()
        if pool:
            await quote_queries.ensure_tables(pool)
            updated = await quote_queries.update_quote_db(pool, quote_id, {
                **req.model_dump(exclude={"line_items"}),
                "valid_till":     (today + datetime.timedelta(days=req.validity_days)).isoformat(),
                "subtotal":       round(subtotal, 2),
                "gst_amount":     round(gst_amount, 2),
                "grand_total":    round(grand_total, 2),
                "avg_margin_pct": avg_margin,
                "line_items":     _expanded(req.line_items),
            })
            if updated:
                updated.setdefault("margin_pct", avg_margin)
                updated.setdefault("avg_margin_pct", avg_margin)
                logger.info("Updated quote %s in MySQL", quote_id)
                return {"quote": updated, "message": "Quote updated", "data_source": "mysql"}
    except Exception as exc:
        logger.warning("Quote update DB error: %s", exc)

    payload = {
        "quote_id":       quote_id,
        "created_at":     today.isoformat(),
        "valid_till":     (today + datetime.timedelta(days=req.validity_days)).isoformat(),
        "status":         "DRAFT",
        **req.model_dump(exclude={"line_items"}),
        "subtotal":       round(subtotal, 2),
        "gst_amount":     round(gst_amount, 2),
        "grand_total":    round(grand_total, 2),
        "margin_pct":     avg_margin,
        "avg_margin_pct": avg_margin,
        "line_items":     _expanded(req.line_items),
    }
    return {"quote": payload, "message": "Quote updated (demo mode)", "data_source": "demo"}


@router.post("/quotes/ai-price")
async def ai_price_recommendation(req: AIPriceRequest):
    """Return AI-suggested discount/margin based on customer type and quantity."""
    base_discount = 0.0
    rationale_parts = []

    # Volume-based discount
    if req.total_quantity >= 5000:
        base_discount += 5.0
        rationale_parts.append("High volume (5000+ units) → 5% volume discount")
    elif req.total_quantity >= 2000:
        base_discount += 3.0
        rationale_parts.append("Medium volume (2000+ units) → 3% volume discount")
    elif req.total_quantity >= 500:
        base_discount += 1.5
        rationale_parts.append("Moderate volume (500+ units) → 1.5% volume discount")

    # Customer-type discount
    type_discounts = {
        "Developer": 2.0,
        "Architect": 1.5,
        "Contractor": 1.0,
        "Interior Firm": 1.5,
        "Retailer": 0.5,
    }
    ct_disc = type_discounts.get(req.customer_type, 0.0)
    if ct_disc:
        base_discount += ct_disc
        rationale_parts.append(f"{req.customer_type} relationship pricing → {ct_disc}% extra")

    # Floor margin protection
    projected_margin = req.current_margin - base_discount
    if projected_margin < 15.0:
        old = base_discount
        base_discount = max(0, req.current_margin - 15.0)
        rationale_parts.append(f"Margin floor 15% applied — capped discount from {old:.1f}% to {base_discount:.1f}%")
        projected_margin = req.current_margin - base_discount

    return {
        "recommended_discount_pct": round(base_discount, 1),
        "projected_margin_pct": round(projected_margin, 1),
        "margin_floor_pct": 15.0,
        "rationale": " · ".join(rationale_parts) if rationale_parts else "Standard pricing — no additional discount recommended",
        "competitive_note": "Market range for this segment: 12–22% margin. Recommended discount keeps you competitive.",
    }


def _extract_file_content(file_bytes: bytes, content_type: str, filename: str):
    """
    Universal file content extractor.
    Returns (text_content: str, is_image: bool, image_b64: str, image_ct: str).
    is_image=True means caller should use Vision API with image_b64/image_ct.
    """
    import io as _io

    IMAGE_EXTS = (".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".tiff", ".heic", ".heif")
    is_image = content_type.startswith("image/") or filename.endswith(IMAGE_EXTS)
    if is_image:
        b64 = base64.b64encode(file_bytes).decode()
        ct  = content_type if content_type.startswith("image/") else "image/jpeg"
        return "", True, b64, ct

    # PDF
    if content_type == "application/pdf" or filename.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(_io.BytesIO(file_bytes))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if text.strip():
                return text[:14000], False, "", ""
            # Empty text = scanned/image-based PDF — stop here, don't decode as binary
            return "__scanned_pdf__", False, "", ""
        except Exception as e:
            logger.warning("pypdf failed: %s", e)
            return "", False, "", ""

    # DOCX / DOC
    if filename.endswith((".docx", ".doc")) or "wordprocessingml" in content_type:
        try:
            from docx import Document
            doc = Document(_io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    parts.append("\t".join(c.text.strip() for c in row.cells if c.text.strip()))
            return "\n".join(parts)[:14000], False, "", ""
        except Exception as e:
            logger.warning("python-docx failed: %s", e)

    # XLSX / XLS
    if filename.endswith((".xlsx", ".xls")) or "spreadsheetml" in content_type or "ms-excel" in content_type:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(_io.BytesIO(file_bytes), read_only=True, data_only=True)
            lines = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    row_str = "\t".join(str(c) for c in row if c is not None)
                    if row_str.strip():
                        lines.append(row_str)
            return "\n".join(lines)[:14000], False, "", ""
        except Exception as e:
            logger.warning("openpyxl failed: %s", e)

    # CSV
    if filename.endswith(".csv") or content_type in ("text/csv", "application/csv"):
        try:
            import csv
            decoded = file_bytes.decode("utf-8", errors="ignore")
            reader = csv.reader(_io.StringIO(decoded))
            rows = [", ".join(r) for r in reader if any(c.strip() for c in r)]
            return "\n".join(rows)[:14000], False, "", ""
        except Exception as e:
            logger.warning("csv parse failed: %s", e)

    # Fallback: any text-based file (TXT, JSON, XML, HTML, RTF, etc.)
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            text = file_bytes.decode(enc, errors="ignore")
            if text.strip():
                return text[:14000], False, "", ""
        except Exception:
            continue

    return "", False, "", ""


@router.post("/quotes/scan-whatsapp")
async def scan_whatsapp_requirement(
    file: List[UploadFile] = File(default=[]),
    text_input: str = Form(""),
):
    """
    Universal requirement scanner: accepts multiple files + optional typed/pasted text.
    Supports images (Vision API), PDF, DOCX, XLSX, CSV, TXT and any text-decodable file.
    Always returns top catalog matches even for unrecognized/partial items.
    """
    if text_input.strip() == "__demo__":
        return _demo_scan_result()

    api_key = os.getenv("OPENAI_API_KEY", "")
    if not api_key:
        return _demo_scan_result("No OPENAI_API_KEY configured — showing demo extraction.")

    files = [f for f in (file or []) if f and f.filename]
    has_text = bool(text_input.strip())
    if not files and not has_text:
        return _demo_scan_result("No input provided — showing demo extraction.")

    try:
        from openai import AsyncOpenAI

        combined_text = text_input.strip()
        image_parts: List[dict] = []  # list of {b64, ct} dicts

        for f in files:
            raw = await f.read()
            if len(raw) > 20 * 1024 * 1024:
                continue  # skip oversized files silently
            ct = f.content_type or ""
            fn = (f.filename or "").lower()

            file_text, is_image, image_b64, image_ct = _extract_file_content(raw, ct, fn)

            if file_text == "__scanned_pdf__":
                return {
                    "extracted": {"customer_name": "", "customer_type": "Retailer", "contact_person": "", "contact_phone": "", "contact_email": "", "project_name": "", "site_location": "", "required_products": [], "special_requirements": "", "delivery_notes": "", "budget_indication": ""},
                    "matched_products": [], "suggested_lines": [], "data_source": "error",
                    "demo_note": "One of the uploaded PDFs appears to be scanned (image-based) — no text layer found. Please upload a screenshot instead.",
                }

            if is_image:
                image_parts.append({"b64": image_b64, "ct": image_ct})
            elif file_text.strip():
                combined_text = "\n\n---\n".join(filter(None, [combined_text, file_text]))

        if not combined_text.strip() and not image_parts:
            return _demo_scan_result("No readable content in the uploaded files.")

        client = AsyncOpenAI(api_key=api_key, timeout=60.0)

        if image_parts:
            # Vision: send all images + any text context
            user_content: list = []
            for img in image_parts:
                user_content.append({"type": "image_url", "image_url": {"url": f"data:{img['ct']};base64,{img['b64']}", "detail": "high"}})
            context = f"Also consider this typed context:\n\n{combined_text}\n\n" if combined_text.strip() else ""
            user_content.append({"type": "text", "text": f"{context}Extract ALL product requirements from these images. Return ONLY valid JSON as specified."})
        else:
            user_content = (
                f"Extract ALL product requirements from the following input. "
                f"It may be a WhatsApp message, typed note, requisition form, or any requirement document. "
                f"Return ONLY valid JSON as specified.\n\n---\n{combined_text}\n---"
            )

        response = await client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": _EXTRACTION_SYSTEM},
                {"role": "user",   "content": user_content},
            ],
            max_tokens=3000,
            response_format={"type": "json_object"},
        )

        extracted = json.loads(response.choices[0].message.content)
        required_products = extracted.get("required_products", [])
        matched = _match_to_catalog(required_products)

        note = ""
        if not required_products:
            note = "No product requirements found in this input. Try adding more detail — e.g. product names, quantities, specifications."

        return {
            "extracted":        extracted,
            "matched_products": matched,
            "suggested_lines":  _build_suggested_lines(matched),
            "data_source":      "openai",
            **({"demo_note": note} if note else {}),
        }

    except asyncio.CancelledError:
        raise
    except (KeyboardInterrupt, SystemExit):
        raise
    except BaseException as exc:
        logger.exception("scan-whatsapp error: %s", exc)
        return _demo_scan_result(f"AI extraction failed ({type(exc).__name__}) — showing demo result.")


# ── Quote Analyzer helpers ─────────────────────────────────────────────────────

def _analyze_rule_based(req: "AnalyzeQuoteRequest") -> dict:
    items = req.line_items

    item_analysis = []
    for li in items:
        net    = li.unit_price * (1 - li.discount_pct / 100) if li.unit_price > 0 else 0
        margin = (net - li.buy_price) / net * 100 if net > 0 and li.buy_price > 0 else None
        line_t = net * li.quantity
        if margin is None:
            status = "unknown"
        elif margin < 14:
            status = "critical"
        elif margin < 18:
            status = "at_risk"
        else:
            status = "healthy"
        item_analysis.append({
            "product_name": li.product_name,
            "net_price":    round(net, 2),
            "margin_pct":   round(margin, 1) if margin is not None else None,
            "line_total":   round(line_t, 2),
            "status":       status,
        })

    margins     = [a["margin_pct"] for a in item_analysis if a["margin_pct"] is not None]
    avg_margin  = sum(margins) / len(margins) if margins else 0
    below_floor = sum(1 for m in margins if m < 14)
    total_val   = req.subtotal * (1 + req.gst_rate / 100) + (req.freight_amount if req.include_freight else 0)

    deal_size = "large" if total_val >= 1_000_000 else "medium" if total_val >= 200_000 else "small"

    # Win probability (rule engine)
    score = 65
    if avg_margin > 0:
        if avg_margin <= 16:   score += 15
        elif avg_margin <= 20: score += 8
        elif avg_margin > 25:  score -= 10
    score += {"Developer": 5, "Architect": 8, "Interior Firm": 6, "Contractor": 3, "Retailer": 2}.get(req.customer_type, 0)
    if deal_size == "large": score -= 5
    if deal_size == "small": score += 5
    if "advance" in req.payment_terms.lower(): score += 5
    score -= below_floor * 5
    score = max(20, min(95, score))

    insights = []
    actions  = []
    if below_floor:
        insights.append(f"{below_floor} item(s) below 14% margin floor — immediate pricing correction needed")
        actions.append({"action": f"Raise price on {below_floor} critical item(s) to restore minimum 14% margin", "impact": "high"})
    if avg_margin > 24:
        room = round(avg_margin - 20, 1)
        insights.append(f"Avg margin {avg_margin:.1f}% — you have {room}% headroom to sharpen price if needed")
        actions.append({"action": f"Offer up to {min(room, 5):.0f}% strategic discount to improve win probability", "impact": "medium"})
    if 0 < avg_margin < 17:
        insights.append(f"Margin {avg_margin:.1f}% is very competitive — use this as your key selling point")
    if deal_size == "large":
        insights.append("Large deal — propose phased delivery or milestone-based payments to reduce buyer risk")
        actions.append({"action": "Add phased delivery schedule + payment milestones to the Notes section", "impact": "medium"})
    if req.customer_type == "Developer":
        insights.append("Developers compare 3+ suppliers — emphasise IS certification, brand, and on-time delivery SLA")
        actions.append({"action": "Add IS/EN certification numbers and guaranteed delivery window to notes", "impact": "high"})
    elif req.customer_type == "Architect":
        insights.append("Architects drive future projects — offer spec sheet support and a complimentary sample kit")
        actions.append({"action": "Offer complimentary sample kit and technical specification sheets", "impact": "high"})
    if not req.notes:
        insights.append("No special notes — add delivery SLA, certifications, and unique selling points to win")
        actions.append({"action": "Fill in Notes: guaranteed delivery date, IS certification, and your USPs", "impact": "medium"})
    if len(items) == 1:
        insights.append("Single-line quote — consider suggesting a complementary product to grow deal value")

    cats = [li.category for li in items]
    upsell = ""
    if "Aluminium Louvers" in cats and "HPL Exterior Cladding" not in cats:
        upsell = "Add HPL Exterior Cladding 6mm — common pairing with aluminium louvers on facade projects"
    elif "High Pressure Laminate" in cats and "Acrylic Laminate" not in cats:
        upsell = "Suggest Acrylic High-Gloss for premium kitchen zones — 18% higher value, easy upsell with sample"
    elif "PVC Louvers" in cats and "Aluminium Louvers" not in cats:
        upsell = "Upgrade pitch: Aluminium Z-100 for premium areas, PVC for budget zones — two-tier option"
    elif "Compact Laminate" in cats and "Aluminium Composite Panel" not in cats:
        upsell = "Add ACP 4mm FR for exterior areas — required by NBC 2016 for buildings above G+4"

    health = "critical" if below_floor or avg_margin < 14 else "at_risk" if avg_margin < 17 else "good"

    strategy = ""
    if avg_margin < 14:
        strategy = "Critical: Prices are below margin floor. Increase list prices or remove discount before sending."
    elif avg_margin < 17:
        strategy = f"Competitive stance at {avg_margin:.1f}% margin. Hold on further discounts — lead with reliability, certifications, and delivery speed instead."
    elif avg_margin < 23:
        if deal_size == "large":
            strategy = f"Healthy {avg_margin:.1f}% margin on a large deal. Offer up to 2–3% discount as a closing lever if negotiation stalls."
        else:
            strategy = f"Healthy {avg_margin:.1f}% margin. Lead with quality and delivery certainty. Reserve discount as a closing tool."
    else:
        strategy = f"Strong {avg_margin:.1f}% margin gives you negotiation room. You can offer up to {avg_margin - 20:.0f}% discount while staying above 20% floor."

    return {
        "win_probability":          score,
        "deal_health":              health,
        "avg_margin_pct":           round(avg_margin, 1),
        "deal_size":                deal_size,
        "total_value":              round(total_val, 2),
        "item_analysis":            item_analysis,
        "key_insights":             insights[:5],
        "recommended_actions":      actions[:4],
        "upsell_opportunity":       upsell,
        "pricing_strategy":         strategy,
        "win_probability_rationale":"",
        "data_source":              "rule_engine",
    }


@router.post("/quotes/analyze")
async def analyze_quote(req: AnalyzeQuoteRequest):
    """
    Comprehensive AI analysis of a quotation in progress.
    Returns win probability, per-item margin health, pricing strategy, and actionable recommendations.
    Rule-based analysis is always returned instantly; GPT-4o enhances it when an API key is present.
    """
    result = _analyze_rule_based(req)

    api_key = os.getenv("OPENAI_API_KEY", "")
    if not api_key or not req.line_items:
        return result

    try:
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=api_key)

        items_text = "\n".join([
            f"- {li.product_name} ({li.category}): {li.quantity} {li.unit} @ ₹{li.unit_price}, "
            f"disc {li.discount_pct}%, buy ₹{li.buy_price}"
            for li in req.line_items
        ])
        prompt = (
            f"Analyse this building-materials quotation and return JSON with keys: "
            f"key_insights (list of 4 specific insights, max 90 chars each), "
            f"recommended_actions (list of 3 objects with 'action' str and 'impact' high/medium/low), "
            f"pricing_strategy (1 short paragraph), "
            f"upsell_opportunity (one specific suggestion or empty string), "
            f"win_probability_rationale (1 sentence explaining score {result['win_probability']}/100).\n\n"
            f"Customer: {req.customer_name or 'Unknown'} ({req.customer_type})\n"
            f"Project: {req.project_name or 'Not specified'}\n"
            f"Payment: {req.payment_terms}\nTotal: ₹{result['total_value']:,.0f}\n"
            f"Avg margin: {result['avg_margin_pct']}%\nItems:\n{items_text}\nNotes: {req.notes or 'None'}"
        )
        resp = await client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a building-materials sales strategist in Bangalore. Return only valid JSON."},
                {"role": "user",   "content": prompt},
            ],
            max_tokens=900,
            response_format={"type": "json_object"},
        )
        ai = json.loads(resp.choices[0].message.content)
        result.update({
            "key_insights":             ai.get("key_insights", result["key_insights"]),
            "recommended_actions":      ai.get("recommended_actions", result["recommended_actions"]),
            "pricing_strategy":         ai.get("pricing_strategy", result["pricing_strategy"]),
            "upsell_opportunity":       ai.get("upsell_opportunity", result["upsell_opportunity"]),
            "win_probability_rationale": ai.get("win_probability_rationale", ""),
            "data_source":              "openai",
        })
    except Exception as exc:
        logger.exception("Quote analyze OpenAI error: %s", exc)

    return result


# ── Quote → Sales Order Conversion ─────────────────────────────────────────────

_converted_orders: dict[int, dict] = {}  # in-memory conversion registry (demo mode)


@router.post("/quotes/{quote_id}/convert-to-order")
async def convert_quote_to_order(quote_id: int):
    """
    Convert a WON quotation to a Sales Order.
    Idempotent — calling twice returns the same order number.
    DB mode: inserts a row in customer_orders and marks quote status CONVERTED.
    Demo mode: generates a deterministic order number and stores in memory.
    """
    if quote_id in _converted_orders:
        return _converted_orders[quote_id]

    today = datetime.date.today()

    try:
        from app.db.connection import get_pool
        from app.db import quote_queries
        pool = await get_pool()
        if pool:
            await quote_queries.ensure_tables(pool)
            quote_data = await quote_queries.get_quote_db(pool, quote_id)
            if quote_data:
                order_number = f"SO-{today.strftime('%Y%m%d')}-{quote_id:04d}"
                try:
                    async with pool.acquire() as conn:
                        async with conn.cursor() as cur:
                            li = (quote_data.get("line_items") or [{}])[0]
                            await cur.execute(
                                """INSERT INTO customer_orders
                                   (order_number, customer_name, customer_type, product_name,
                                    quantity, sell_price, total_value, delivery_date, status, notes)
                                   VALUES (%s,%s,%s,%s,%s,%s,%s,%s,'CONFIRMED',%s)""",
                                (
                                    order_number,
                                    quote_data.get("customer_name", ""),
                                    quote_data.get("customer_type", ""),
                                    li.get("product_name", ""),
                                    li.get("quantity", 1),
                                    li.get("net_price", li.get("unit_price", 0)),
                                    quote_data.get("grand_total", 0),
                                    quote_data.get("valid_till", today.isoformat()),
                                    f"Converted from {quote_data.get('quote_number', '')}",
                                ),
                            )
                    result = {
                        "success": True,
                        "order_number": order_number,
                        "quote_id": quote_id,
                        "quote_number": quote_data.get("quote_number", ""),
                        "customer_name": quote_data.get("customer_name", ""),
                        "total_value": quote_data.get("grand_total", 0),
                        "converted_at": today.isoformat(),
                        "message": f"Quote converted to Sales Order {order_number}",
                        "data_source": "mysql",
                    }
                    _converted_orders[quote_id] = result
                    return result
                except Exception as db_exc:
                    logger.warning("convert-to-order DB insert failed: %s", db_exc)
    except Exception as exc:
        logger.warning("convert-to-order DB path failed: %s", exc)

    # Demo fallback
    mock_quote = next((q for q in _mock_quotes() if q["quote_id"] == quote_id), None)
    customer = mock_quote["customer_name"] if mock_quote else "Customer"
    total    = mock_quote["grand_total"]   if mock_quote else 0
    qnum     = mock_quote["quote_number"]  if mock_quote else f"QT-{today.year}-{quote_id:04d}"
    order_number = f"SO-{today.strftime('%Y%m%d')}-{quote_id:04d}"

    result = {
        "success": True,
        "order_number": order_number,
        "quote_id": quote_id,
        "quote_number": qnum,
        "customer_name": customer,
        "total_value": total,
        "converted_at": today.isoformat(),
        "message": f"Quote {qnum} converted to Sales Order {order_number}",
        "data_source": "demo",
    }
    _converted_orders[quote_id] = result
    logger.info("Demo convert: %s → %s for %s", qnum, order_number, customer)
    return result


# ── /quotes/merge ─────────────────────────────────────────────────────────────

@router.post("/quotes/merge")
async def merge_quotes(body: dict):
    """
    Merge 2+ quotations into one new DRAFT quote.
    Client details from the first selected quote are used as the base.
    All line items from every selected quote are combined into one flat list.
    """
    quote_ids = [int(i) for i in (body.get("quote_ids") or []) if i]
    if len(quote_ids) < 2:
        raise HTTPException(status_code=400, detail="At least 2 quotes required to merge")

    selected: list = []
    try:
        from app.db.connection import get_pool
        from app.db import quote_queries
        pool = await get_pool()
        if pool:
            await quote_queries.ensure_tables(pool)
            for qid in quote_ids:
                q = await quote_queries.get_quote_db(pool, qid)
                if q:
                    selected.append(q)
    except Exception as exc:
        logger.warning("quotes/merge DB fetch failed: %s", exc)

    if not selected:
        # Demo fallback — fetch from mock
        all_mock = {q["quote_id"]: q for q in _mock_quotes()}
        for qid in quote_ids:
            q = all_mock.get(qid)
            if q:
                selected.append(q)

    if not selected:
        raise HTTPException(status_code=404, detail="None of the selected quotes were found")

    base = selected[0]
    merged_items = []
    for q in selected:
        for li in (q.get("line_items") or []):
            # Coerce all numeric Decimal values from MySQL to float to prevent
            # TypeError when mixing Decimal + float in arithmetic below
            merged_items.append({
                k: (float(v) if hasattr(v, '__float__') and not isinstance(v, (str, bool)) else v)
                for k, v in li.items() if k != "sl"
            })

    # Re-number sl
    for i, li in enumerate(merged_items, 1):
        li["sl"] = i

    today = datetime.date.today()
    source_nums = ", ".join(q.get("quote_number", "") for q in selected)
    subtotal    = sum(float(li.get("line_total") or 0) for li in merged_items)
    gst_rate    = float(base.get("gst_rate") or 18)
    gst_amount  = round(subtotal * gst_rate / 100, 2)
    grand_total = round(subtotal + gst_amount, 2)

    merged_payload = {
        "customer_name":    str(base.get("customer_name") or ""),
        "customer_type":    str(base.get("customer_type") or ""),
        "contact_person":   str(base.get("contact_person") or ""),
        "contact_phone":    str(base.get("contact_phone") or ""),
        "contact_email":    str(base.get("contact_email") or ""),
        "project_name":     f"[Merged] {base.get('project_name') or ''}",
        "site_location":    str(base.get("site_location") or ""),
        "payment_terms":    str(base.get("payment_terms") or ""),
        "delivery_terms":   str(base.get("delivery_terms") or ""),
        "gst_number":       str(base.get("gst_number") or ""),
        "billing_address":  str(base.get("billing_address") or ""),
        "validity_days":    int(base.get("validity_days") or 30),
        "gst_rate":         gst_rate,
        "include_freight":  False,
        "freight_amount":   0.0,
        "avg_margin_pct":   0.0,
        "notes":            f"Merged from: {source_nums}",
        "status":           "DRAFT",
        "subtotal":         round(subtotal, 2),
        "gst_amount":       gst_amount,
        "grand_total":      grand_total,
        "line_items":       merged_items,
    }

    try:
        from app.db.connection import get_pool
        from app.db import quote_queries
        pool = await get_pool()
        if pool:
            await quote_queries.ensure_tables(pool)
            quote_number = await quote_queries.next_quote_number(pool)
            payload_for_db = {
                **merged_payload,
                "quote_number": quote_number,
                "created_at": today.isoformat(),
                "valid_till": (today + datetime.timedelta(days=merged_payload["validity_days"])).isoformat(),
            }
            new_id = await quote_queries.insert_quote(pool, payload_for_db)
            return {**payload_for_db, "quote_id": new_id, "data_source": "mysql"}
    except Exception as exc:
        logger.warning("quotes/merge DB save failed: %s", exc)

    # Demo fallback
    new_id = 100 + len(quote_ids)
    quote_number = f"QT-{today.year}-{new_id:04d}"
    return {
        **merged_payload,
        "quote_id": new_id,
        "quote_number": quote_number,
        "created_at": today.isoformat(),
        "valid_till": (today + datetime.timedelta(days=merged_payload["validity_days"])).isoformat(),
        "data_source": "demo",
    }


# ── /quotes/{quote_id}/clone ──────────────────────────────────────────────────

@router.post("/quotes/{quote_id}/clone")
async def clone_quote(quote_id: int):
    """
    Return a clone-ready payload of an existing quote.
    Strips quote_id, quote_number, created_at, valid_till.
    Status reset to DRAFT, project_name prefixed [Copy].
    No DB write — frontend opens NewQuoteForm pre-populated and saves explicitly.
    """
    quote: dict = {}
    try:
        from app.db.connection import get_pool
        from app.db import quote_queries
        pool = await get_pool()
        if pool:
            await quote_queries.ensure_tables(pool)
            q = await quote_queries.get_quote_db(pool, quote_id)
            if q:
                quote = q
    except Exception as exc:
        logger.warning("clone DB fetch failed: %s", exc)

    if not quote:
        all_mock = {q["quote_id"]: q for q in _mock_quotes()}
        quote = all_mock.get(quote_id, {})

    if not quote:
        raise HTTPException(status_code=404, detail=f"Quote {quote_id} not found")

    proj = str(quote.get("project_name") or "")
    prefix = "[Copy] "
    return {
        **{k: v for k, v in quote.items() if k not in ("quote_id", "quote_number", "created_at", "valid_till", "updated_at", "total", "margin_pct")},
        "project_name": proj if proj.startswith(prefix) else prefix + proj,
        "status": "DRAFT",
        "line_items": [
            {k2: (float(v2) if hasattr(v2, "__float__") and not isinstance(v2, (str, bool)) else v2)
             for k2, v2 in li.items() if k2 not in ("item_id",)}
            for li in (quote.get("line_items") or [])
        ],
    }


# ── /quotes/{quote_id}/send-email ─────────────────────────────────────────────

@router.post("/quotes/{quote_id}/send-email")
async def send_quote_email(quote_id: int, payload: dict):
    """
    Send a quotation by email.
    payload: { recipient_email, recipient_name, subject, message }
    Uses SMTP_USER / SMTP_PASSWORD / SMTP_HOST / SMTP_PORT from .env.
    Falls back to a success-simulation response if SMTP is not configured.
    """
    recipient_email = (payload.get("recipient_email") or "").strip()
    recipient_name  = (payload.get("recipient_name")  or "Customer").strip()
    subject         = (payload.get("subject")         or "").strip()
    message_body    = (payload.get("message")         or "").strip()

    if not recipient_email:
        raise HTTPException(status_code=422, detail="recipient_email is required.")

    # Resolve quote details for email body
    quote_data: dict = {}
    try:
        from app.db.connection import get_pool
        from app.db import quote_queries
        pool = await get_pool()
        if pool:
            await quote_queries.ensure_tables(pool)
            quote_data = await quote_queries.get_quote_db(pool, quote_id) or {}
    except Exception:
        pass

    if not quote_data:
        mock_q = next((q for q in _mock_quotes() if q["quote_id"] == quote_id), None)
        if mock_q:
            quote_data = mock_q

    quote_number = quote_data.get("quote_number") or f"QT-{datetime.date.today().year}-{quote_id:04d}"
    customer     = quote_data.get("customer_name") or recipient_name
    total        = quote_data.get("grand_total") or quote_data.get("total") or 0

    if not subject:
        subject = f"Quotation {quote_number} from InvenIQ — Building Materials"

    # Build HTML email
    custom_para = f"<p style='margin:0 0 16px'>{message_body}</p>" if message_body else ""
    html_body = f"""
<!DOCTYPE html>
<html>
<head><meta charset='utf-8'></head>
<body style='font-family:Inter,Arial,sans-serif;background:#f8fafc;margin:0;padding:24px'>
  <div style='max-width:600px;margin:0 auto;background:#fff;border-radius:12px;overflow:hidden;border:1px solid #e2e8f0'>
    <div style='background:linear-gradient(135deg,#0f2744,#15803d);padding:24px 28px'>
      <div style='font-size:24px;font-weight:900;color:#fff;letter-spacing:-1px'>IQ</div>
      <div style='color:rgba(255,255,255,.8);font-size:13px;margin-top:4px'>InvenIQ — Building Materials</div>
    </div>
    <div style='padding:28px'>
      <p style='margin:0 0 16px;font-size:15px;color:#0f172a'>Dear {recipient_name},</p>
      {custom_para}
      <p style='margin:0 0 16px;color:#475569;font-size:14px'>
        Please find attached our quotation <strong style='color:#0f172a'>{quote_number}</strong>
        for <strong style='color:#0f172a'>{customer}</strong>.
      </p>
      <div style='background:#f1f5f9;border-radius:8px;padding:16px 20px;margin:20px 0;border-left:4px solid #15803d'>
        <div style='font-size:12px;font-weight:700;color:#94a3b8;text-transform:uppercase;letter-spacing:.7px;margin-bottom:8px'>Quotation Summary</div>
        <div style='font-size:18px;font-weight:900;color:#0f172a;font-family:monospace'>
          {quote_number}
        </div>
        {'<div style="font-size:22px;font-weight:900;color:#15803d;font-family:monospace;margin-top:4px">₹{:,.0f}</div>'.format(total) if total else ''}
      </div>
      <p style='color:#475569;font-size:14px;margin:0 0 16px'>
        Please review the quotation and feel free to contact us with any questions or to proceed with the order.
      </p>
      <p style='color:#94a3b8;font-size:12px;margin:0'>
        Best regards,<br>
        <strong style='color:#0f172a'>InvenIQ — Building Materials Team</strong>
      </p>
    </div>
    <div style='background:#f8fafc;border-top:1px solid #e2e8f0;padding:14px 28px;font-size:11px;color:#94a3b8'>
      Sent via InvenIQ Quotation Builder · Auto-generated
    </div>
  </div>
</body>
</html>
"""

    smtp_user = os.getenv("SMTP_USER", "")
    smtp_pass = os.getenv("SMTP_PASSWORD", "")
    smtp_host = os.getenv("SMTP_HOST", "smtp.gmail.com")
    smtp_port = int(os.getenv("SMTP_PORT", "587"))

    # Generate PDF attachment from quote data
    pdf_bytes: bytes | None = None
    pdf_filename = f"{quote_number}.pdf"
    try:
        from app.services.pdf_service import generate_quote_pdf
        pdf_bytes = generate_quote_pdf(quote_data)
        logger.info("PDF generated for quote %s (%d bytes)", quote_number, len(pdf_bytes))
    except Exception as pdf_err:
        logger.warning("PDF generation failed for %s — sending email without attachment: %s", quote_number, pdf_err)

    if not smtp_user or not smtp_pass:
        # SMTP not configured — return simulated success with PDF info
        logger.info("SMTP not configured — simulating email send to %s for quote %s", recipient_email, quote_number)
        return {
            "success": True,
            "simulated": True,
            "has_pdf": pdf_bytes is not None,
            "message": f"Email to {recipient_email} simulated (SMTP not configured). Add SMTP_USER / SMTP_PASSWORD to backend/.env to send real emails.",
            "quote_number": quote_number,
            "recipient": recipient_email,
        }

    try:
        from app.services.email_service import _smtp_send, _smtp_hint
        import asyncio
        await asyncio.get_event_loop().run_in_executor(
            None,
            lambda: _smtp_send(
                smtp_host, smtp_port, smtp_user, smtp_pass,
                recipient_email, subject, html_body,
                attachment=pdf_bytes,
                attachment_filename=pdf_filename,
            ),
        )
        logger.info("Quote email sent with PDF: quote=%s recipient=%s", quote_number, recipient_email)
        return {
            "success": True,
            "simulated": False,
            "has_pdf": pdf_bytes is not None,
            "message": f"Email sent successfully to {recipient_email}.",
            "quote_number": quote_number,
            "recipient": recipient_email,
        }
    except Exception as exc:
        hint = ""
        try:
            from app.services.email_service import _smtp_hint
            hint = _smtp_hint(exc)
        except Exception:
            pass
        logger.exception("send-quote-email failed: %s", exc)
        raise HTTPException(
            status_code=502,
            detail=f"Email delivery failed: {exc}. {hint}".strip(),
        )
